import streamlit as st
import json
import os
import re
import uuid
import random
import time
import hashlib
import html as _html
import urllib.request
import urllib.error
import threading
from contextlib import nullcontext
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px

# ─── LLM Config ──────────────────────────────────────────────────────────────
_FORGE_LLM_API_KEY  = os.environ.get("FORGE_LLM_API_KEY", "")
_FORGE_LLM_BASE_URL = os.environ.get("FORGE_LLM_BASE_URL", "https://api.x.ai/v1").rstrip("/")
_FORGE_LLM_MODEL    = os.environ.get("FORGE_LLM_MODEL", "grok-3")
_LLM_AVAILABLE      = bool(_FORGE_LLM_API_KEY)   # True only when env var is set

# ─── Persistent key file (never committed — listed in .gitignore) ─────────────
_KEY_FILE = os.path.join(os.path.dirname(__file__), ".forge_api_key")


def _load_saved_key() -> str:
    """Read the persisted API key from disk. Returns '' if absent or unreadable."""
    try:
        with open(_KEY_FILE, "r", encoding="utf-8") as f:
            return f.read().strip()
    except (OSError, IOError):
        return ""


def _save_key_to_disk(key: str) -> None:
    """Write the API key to disk so it survives page refreshes."""
    try:
        with open(_KEY_FILE, "w", encoding="utf-8") as f:
            f.write(key.strip())
    except (OSError, IOError):
        pass   # silently ignore (e.g. read-only FS)


def _clear_saved_key() -> None:
    """Delete the persisted key file from disk."""
    try:
        os.remove(_KEY_FILE)
    except (OSError, FileNotFoundError):
        pass


def _effective_llm_key() -> str:
    """Return the API key to use: env var takes priority, then session-entered key."""
    if _FORGE_LLM_API_KEY:
        return _FORGE_LLM_API_KEY
    return st.session_state.get("session_llm_key", "").strip()


def _llm_ready() -> bool:
    """True when a usable API key is available (env var or session entry)."""
    return bool(_effective_llm_key())

# ─── Page Config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ForgeOS",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Design System & CSS ─────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ══ Reset & base ══════════════════════════════════════════ */
*, *::before, *::after { box-sizing: border-box; }
html, body, .stApp,
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stSidebar"] label,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] [data-baseweb="select"] {
    font-family: 'Inter', -apple-system, sans-serif !important;
}
.stApp { background: #0d1117; }

/* ══ Sidebar ════════════════════════════════════════════════ */
[data-testid="stSidebar"] {
    background: #010409 !important;
    border-right: 1px solid #21262d !important;
    min-width: 220px !important;
    max-width: 220px !important;
}
[data-testid="stSidebar"] > div { padding: 0 !important; }
[data-testid="stSidebarContent"] { padding: 0 !important; }

/* Hide default radio label */
[data-testid="stSidebar"] .stRadio > label { display: none !important; }
[data-testid="stSidebar"] .stRadio > div {
    display: flex !important;
    flex-direction: column !important;
    gap: 1px !important;
}
[data-testid="stSidebar"] .stRadio > div > label {
    display: flex !important;
    align-items: center !important;
    padding: 6px 12px 6px 16px !important;
    border-radius: 6px !important;
    margin: 1px 6px !important;
    cursor: pointer !important;
    font-size: 13px !important;
    font-weight: 500 !important;
    color: #b0b8c4 !important;
    transition: background 0.15s, color 0.15s !important;
    border: none !important;
}
[data-testid="stSidebar"] .stRadio > div > label:hover {
    background: #161b22 !important;
    color: #f0f6fc !important;
}
[data-testid="stSidebar"] .stRadio > div > label[data-baseweb="radio"] { }
[data-testid="stSidebar"] .stRadio [aria-checked="true"] ~ span,
[data-testid="stSidebar"] .stRadio input:checked ~ span {
    color: #e6edf3 !important;
}
[data-testid="stSidebar"] .stRadio > div > label > div:first-child {
    display: none !important;
}
[data-testid="stSidebar"] .stRadio > div > label > div:last-child > p {
    font-size: 13px !important;
    font-weight: 500 !important;
    color: inherit !important;
    margin: 0 !important;
}

/* ══ Main content padding ═══════════════════════════════════ */
.main .block-container {
    padding: 0 !important;
    max-width: 100% !important;
}

/* ══ Top bar ════════════════════════════════════════════════ */
.forge-topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 12px 28px;
    border-bottom: 1px solid #21262d;
    background: #0d1117;
    position: sticky;
    top: 0;
    z-index: 100;
}
.forge-topbar-left {
    display: flex;
    align-items: center;
    gap: 8px;
}
.forge-breadcrumb {
    font-size: 12px;
    color: #7d8490;
    font-weight: 500;
}
.forge-breadcrumb span {
    color: #f0f6fc;
    font-weight: 600;
    font-size: 14px;
}
.forge-topbar-actions {
    display: flex;
    align-items: center;
    gap: 8px;
}

/* ══ Page content wrapper ═══════════════════════════════════ */
.page-content { padding: 24px 28px; }

/* ══ Stat strip (Apollo-style) ══════════════════════════════ */
.stat-strip {
    display: flex;
    border: 1px solid #21262d;
    border-radius: 8px;
    overflow: hidden;
    background: #161b22;
    margin-bottom: 20px;
}
.stat-item {
    flex: 1;
    padding: 16px 20px;
    border-right: 1px solid #21262d;
    min-width: 0;
}
.stat-item:last-child { border-right: none; }
.stat-label {
    font-size: 11px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    color: #7d8490;
    margin-bottom: 6px;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}
.stat-value {
    font-size: 22px;
    font-weight: 700;
    color: #f0f6fc;
    line-height: 1;
    margin-bottom: 4px;
}
.stat-sub {
    font-size: 11px;
    color: #7d8490;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}
.stat-delta-up   { color: #3fb950; font-size: 11px; font-weight: 600; }
.stat-delta-down { color: #f85149; font-size: 11px; font-weight: 600; }

/* ══ Section header ══════════════════════════════════════════ */
.section-hd {
    font-size: 11px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: #8b949e;
    margin: 24px 0 12px 0;
    padding-bottom: 8px;
    border-bottom: 1px solid #21262d;
}

/* ══ Card ════════════════════════════════════════════════════ */
.forge-card {
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 8px;
    padding: 16px 20px;
    transition: border-color 0.15s;
}
.forge-card:hover { border-color: #6e7681; }

/* ══ Score badge ════════════════════════════════════════════ */
.badge-score {
    display: inline-flex;
    align-items: center;
    gap: 4px;
    padding: 2px 8px;
    border-radius: 4px;
    font-size: 12px;
    font-weight: 700;
    font-variant-numeric: tabular-nums;
}
.badge-green  { background: #0d2b1a; color: #3fb950; border: 1px solid #238636; }
.badge-yellow { background: #2b1f05; color: #d29922; border: 1px solid #9e6a03; }
.badge-red    { background: #2b0f0f; color: #f85149; border: 1px solid #6e1818; }

/* ══ Status pill ════════════════════════════════════════════ */
.pill {
    display: inline-block;
    padding: 2px 8px;
    border-radius: 9999px;
    font-size: 11px;
    font-weight: 600;
    letter-spacing: 0.02em;
    white-space: nowrap;
}
.pill-new      { background: #0c1e35; color: #58a6ff; border: 1px solid #1f6feb44; }
.pill-scored   { background: #0d2b1a; color: #3fb950; border: 1px solid #23863644; }
.pill-review   { background: #2b1f05; color: #d29922; border: 1px solid #9e6a0344; }
.pill-approved { background: #1b0f2e; color: #a371f7; border: 1px solid #6e40c944; }
.pill-rejected { background: #2b0f0f; color: #f85149; border: 1px solid #6e181844; }

/* ══ Table ═══════════════════════════════════════════════════ */
.forge-table { width: 100%; border-collapse: collapse; }
.forge-th {
    font-size: 11px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    color: #8b949e;
    padding: 8px 12px;
    text-align: left;
    border-bottom: 1px solid #21262d;
    white-space: nowrap;
}
.forge-tr {
    border-bottom: 1px solid #161b22;
    transition: background 0.1s;
}
.forge-tr:hover { background: #161b22; }
.forge-td {
    padding: 10px 12px;
    font-size: 13px;
    color: #b0b8c4;
    vertical-align: middle;
}
.forge-td-primary { color: #f0f6fc; font-weight: 500; }
.forge-id {
    font-size: 11px;
    font-weight: 600;
    color: #8b949e;
    font-family: 'SF Mono', monospace !important;
    background: #161b22;
    padding: 2px 6px;
    border-radius: 4px;
    border: 1px solid #6e7681;
}

/* ══ Kanban column ══════════════════════════════════════════ */
.kanban-col-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 8px 0 10px 0;
    margin-bottom: 10px;
    border-bottom: 1px solid #21262d;
}
.kanban-col-name {
    font-size: 12px;
    font-weight: 600;
    color: #b0b8c4;
    text-transform: uppercase;
    letter-spacing: 0.06em;
}
.kanban-count {
    font-size: 11px;
    font-weight: 600;
    color: #8b949e;
    background: #161b22;
    border: 1px solid #6e7681;
    border-radius: 9999px;
    padding: 1px 7px;
}
.kanban-card {
    background: #161b22;
    border: 1px solid #6e7681;
    border-radius: 6px;
    padding: 10px 12px;
    margin-bottom: 8px;
    transition: border-color 0.15s, box-shadow 0.15s;
    cursor: default;
}
.kanban-card:hover {
    border-color: #8b949e;
    box-shadow: 0 2px 12px #00000040;
}
.kanban-card-title {
    font-size: 12px;
    font-weight: 500;
    color: #f0f6fc;
    margin-bottom: 6px;
    line-height: 1.4;
}
.kanban-card-meta {
    font-size: 11px;
    color: #7d8490;
    display: flex;
    align-items: center;
    justify-content: space-between;
}
.kanban-empty {
    font-size: 12px;
    color: #8b949e;
    text-align: center;
    padding: 20px 8px;
    border: 1px dashed #6e7681;
    border-radius: 6px;
}

/* ══ Upload zone ════════════════════════════════════════════ */
[data-testid="stFileUploader"] {
    background: #161b22 !important;
    border: 1px dashed #6e7681 !important;
    border-radius: 8px !important;
    transition: border-color 0.15s !important;
}
[data-testid="stFileUploader"]:hover { border-color: #58a6ff !important; }
[data-testid="stFileUploader"] label { color: #8d96a3 !important; }

/* ══ Buttons ════════════════════════════════════════════════ */
.stButton > button {
    background: #21262d !important;
    color: #c9d1d9 !important;
    border: 1px solid #6e7681 !important;
    border-radius: 6px !important;
    font-size: 12px !important;
    font-weight: 600 !important;
    padding: 5px 14px !important;
    transition: background 0.15s, border-color 0.15s !important;
    white-space: nowrap !important;
}
.stButton > button:hover {
    background: #6e7681 !important;
    border-color: #8b949e !important;
    color: #e6edf3 !important;
}
.stButton > button[kind="primary"],
.stButton > button:first-child[data-primary="true"] {
    background: #1f6feb !important;
    border-color: #1f6feb !important;
    color: white !important;
}

/* ══ Inputs ════════════════════════════════════════════════ */
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] > div > div {
    background: #161b22 !important;
    border: 1px solid #6e7681 !important;
    border-radius: 6px !important;
    color: #c9d1d9 !important;
    font-size: 13px !important;
}
[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
    border-color: #1f6feb !important;
    box-shadow: 0 0 0 3px #1f6feb25 !important;
}

/* ══ Expander ═══════════════════════════════════════════════ */
[data-testid="stExpander"] {
    background: #161b22 !important;
    border: 1px solid #21262d !important;
    border-radius: 8px !important;
}
[data-testid="stExpander"] summary { color: #c9d1d9 !important; }

/* ══ Progress bar ═══════════════════════════════════════════ */
[data-testid="stProgress"] > div > div { border-radius: 2px !important; }
[data-testid="stProgress"] { border-radius: 2px !important; }

/* ══ Alerts ════════════════════════════════════════════════ */
[data-testid="stAlert"] {
    background: #161b22 !important;
    border: 1px solid #21262d !important;
    border-radius: 8px !important;
}

/* ══ Rubric criterion card ══════════════════════════════════ */
.crit-card {
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 8px;
    padding: 16px 18px;
    margin-bottom: 10px;
    transition: border-color 0.15s;
}
.crit-card:hover { border-color: #6e7681; }
.crit-name {
    font-size: 13px;
    font-weight: 600;
    color: #e6edf3;
    margin-bottom: 2px;
}
.crit-desc { font-size: 12px; color: #8d96a3; }
.crit-weight {
    font-size: 11px;
    font-weight: 700;
    padding: 2px 8px;
    border-radius: 4px;
    background: #0c1e35;
    color: #58a6ff;
    border: 1px solid #1f6feb33;
}
.anchor-band {
    background: #0d1117;
    border-radius: 6px;
    padding: 8px 12px;
    margin: 4px 0;
    font-size: 12px;
}
.anchor-low  { border-left: 3px solid #f85149; }
.anchor-mid  { border-left: 3px solid #d29922; }
.anchor-high { border-left: 3px solid #3fb950; }
.subfactor-tag {
    display: inline-block;
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 4px;
    padding: 2px 8px;
    font-size: 11px;
    color: #8d96a3;
    margin: 2px;
}
.redflag-item {
    display: flex;
    align-items: center;
    gap: 6px;
    font-size: 12px;
    color: #f85149;
    margin: 3px 0;
}

/* ══ Gating rule ════════════════════════════════════════════ */
.gate-rule {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 8px 12px;
    background: #2b0f0f18;
    border: 1px solid #6e181822;
    border-radius: 6px;
    margin-bottom: 6px;
    font-size: 12px;
    color: #e6edf3;
}

/* ══ Sidebar custom HTML ════════════════════════════════════ */
.sb-logo {
    padding: 16px 14px 12px 14px;
    border-bottom: 1px solid #21262d;
    margin-bottom: 8px;
}
.sb-wordmark {
    font-size: 15px;
    font-weight: 800;
    color: #e6edf3;
    letter-spacing: -0.03em;
    display: flex;
    align-items: center;
    gap: 8px;
}
.sb-tag {
    font-size: 9px;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    color: #6e7681;
    margin-top: 2px;
}
.sb-section-label {
    font-size: 10px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    color: #6e7681;
    padding: 12px 16px 4px 16px;
}
.sb-divider { border-top: 1px solid #21262d; margin: 8px 0; }
.sb-stat-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 4px 16px;
    font-size: 12px;
}
.sb-stat-label { color: #8b949e; }
.sb-stat-val   { color: #b0b8c4; font-weight: 600; }

/* ══ Empty state ════════════════════════════════════════════ */
.empty-state {
    text-align: center;
    padding: 60px 24px;
}
.empty-icon  { font-size: 32px; margin-bottom: 12px; opacity: 0.3; }
.empty-title { font-size: 15px; font-weight: 600; color: #8b949e; margin-bottom: 6px; }
.empty-sub   { font-size: 13px; color: #8b949e; }

/* ══ Overrides ══════════════════════════════════════════════ */
hr { border-color: #6e7681 !important; margin: 12px 0 !important; }
[data-testid="stMarkdownContainer"] p { font-size: 13px; color: #b0b8c4; }
</style>
""", unsafe_allow_html=True)

# ─── Design System v2 — Premium Redesign ─────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Material+Symbols+Rounded:opsz,wght,FILL,GRAD@24,400,0,0&display=swap');

/* ══ Design Tokens ══════════════════════════════════════════════ */
:root {
  --bg:            #07101f;
  --surface:       #0b1626;
  --card:          #0f1e30;
  --card-hover:    #132438;
  --border:        #1a2f49;
  --border-soft:   #111e2f;
  --text-1:        #ddeaf8;
  --text-2:        #8aa4c0;
  --text-3:        #4e6680;
  --accent:        #3b82f6;
  --accent-dim:    rgba(59,130,246,0.10);
  --accent-glow:   rgba(59,130,246,0.20);
  --green:         #22c55e;
  --green-bg:      rgba(34,197,94,0.09);
  --green-border:  rgba(34,197,94,0.24);
  --amber:         #f59e0b;
  --amber-bg:      rgba(245,158,11,0.09);
  --amber-border:  rgba(245,158,11,0.24);
  --red:           #f43f5e;
  --red-bg:        rgba(244,63,94,0.09);
  --red-border:    rgba(244,63,94,0.24);
  --purple:        #a78bfa;
  --purple-bg:     rgba(167,139,250,0.09);
  --purple-border: rgba(167,139,250,0.24);
  --shadow-sm: 0 1px 5px rgba(0,0,16,0.55), 0 0 0 1px rgba(255,255,255,0.025);
  --shadow-md: 0 4px 20px rgba(0,0,16,0.60), 0 0 0 1px rgba(255,255,255,0.038);
  --shadow-lg: 0 10px 40px rgba(0,0,16,0.70), 0 0 0 1px rgba(255,255,255,0.05);
  --r-sm: 6px; --r-md: 10px; --r-lg: 14px;
  --ease: cubic-bezier(0.4,0,0.2,1); --dur: 140ms;
}

/* ══ Base overrides ═════════════════════════════════════════════ */
*, *::before, *::after { box-sizing: border-box; }
html, body, .stApp,
[data-testid="stAppViewContainer"],
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stSidebar"] label,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] [data-baseweb="select"],
.stButton > button {
  font-family: 'Inter', -apple-system, sans-serif !important;
  -webkit-font-smoothing: antialiased !important;
}
.stApp { background: var(--bg) !important; }

/* Streamlit Material Symbols — must not inherit Inter (shows "arrowright" text) */
[data-testid="stIconMaterial"],
[data-testid="stIconMaterial"] span {
  font-family: "Material Symbols Rounded" !important;
  font-weight: normal !important;
  font-style: normal !important;
  letter-spacing: normal !important;
  text-transform: none !important;
  line-height: 1 !important;
  -webkit-font-feature-settings: "liga" !important;
  font-feature-settings: "liga" !important;
}
[data-testid="stExpander"] details summary [data-testid="stIconMaterial"],
[data-testid="stSelectbox"] [data-testid="stIconMaterial"],
[data-testid="stDateInput"] [data-testid="stIconMaterial"] {
  font-family: "Material Symbols Rounded" !important;
}

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 99px; }
::-webkit-scrollbar-thumb:hover { background: var(--text-3); }

/* ══ Sidebar ════════════════════════════════════════════════════ */
[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #060f1e 0%, #050c18 100%) !important;
  border-right: 1px solid var(--border) !important;
  min-width: 220px !important; max-width: 220px !important;
}
[data-testid="stSidebar"] > div { padding: 0 !important; }
[data-testid="stSidebarContent"] { padding: 0 !important; }
[data-testid="stSidebar"] .stRadio > label { display: none !important; }
[data-testid="stSidebar"] .stRadio > div {
  display: flex !important; flex-direction: column !important;
  gap: 1px !important; padding: 0 8px !important;
}
[data-testid="stSidebar"] .stRadio > div > label {
  display: flex !important; align-items: center !important;
  padding: 7px 10px 7px 12px !important; border-radius: var(--r-sm) !important;
  margin: 1px 0 !important; cursor: pointer !important;
  font-size: 13px !important; font-weight: 500 !important;
  color: var(--text-2) !important;
  transition: all var(--dur) var(--ease) !important;
  border: none !important; border-left: 2px solid transparent !important;
}
[data-testid="stSidebar"] .stRadio > div > label:hover {
  background: rgba(59,130,246,0.08) !important;
  color: var(--text-1) !important;
  border-left-color: rgba(59,130,246,0.4) !important;
}
[data-testid="stSidebar"] .stRadio > div > label > div:first-child { display: none !important; }
[data-testid="stSidebar"] .stRadio > div > label > div:last-child > p {
  font-size: 13px !important; font-weight: 500 !important;
  color: inherit !important; margin: 0 !important;
}

/* ══ Top bar ════════════════════════════════════════════════════ */
.forge-topbar {
  display: flex; align-items: center; justify-content: space-between;
  padding: 0 28px; height: 50px;
  border-bottom: 1px solid var(--border);
  background: rgba(7,16,31,0.93);
  backdrop-filter: blur(20px); -webkit-backdrop-filter: blur(20px);
  position: sticky; top: 0; z-index: 100;
}
.forge-topbar-left { display: flex; align-items: center; gap: 12px; }
.forge-breadcrumb {
  font-size: 12px; color: var(--text-3); font-weight: 500;
  display: flex; align-items: center; gap: 7px;
}
.forge-breadcrumb span { color: var(--text-1); font-weight: 600; font-size: 14px; }
.forge-sep { color: var(--border); }
.forge-page-tag {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.09em; color: var(--accent);
  background: var(--accent-dim); border: 1px solid rgba(59,130,246,0.22);
  padding: 2px 9px; border-radius: 99px;
}
.forge-topbar-status {
  display: flex; align-items: center; gap: 6px;
  font-size: 11px; color: var(--text-3); font-weight: 500;
}
.forge-status-dot {
  width: 6px; height: 6px; border-radius: 50%;
  background: var(--green); box-shadow: 0 0 7px var(--green);
  animation: pulse-dot 2.5s ease-in-out infinite;
}
@keyframes pulse-dot { 0%,100%{opacity:1} 50%{opacity:0.35} }

/* ══ Page content ═══════════════════════════════════════════════ */
.page-content { padding: 24px 28px 48px !important; }

/* ══ KPI card grid ══════════════════════════════════════════════ */
.kpi-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(152px, 1fr));
  gap: 12px; margin-bottom: 28px;
}
.kpi-card {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-md); padding: 18px 18px 16px;
  box-shadow: var(--shadow-sm); transition: all var(--dur) var(--ease);
  position: relative; overflow: hidden; cursor: default;
}
.kpi-card::after {
  content: ''; position: absolute; top: 0; left: 0; right: 0; height: 1px;
  background: linear-gradient(90deg, transparent, rgba(59,130,246,0.55), transparent);
  opacity: 0; transition: opacity var(--dur) var(--ease);
}
.kpi-card:hover {
  border-color: rgba(59,130,246,0.3); transform: translateY(-2px);
  box-shadow: var(--shadow-md);
}
.kpi-card:hover::after { opacity: 1; }
.kpi-icon { font-size: 18px; margin-bottom: 10px; display: block; }
.kpi-label {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.09em; color: var(--text-3); margin-bottom: 8px;
}
.kpi-value {
  font-size: 30px; font-weight: 800; color: var(--text-1); line-height: 1;
  margin-bottom: 6px; font-variant-numeric: tabular-nums; letter-spacing: -0.03em;
}
.kpi-sub { font-size: 11px; color: var(--text-3); line-height: 1.4; }

/* ══ Stat strip (legacy fallback) ═══════════════════════════════ */
.stat-strip {
  border: 1px solid var(--border) !important; border-radius: var(--r-md) !important;
  background: var(--card) !important; box-shadow: var(--shadow-sm) !important;
  overflow: hidden !important;
}
.stat-item {
  padding: 18px 20px !important; border-right: 1px solid var(--border) !important;
  transition: background var(--dur) var(--ease) !important;
}
.stat-item:hover { background: var(--card-hover) !important; }
.stat-label {
  font-size: 10px !important; font-weight: 700 !important;
  text-transform: uppercase !important; letter-spacing: 0.09em !important;
  color: var(--text-3) !important; margin-bottom: 8px !important;
}
.stat-value {
  font-size: 26px !important; font-weight: 800 !important;
  color: var(--text-1) !important; letter-spacing: -0.025em !important;
  line-height: 1 !important; margin-bottom: 4px !important;
}
.stat-sub { font-size: 11px !important; color: var(--text-3) !important; }

/* ══ Section header ═════════════════════════════════════════════ */
.section-hd {
  font-size: 10px !important; font-weight: 700 !important;
  text-transform: uppercase !important; letter-spacing: 0.10em !important;
  color: var(--text-3) !important;
  margin: 28px 0 14px !important; padding-bottom: 10px !important;
  border-bottom: 1px solid var(--border-soft) !important;
  display: flex !important; align-items: center !important; gap: 10px !important;
}
.section-hd::before {
  content: ''; display: block; width: 3px; height: 12px;
  background: linear-gradient(180deg, var(--accent) 0%, #6366f1 100%);
  border-radius: 2px; flex-shrink: 0;
}

/* ══ Forge card ═════════════════════════════════════════════════ */
.forge-card {
  background: var(--card) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r-md) !important; box-shadow: var(--shadow-sm) !important;
  transition: all var(--dur) var(--ease) !important;
}
.forge-card:hover {
  border-color: rgba(59,130,246,0.3) !important; box-shadow: var(--shadow-md) !important;
}

/* ══ Score badge ════════════════════════════════════════════════ */
.badge-score {
  display: inline-flex !important; align-items: center !important; gap: 4px !important;
  padding: 3px 9px !important; border-radius: var(--r-sm) !important;
  font-size: 12px !important; font-weight: 700 !important;
  font-variant-numeric: tabular-nums !important; letter-spacing: -0.01em !important;
}
.badge-green  { background: var(--green-bg)  !important; color: var(--green)  !important; border: 1px solid var(--green-border)  !important; }
.badge-yellow { background: var(--amber-bg)  !important; color: var(--amber)  !important; border: 1px solid var(--amber-border)  !important; }
.badge-red    { background: var(--red-bg)    !important; color: var(--red)    !important; border: 1px solid var(--red-border)    !important; }

/* ══ Status pill ════════════════════════════════════════════════ */
.pill {
  display: inline-block !important; padding: 3px 9px !important;
  border-radius: 99px !important; font-size: 11px !important;
  font-weight: 600 !important; letter-spacing: 0.02em !important;
}
.pill-new      { background: var(--accent-dim)  !important; color: #60a5fa     !important; border: 1px solid rgba(59,130,246,0.25)  !important; }
.pill-scored   { background: var(--green-bg)    !important; color: var(--green) !important; border: 1px solid var(--green-border)    !important; }
.pill-review   { background: var(--amber-bg)    !important; color: var(--amber) !important; border: 1px solid var(--amber-border)    !important; }
.pill-approved { background: var(--purple-bg)   !important; color: var(--purple)!important; border: 1px solid var(--purple-border)   !important; }
.pill-rejected { background: var(--red-bg)      !important; color: var(--red)   !important; border: 1px solid var(--red-border)      !important; }

/* ══ Table ══════════════════════════════════════════════════════ */
.forge-table { width: 100%; border-collapse: collapse; }
.forge-th {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.09em; color: var(--text-3); padding: 10px 14px;
  text-align: left; border-bottom: 1px solid var(--border);
  white-space: nowrap; background: rgba(7,16,31,0.55);
}
.forge-tr { border-bottom: 1px solid var(--border-soft); transition: background var(--dur) var(--ease); }
.forge-tr:last-child { border-bottom: none; }
.forge-tr:hover { background: rgba(59,130,246,0.05) !important; }
.forge-tr:nth-child(even) { background: rgba(255,255,255,0.012); }
.forge-td { padding: 11px 14px; font-size: 13px; color: var(--text-2); vertical-align: middle; }
.forge-td-primary { color: var(--text-1); font-weight: 500; }
.forge-id {
  font-size: 10px; font-weight: 700; color: var(--text-3);
  font-family: 'SF Mono','Fira Code','Consolas',monospace !important;
  background: var(--surface); padding: 3px 7px; border-radius: 4px;
  border: 1px solid var(--border); letter-spacing: 0.04em;
}

/* ══ Kanban ═════════════════════════════════════════════════════ */
.kanban-col-header {
  display: flex; align-items: center; justify-content: space-between;
  padding: 0 0 11px; margin-bottom: 12px; border-bottom: 1px solid var(--border-soft);
}
.kanban-col-name {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.10em; color: var(--text-2); display: flex; align-items: center; gap: 6px;
}
.kanban-count {
  font-size: 10px; font-weight: 700; color: var(--text-3);
  background: var(--surface); border: 1px solid var(--border);
  border-radius: 99px; padding: 1px 8px; min-width: 22px; text-align: center;
  font-variant-numeric: tabular-nums;
}
.kanban-card {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-sm); padding: 11px 13px; margin-bottom: 7px;
  box-shadow: var(--shadow-sm); transition: all var(--dur) var(--ease);
  cursor: pointer; border-left-width: 3px; position: relative; overflow: hidden;
  display: block; text-decoration: none !important;
}
.kanban-card:hover {
  border-color: rgba(59,130,246,0.42); box-shadow: 0 14px 32px rgba(0,0,0,0.34);
  transform: translateY(-4px);
}
.kanban-card::after {
  content: '';
  position: absolute;
  inset: 0;
  background: linear-gradient(135deg, rgba(59,130,246,0.10) 0%, rgba(59,130,246,0.0) 55%);
  opacity: 0;
  transition: opacity var(--dur) var(--ease);
  pointer-events: none;
}
.kanban-card:hover::after {
  opacity: 1;
}
div[data-testid="stMarkdown"]:has(.kanban-card) + div[data-testid="stButton"] {
  margin-top: -100px !important;
  margin-bottom: 29px !important;
  position: relative;
  z-index: 4;
}
div[data-testid="stMarkdown"]:has(.kanban-card) + div[data-testid="stButton"] > button {
  min-height: 100px !important;
  height: 100px !important;
  opacity: 0 !important;
  background: transparent !important;
  border: none !important;
  box-shadow: none !important;
  cursor: pointer !important;
}
div[data-testid="stMarkdown"]:has(.kanban-card) + div[data-testid="stButton"] > button:hover {
  transform: none !important;
  box-shadow: none !important;
}
.kanban-card-title {
  font-size: 12px; font-weight: 500; color: var(--text-1);
  margin-bottom: 8px; line-height: 1.45;
}
.kanban-card-meta {
  font-size: 11px; color: var(--text-3);
  display: flex; align-items: center; flex-wrap: wrap; gap: 4px;
}
.kanban-empty {
  font-size: 12px; color: var(--text-3); text-align: center;
  padding: 22px 8px; border: 1px dashed var(--border);
  border-radius: var(--r-sm); background: rgba(255,255,255,0.01);
}
.pipeline-detail-shell {
  background: linear-gradient(180deg, rgba(15,23,36,0.96) 0%, rgba(12,19,31,0.96) 100%);
  border: 1px solid var(--border);
  border-radius: var(--r-lg);
  padding: 20px 22px;
  box-shadow: var(--shadow-md);
  margin-bottom: 18px;
}
.pipeline-detail-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 10px;
  align-items: center;
  margin-top: 10px;
}
.pipeline-detail-pill {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: 999px;
  padding: 5px 10px;
  font-size: 11px;
  color: var(--text-2);
  font-weight: 600;
}
.profile-stat-card {
  background: var(--card);
  border: 1px solid var(--border);
  border-radius: var(--r-md);
  padding: 14px 16px;
  box-shadow: var(--shadow-sm);
  min-height: 110px;
}
.profile-stat-label {
  font-size: 10px;
  font-weight: 700;
  color: var(--text-3);
  text-transform: uppercase;
  letter-spacing: 0.08em;
}
.profile-stat-value {
  margin-top: 8px;
  font-size: 22px;
  font-weight: 800;
  color: var(--text-1);
}
.profile-stat-sub {
  margin-top: 6px;
  font-size: 12px;
  color: var(--text-3);
  line-height: 1.5;
}
.timeline-card {
  background: var(--card);
  border: 1px solid var(--border);
  border-radius: var(--r-md);
  padding: 14px 16px;
  box-shadow: var(--shadow-sm);
}
.timeline-step {
  display: flex;
  gap: 12px;
  align-items: flex-start;
  padding: 10px 0;
  border-bottom: 1px solid rgba(255,255,255,0.04);
}
.timeline-step:last-child {
  border-bottom: none;
  padding-bottom: 0;
}
.timeline-dot {
  width: 12px;
  height: 12px;
  border-radius: 999px;
  flex-shrink: 0;
  margin-top: 4px;
  box-shadow: 0 0 0 3px rgba(255,255,255,0.03);
}
.timeline-step-title {
  font-size: 13px;
  font-weight: 600;
  color: var(--text-1);
}
.timeline-step-date {
  font-size: 11px;
  color: var(--text-3);
  margin-top: 3px;
}

/* ══ Stage flow card (dashboard pipeline) ═══════════════════════ */
.stage-flow-card {
  background: var(--card); border: 1px solid var(--border);
  border-top-width: 2px; border-radius: var(--r-md);
  padding: 14px 10px 13px; text-align: center;
  box-shadow: var(--shadow-sm); transition: all var(--dur) var(--ease);
}
.stage-flow-card:hover { transform: translateY(-2px); box-shadow: var(--shadow-md); }
.stage-flow-icon { font-size: 20px; margin-bottom: 6px; display: block; }
.stage-flow-num {
  font-size: 28px; font-weight: 900; line-height: 1;
  margin-bottom: 5px; letter-spacing: -0.04em;
  font-variant-numeric: tabular-nums;
}
.stage-flow-name {
  font-size: 9px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.09em; color: var(--text-3);
}

/* ══ Upload zone ════════════════════════════════════════════════ */
[data-testid="stFileUploader"] {
  background: var(--card) !important;
  border: 1.5px dashed var(--border) !important;
  border-radius: var(--r-md) !important;
  transition: all var(--dur) var(--ease) !important;
}
[data-testid="stFileUploader"]:hover {
  border-color: var(--accent) !important;
  background: rgba(59,130,246,0.03) !important;
}

/* ══ Buttons ════════════════════════════════════════════════════ */
.stButton > button {
  background: var(--surface) !important; color: var(--text-2) !important;
  border: 1px solid var(--border) !important; border-radius: var(--r-sm) !important;
  font-size: 12px !important; font-weight: 600 !important; padding: 6px 14px !important;
  transition: all var(--dur) var(--ease) !important; letter-spacing: 0.01em !important;
}
.stButton > button:hover {
  background: var(--card-hover) !important; border-color: rgba(59,130,246,0.38) !important;
  color: var(--text-1) !important; transform: translateY(-1px) !important;
  box-shadow: 0 2px 10px rgba(0,0,0,0.35) !important;
}
.stButton > button[kind="primary"] {
  background: linear-gradient(135deg,#3b82f6 0%,#2563eb 100%) !important;
  border-color: rgba(59,130,246,0.55) !important; color: #fff !important;
  box-shadow: 0 2px 10px rgba(59,130,246,0.30) !important;
}
.stButton > button[kind="primary"]:hover {
  box-shadow: 0 4px 22px rgba(59,130,246,0.48) !important;
  transform: translateY(-1px) !important;
}

/* ══ Inputs ═════════════════════════════════════════════════════ */
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stSelectbox"] > div > div,
[data-testid="stNumberInput"] input {
  background: var(--surface) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r-sm) !important; color: var(--text-1) !important;
  font-size: 13px !important; transition: all var(--dur) var(--ease) !important;
}
[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(59,130,246,0.15) !important; outline: none !important;
}

/* ══ Expander ═══════════════════════════════════════════════════ */
[data-testid="stExpander"] {
  background: var(--card) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r-md) !important; box-shadow: var(--shadow-sm) !important;
  transition: border-color var(--dur) var(--ease) !important;
}
[data-testid="stExpander"]:hover { border-color: rgba(59,130,246,0.28) !important; }
[data-testid="stExpander"] summary {
  color: var(--text-2) !important; padding: 13px 18px !important;
}

/* ══ Progress & alerts ══════════════════════════════════════════ */
[data-testid="stProgress"] > div > div { border-radius: 99px !important; }
[data-testid="stAlert"] {
  background: var(--card) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r-md) !important;
}

/* ══ Rubric criterion card ══════════════════════════════════════ */
.crit-card {
  background: var(--card) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r-md) !important; box-shadow: var(--shadow-sm) !important;
  transition: all var(--dur) var(--ease) !important;
}
.crit-card:hover { border-color: rgba(59,130,246,0.30) !important; box-shadow: var(--shadow-md) !important; }
.crit-breakdown-row {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-md); padding: 14px 16px; margin-bottom: 10px;
  box-shadow: var(--shadow-sm);
}
.crit-breakdown-row:hover { border-color: rgba(59,130,246,0.28); }
.crit-detail-md h2 { font-size: 15px !important; color: var(--text-1) !important; margin: 1.1em 0 0.45em !important; }
.crit-detail-md h3 { font-size: 13px !important; color: var(--accent) !important; margin: 0.9em 0 0.35em !important; }
.crit-detail-md p, .crit-detail-md li { font-size: 13px !important; color: var(--text-2) !important; line-height: 1.65 !important; }
.crit-detail-md blockquote {
  border-left: 3px solid var(--accent); margin: 8px 0; padding: 6px 12px;
  background: rgba(59,130,246,0.06); color: var(--text-2) !important; font-size: 12px !important;
}
.crit-name { font-size: 13px !important; font-weight: 600 !important; color: var(--text-1) !important; }
.crit-desc { font-size: 12px !important; color: var(--text-3) !important; line-height: 1.55 !important; }
.crit-weight {
  background: var(--accent-dim) !important; color: var(--accent) !important;
  border: 1px solid rgba(59,130,246,0.25) !important;
}
.anchor-band {
  background: var(--surface) !important; border-radius: var(--r-sm) !important;
  padding: 8px 12px !important; color: var(--text-2) !important; font-size: 12px !important;
}
.anchor-low  { border-left: 3px solid var(--red)   !important; }
.anchor-mid  { border-left: 3px solid var(--amber) !important; }
.anchor-high { border-left: 3px solid var(--green) !important; }
.subfactor-tag {
  background: var(--surface) !important; border: 1px solid var(--border) !important;
  color: var(--text-3) !important;
}
.gate-rule {
  background: var(--red-bg) !important; border: 1px solid var(--red-border) !important;
  border-radius: var(--r-sm) !important;
}

/* ══ Sidebar custom HTML ════════════════════════════════════════ */
.sb-logo { padding: 16px 14px 14px; border-bottom: 1px solid var(--border); margin-bottom: 4px; }
.sb-wordmark {
  font-size: 15px; font-weight: 800; color: var(--text-1);
  letter-spacing: -0.04em; display: flex; align-items: center; gap: 9px;
}
.sb-logo-icon {
  width: 28px; height: 28px; flex-shrink: 0;
  background: linear-gradient(135deg, #3b82f6 0%, #6366f1 100%);
  border-radius: 8px; display: flex; align-items: center; justify-content: center;
  font-size: 15px; box-shadow: 0 2px 14px rgba(59,130,246,0.40);
}
.sb-tag { font-size: 9px; letter-spacing: 0.07em; color: var(--text-3); margin-top: 2px; font-weight: 500; text-transform: uppercase; }
.sb-section-label {
  font-size: 9px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.13em; color: var(--text-3); padding: 14px 14px 5px;
}
.sb-divider { border-top: 1px solid var(--border); margin: 8px 0; }
.sb-stat-row { display: flex; justify-content: space-between; align-items: center; padding: 4px 14px; font-size: 12px; }
.sb-stat-label { color: var(--text-3); }
.sb-stat-val   { color: var(--text-2); font-weight: 600; font-variant-numeric: tabular-nums; }

/* ══ Empty state ════════════════════════════════════════════════ */
.empty-state { text-align: center; padding: 72px 24px; }
.empty-icon  { font-size: 40px; margin-bottom: 16px; opacity: 0.13; display: block; }
.empty-title { font-size: 16px; font-weight: 600; color: var(--text-2); margin-bottom: 8px; }
.empty-sub   { font-size: 13px; color: var(--text-3); line-height: 1.65; }

/* ══ Misc overrides ═════════════════════════════════════════════ */
hr { border-color: var(--border) !important; margin: 14px 0 !important; }
[data-testid="stMarkdownContainer"] p { font-size: 13px; color: var(--text-2); line-height: 1.6; }
[data-testid="stCheckbox"] span { color: var(--text-2) !important; font-size: 13px !important; }
[data-testid="stFileUploader"] label { color: var(--text-3) !important; }

/* ══ Submissions table rows ═══════════════════════════════════════ */
.sub-table-header {
  padding: 8px 4px 10px; margin-bottom: 4px;
  border-bottom: 1px solid var(--border);
}
.sub-table-th {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.09em; color: var(--text-3);
}
.sub-row-card {
  background: var(--card); border: 1px solid var(--border-soft);
  border-radius: var(--r-sm); padding: 10px 8px 8px; margin-bottom: 8px;
}
.sub-row-card:hover { border-color: var(--border); }
.sub-cell-empty { color: var(--text-3); font-size: 13px; }
.sub-actions [data-testid="column"] { padding: 0 2px !important; }
.sub-actions .stButton > button {
  width: 100% !important; min-height: 32px !important; padding: 4px 6px !important;
  font-size: 11px !important;
}
[data-testid="stVerticalBlockBorderWrapper"] {
  margin-bottom: 10px !important;
  background: var(--card) !important;
  border-color: var(--border-soft) !important;
  border-radius: var(--r-sm) !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:hover {
  border-color: var(--border) !important;
}

/* ══ Investment memo preview ══════════════════════════════════════ */
.memo-preview-wrap {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--r-md);
  padding: 20px 24px;
  max-height: 52vh;
  overflow-y: auto;
  margin: 12px 0 16px 0;
  box-shadow: var(--shadow-sm);
}
.memo-preview-wrap h1 { font-size: 20px !important; color: var(--text-1) !important; margin-top: 0 !important; }
.memo-preview-wrap h2 { font-size: 15px !important; color: var(--text-1) !important; margin-top: 1.2em !important; }
.memo-preview-wrap h3 { font-size: 13px !important; color: var(--text-2) !important; }
.memo-preview-wrap p, .memo-preview-wrap li { font-size: 13px !important; color: var(--text-2) !important; line-height: 1.65 !important; }
.memo-preview-wrap strong { color: var(--text-1) !important; }
.memo-source-tag {
  font-size: 10px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.08em; color: var(--text-3); margin-bottom: 8px;
}

/* ══ Shortlist folder cards ═══════════════════════════════════════ */
.shortlist-folder-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
  gap: 14px;
  margin-bottom: 8px;
}
.shortlist-folder-card {
  background: var(--card);
  border: 1px solid var(--border);
  border-radius: var(--r-md);
  padding: 18px 16px 16px;
  box-shadow: var(--shadow-sm);
  transition: all var(--dur) var(--ease);
  cursor: default;
  position: relative;
  overflow: hidden;
}
.shortlist-folder-card::before {
  content: '';
  position: absolute;
  top: 0; left: 0; right: 0;
  height: 3px;
  background: var(--shortlist-accent, var(--accent));
  opacity: 0.85;
}
.shortlist-folder-card:hover {
  border-color: rgba(59,130,246,0.35);
  transform: translateY(-2px);
  box-shadow: var(--shadow-md);
}
.shortlist-folder-icon {
  font-size: 26px;
  line-height: 1;
  margin-bottom: 10px;
  display: block;
}
.shortlist-folder-name {
  font-size: 13px;
  font-weight: 600;
  color: var(--text-1);
  line-height: 1.35;
  margin-bottom: 10px;
  min-height: 2.6em;
}
.shortlist-folder-count {
  font-size: 28px;
  font-weight: 800;
  color: var(--text-1);
  letter-spacing: -0.03em;
  font-variant-numeric: tabular-nums;
  line-height: 1;
}
.shortlist-folder-count-label {
  font-size: 10px;
  font-weight: 600;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  color: var(--text-3);
  margin-top: 4px;
}
.shortlist-back-link {
  font-size: 12px;
  color: var(--accent);
  font-weight: 600;
  margin-bottom: 12px;
  display: inline-block;
}

/* ══ Landing page (Home) ════════════════════════════════════════════ */
.landing-page-bg {
  background:
    radial-gradient(ellipse 80% 50% at 50% -20%, rgba(59,130,246,0.22) 0%, transparent 55%),
    radial-gradient(ellipse 60% 40% at 100% 20%, rgba(34,197,94,0.12) 0%, transparent 50%),
    radial-gradient(ellipse 50% 30% at 0% 60%, rgba(99,102,241,0.10) 0%, transparent 45%),
    var(--bg) !important;
}
.landing-nav {
  display: flex; align-items: center; justify-content: space-between;
  padding: 18px 40px; border-bottom: 1px solid var(--border-soft);
  background: rgba(7,16,31,0.85); backdrop-filter: blur(16px);
  position: sticky; top: 0; z-index: 200;
}
.landing-nav-brand {
  display: flex; align-items: center; gap: 10px;
  font-size: 18px; font-weight: 800; color: var(--text-1);
  letter-spacing: -0.04em;
}
.landing-nav-badge {
  font-size: 9px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.1em; color: var(--green);
  background: var(--green-bg); border: 1px solid var(--green-border);
  padding: 3px 8px; border-radius: 99px;
}
.landing-hero-section {
  position: relative; overflow: hidden;
  max-width: 1320px; margin: 0 auto;
  padding: 20px 48px 16px;
}
.landing-hero-glow {
  position: absolute; border-radius: 50%; pointer-events: none; filter: blur(80px);
}
.landing-hero-glow--blue {
  width: 520px; height: 520px; top: -120px; left: -80px;
  background: radial-gradient(circle, rgba(59,130,246,0.28) 0%, transparent 70%);
}
.landing-hero-glow--green {
  width: 440px; height: 440px; top: 40px; right: -60px;
  background: radial-gradient(circle, rgba(34,197,94,0.20) 0%, transparent 70%);
}
.landing-hero-glow--purple {
  width: 320px; height: 320px; bottom: -40px; left: 35%;
  background: radial-gradient(circle, rgba(99,102,241,0.14) 0%, transparent 70%);
}
.landing-hero-inner {
  position: relative; z-index: 2;
  display: grid; grid-template-columns: 1fr 1.15fr;
  gap: 40px; align-items: center;
}
.landing-hero-copy { max-width: 560px; }
.landing-hero-visual {
  position: relative; padding: 8px 0 0;
}
.landing-hero-cta-bar {
  max-width: 1280px; margin: 0 auto;
  padding: 12px 40px 28px;
  display: flex; flex-direction: column; align-items: center;
  gap: 12px; text-align: center;
}
.landing-cta-btn {
  display: inline-block; padding: 12px 20px; border-radius: 10px;
  border: 1px solid var(--border); color: var(--text-1) !important;
  font-weight: 700; font-size: 14px; text-decoration: none !important;
  transition: all 0.2s var(--ease); cursor: pointer;
  background: transparent;
}
.landing-cta-btn:hover {
  border-color: rgba(59,130,246,0.45) !important;
  background: rgba(59,130,246,0.08) !important;
  transform: translateY(-1px);
}
.landing-mock-wrap { position: relative; width: 100%; }
.landing-hero-visual::before {
  content: ""; position: absolute; inset: -20px -10px -10px -30px;
  background: radial-gradient(ellipse at center, rgba(59,130,246,0.12) 0%, transparent 65%);
  pointer-events: none; z-index: 0;
}
.landing-hero-visual > * { position: relative; z-index: 1; }
.landing-hero-actions {
  position: relative; z-index: 2;
  max-width: 1320px; margin: 0 auto;
  padding: 0 48px 28px;
}
.landing-hero-actions-btn-wrap {
  display: flex; flex-direction: column; align-items: center;
  gap: 12px; text-align: center; width: 100%;
}
section.main:has(.landing-nav) .landing-hero-actions [data-testid="column"]:last-child {
  display: flex !important; justify-content: center !important;
}
.landing-hero-actions-btn-wrap .stButton,
.landing-hero-actions-btn-wrap [data-testid="stButton"] {
  margin: 0 auto !important; width: auto !important;
}
.landing-hero-actions-btn-wrap .stButton > button,
.landing-hero-actions-btn-wrap [data-testid="stButton"] button {
  width: auto !important;
}
.landing-hero-actions-row {
  display: grid; grid-template-columns: 1.25fr 1fr; gap: 12px;
}
.landing-hero-micro {
  font-size: 11px; color: var(--text-3); margin: 0; line-height: 1.5;
  max-width: 420px; text-align: center;
}
.landing-hero {
  max-width: 1280px; margin: 0 auto; padding: 56px 40px 40px;
}
.landing-eyebrow {
  font-size: 11px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.14em; color: var(--green); margin-bottom: 16px;
  display: inline-flex; align-items: center; gap: 8px;
}
.landing-eyebrow-dot {
  width: 6px; height: 6px; border-radius: 50%;
  background: var(--green); box-shadow: 0 0 8px var(--green);
}
.landing-h1 {
  font-size: clamp(36px, 5vw, 52px); font-weight: 800;
  line-height: 1.08; letter-spacing: -0.04em; color: var(--text-1) !important;
  margin: 0 0 20px 0;
}
section.main:has(.landing-nav) [data-testid="stMarkdownContainer"] h1.landing-h1 {
  font-size: clamp(36px, 5vw, 52px) !important;
  font-weight: 800 !important; color: var(--text-1) !important;
  line-height: 1.08 !important; letter-spacing: -0.04em !important;
  margin: 0 0 20px 0 !important; border: none !important; padding: 0 !important;
}
.landing-h1 em {
  font-style: normal;
  background: linear-gradient(135deg, #60a5fa 0%, #34d399 100%);
  -webkit-background-clip: text; -webkit-text-fill-color: transparent;
  background-clip: text;
}
.landing-sub {
  font-size: 17px; line-height: 1.65; color: var(--text-2) !important;
  max-width: 540px; margin: 0 0 8px 0;
}
section.main:has(.landing-nav) [data-testid="stMarkdownContainer"] p.landing-sub {
  font-size: 17px !important; line-height: 1.65 !important;
  color: var(--text-2) !important; max-width: 540px; margin: 0 0 8px 0 !important;
}
.landing-mockup-shell {
  background: var(--card);
  border: 1px solid var(--border);
  border-radius: 16px;
  box-shadow: 0 24px 80px rgba(0,0,0,0.55), 0 0 0 1px rgba(255,255,255,0.05),
              0 0 60px rgba(59,130,246,0.08);
  overflow: hidden;
  transform: perspective(1400px) rotateY(-6deg) rotateX(3deg);
  transition: transform 0.45s var(--ease), box-shadow 0.45s var(--ease);
}
.landing-mockup-shell:hover {
  transform: perspective(1400px) rotateY(-2deg) rotateX(1deg);
  box-shadow: 0 32px 100px rgba(0,0,0,0.60), 0 0 0 1px rgba(255,255,255,0.06),
              0 0 80px rgba(34,197,94,0.12);
}
.landing-mock-float {
  position: absolute; bottom: 28px; left: -24px; z-index: 10;
  background: var(--card); border: 1px solid var(--border);
  border-radius: 12px; padding: 12px 14px;
  box-shadow: var(--shadow-lg); min-width: 160px;
}
.landing-mock-float-lbl {
  font-size: 8px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.1em; color: var(--text-3); margin-bottom: 6px;
}
.landing-mock-float-val {
  font-size: 22px; font-weight: 800; color: #22c55e; letter-spacing: -0.03em;
}
.landing-mock-float-sub { font-size: 10px; color: var(--text-2); margin-top: 2px; }
.landing-mock-score-chip {
  position: absolute; top: 20px; right: -18px; z-index: 10;
  background: linear-gradient(135deg, #1e3a5f, #0f1e30);
  border: 1px solid rgba(59,130,246,0.35); border-radius: 10px;
  padding: 10px 12px; box-shadow: var(--shadow-md); text-align: center;
}
.landing-mock-score-chip-num {
  font-size: 20px; font-weight: 800; color: #60a5fa; line-height: 1;
}
.landing-mock-score-chip-lbl {
  font-size: 8px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.08em; color: var(--text-3); margin-top: 4px;
}
.landing-mock-topbar {
  display: flex; align-items: center; justify-content: space-between;
  padding: 10px 14px; border-bottom: 1px solid var(--border-soft);
  background: rgba(11,22,38,0.9);
}
.landing-mock-dots { display: flex; gap: 5px; }
.landing-mock-dot { width: 8px; height: 8px; border-radius: 50%; }
.landing-mock-body { display: flex; min-height: 340px; }
.landing-mock-sidebar {
  width: 128px; border-right: 1px solid var(--border-soft);
  padding: 12px 8px; background: rgba(7,16,31,0.55); flex-shrink: 0;
}
.landing-mock-nav-item {
  font-size: 10px; font-weight: 600; color: var(--text-3);
  padding: 7px 8px; border-radius: 6px; margin-bottom: 2px;
}
.landing-mock-nav-item.active {
  background: var(--accent-dim); color: #60a5fa;
  border: 1px solid rgba(59,130,246,0.25);
}
.landing-mock-main { flex: 1; padding: 12px 14px; min-width: 0; }
.landing-mock-forge-bar {
  display: flex; align-items: center; justify-content: space-between;
  padding: 6px 0 10px; border-bottom: 1px solid var(--border-soft); margin-bottom: 10px;
}
.landing-mock-breadcrumb {
  font-size: 9px; color: var(--text-3); font-weight: 500;
}
.landing-mock-breadcrumb span { color: var(--text-1); font-weight: 600; }
.landing-mock-status {
  font-size: 8px; color: var(--green); font-weight: 700;
  display: flex; align-items: center; gap: 4px;
}
.landing-mock-status-dot {
  width: 5px; height: 5px; border-radius: 50%; background: var(--green);
  box-shadow: 0 0 6px var(--green);
}
.landing-mock-kpis {
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 6px; margin-bottom: 10px;
}
.landing-mock-kpi {
  background: var(--surface); border: 1px solid var(--border-soft);
  border-radius: 8px; padding: 10px;
}
.landing-mock-kpi-lbl { font-size: 8px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.08em; color: var(--text-3); }
.landing-mock-kpi-val { font-size: 18px; font-weight: 800; color: var(--text-1);
  margin-top: 4px; letter-spacing: -0.03em; }
.landing-mock-pipeline {
  display: grid; grid-template-columns: repeat(7, 1fr); gap: 5px;
}
.landing-mock-stage {
  text-align: center; padding: 8px 4px; border-radius: 6px;
  border: 1px solid var(--border-soft); background: var(--surface);
  font-size: 8px; font-weight: 600; color: var(--text-3);
}
.landing-mock-stage-num {
  font-size: 16px; font-weight: 800; display: block; margin-bottom: 2px;
}
.landing-mock-chart {
  display: flex; align-items: flex-end; gap: 3px; height: 44px;
  padding: 8px 10px; background: var(--surface); border: 1px solid var(--border-soft);
  border-radius: 8px; margin-bottom: 10px;
}
.landing-mock-bar {
  flex: 1; border-radius: 3px 3px 0 0; background: linear-gradient(180deg, #3b82f6, #1d4ed8);
  opacity: 0.85;
}
.landing-mock-table {
  width: 100%; border-collapse: collapse; font-size: 8px;
}
.landing-mock-table th {
  text-align: left; padding: 5px 6px; font-size: 7px; font-weight: 700;
  text-transform: uppercase; letter-spacing: 0.06em; color: var(--text-3);
  border-bottom: 1px solid var(--border-soft);
}
.landing-mock-table td {
  padding: 6px; color: var(--text-2); border-bottom: 1px solid var(--border-soft);
}
.landing-mock-table td:first-child { color: var(--text-3); font-weight: 600; }
.landing-mock-table td:nth-child(2) { color: var(--text-1); font-weight: 500; }
.landing-mock-badge {
  display: inline-block; padding: 2px 5px; border-radius: 4px;
  font-size: 7px; font-weight: 700;
}
.landing-mock-badge-green { background: rgba(34,197,94,0.15); color: #22c55e; }
.landing-mock-badge-blue { background: rgba(59,130,246,0.15); color: #60a5fa; }
.landing-mock-badge-amber { background: rgba(245,158,11,0.15); color: #f59e0b; }
.landing-mock-pipeline-hd {
  font-size: 8px; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.1em; color: var(--text-3); margin-bottom: 6px;
}
.landing-mock-kanban {
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 5px;
}
.landing-mock-k-col {
  background: var(--surface); border: 1px solid var(--border-soft);
  border-radius: 6px; padding: 5px; min-height: 52px;
}
.landing-mock-k-hd {
  font-size: 7px; font-weight: 700; color: var(--text-3);
  margin-bottom: 4px; padding-bottom: 3px; border-bottom: 2px solid;
}
.landing-mock-k-card {
  background: var(--card); border: 1px solid var(--border-soft);
  border-radius: 4px; padding: 4px 5px; margin-bottom: 3px;
  font-size: 7px; font-weight: 600; color: var(--text-1); line-height: 1.3;
}
.landing-mock-k-score {
  font-size: 6px; font-weight: 700; margin-top: 2px;
}
@media (max-width: 960px) {
  .landing-hero-inner { grid-template-columns: 1fr; gap: 32px; }
  .landing-hero-section { padding: 40px 24px 16px; }
  .landing-hero-cta-bar { padding: 12px 24px 24px; }
  .landing-hero-actions { padding: 0 24px 32px; }
  .landing-mock-float, .landing-mock-score-chip { display: none; }
}
.landing-benefits {
  max-width: 1280px; margin: 0 auto; padding: 20px 40px 48px;
}
.landing-benefits-hd {
  text-align: center; font-size: 11px; font-weight: 700;
  text-transform: uppercase; letter-spacing: 0.12em;
  color: var(--text-3); margin-bottom: 28px;
}
.landing-benefit-card {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-md); padding: 24px 20px; height: 100%;
  box-shadow: var(--shadow-sm); transition: all var(--dur) var(--ease);
}
.landing-benefit-card:hover {
  border-color: rgba(59,130,246,0.35); transform: translateY(-3px);
  box-shadow: var(--shadow-md);
}
.landing-benefit-icon {
  font-size: 28px; margin-bottom: 12px; display: block;
}
.landing-benefit-title {
  font-size: 15px; font-weight: 700; color: var(--text-1); margin-bottom: 8px;
}
.landing-benefit-desc {
  font-size: 13px; line-height: 1.6; color: var(--text-2); margin: 0;
}
.landing-footer {
  border-top: 1px solid var(--border-soft);
  padding: 28px 40px 36px; text-align: center;
  background: rgba(7,16,31,0.6);
}
.landing-footer p { font-size: 12px; color: var(--text-3); margin: 4px 0; }
section.main:has(.landing-nav) [data-testid="column"] { padding-top: 0; }
.landing-page [data-testid="column"] { padding-top: 0; }
section.main:has(.landing-nav) [data-testid="stButton"] button {
  border-radius: 10px !important; font-weight: 700 !important;
  font-size: 14px !important; padding: 12px 20px !important;
  transition: all 0.2s var(--ease) !important;
}
section.main:has(.landing-nav) [data-testid="stButton"] button[kind="primary"],
section.main:has(.landing-nav) [data-testid="stButton"] button[data-testid="baseButton-primary"] {
  background: linear-gradient(135deg, #2563eb 0%, #059669 100%) !important;
  border: none !important; box-shadow: 0 4px 20px rgba(37,99,235,0.35) !important;
}
section.main:has(.landing-nav) [data-testid="stButton"] button[kind="primary"]:hover,
section.main:has(.landing-nav) [data-testid="stButton"] button[data-testid="baseButton-primary"]:hover {
  transform: translateY(-1px) !important;
  box-shadow: 0 6px 28px rgba(37,99,235,0.45) !important;
}
section.main:has(.landing-nav) [data-testid="stButton"] button[kind="secondary"],
section.main:has(.landing-nav) [data-testid="stButton"] button[data-testid="baseButton-secondary"] {
  background: transparent !important;
  border: 1px solid var(--border) !important; color: var(--text-1) !important;
}
section.main:has(.landing-nav) .landing-hero-actions [data-testid="stButton"] {
  width: auto !important;
  margin-left: auto !important;
  margin-right: auto !important;
}
section.main:has(.landing-nav) .landing-hero-actions-btn-wrap [data-testid="stButton"] button {
  background: transparent !important;
  border: 1px solid var(--border) !important;
  color: var(--text-1) !important;
  box-shadow: none !important;
  width: auto !important;
}
section.main:has(.landing-nav) .landing-hero-actions-btn-wrap [data-testid="stButton"] button:hover {
  border-color: rgba(59,130,246,0.45) !important;
  background: rgba(59,130,246,0.08) !important;
  transform: translateY(-1px) !important;
}
.stApp:has(.landing-nav) {
  padding-top: 0 !important;
  background:
    radial-gradient(ellipse 80% 50% at 50% -20%, rgba(59,130,246,0.22) 0%, transparent 55%),
    radial-gradient(ellipse 60% 40% at 100% 20%, rgba(34,197,94,0.12) 0%, transparent 50%),
    radial-gradient(ellipse 50% 30% at 0% 60%, rgba(99,102,241,0.10) 0%, transparent 45%),
    var(--bg) !important;
}
div[data-testid="stAppViewContainer"]:has(.landing-nav) {
  padding-top: 0 !important;
}
div[data-testid="stAppViewContainer"]:has(.landing-nav) .block-container,
section.main:has(.landing-nav) .block-container {
  padding: 0 !important; max-width: 100% !important;
}
section.main:has(.landing-nav) > div {
  padding-top: 0 !important;
}
section.main:has(.landing-nav) [data-testid="stMainBlockContainer"] {
  padding-top: 0 !important;
  gap: 0 !important;
}
section.main:has(.landing-nav) [data-testid="stVerticalBlock"] {
  gap: 0.25rem !important;
}
div[data-testid="stAppViewContainer"]:has(.landing-nav) [data-testid="stSidebar"] {
  display: none !important;
}
div[data-testid="stAppViewContainer"]:has(.landing-nav) header[data-testid="stHeader"] {
  background: transparent !important;
}
</style>
""", unsafe_allow_html=True)

# ─── Load Rubric ─────────────────────────────────────────────────────────────
@st.cache_data
def load_rubric():
    path = os.path.join(os.path.dirname(__file__), "rubric.json")
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}

rubric = load_rubric()

def _debug_log(message: str, data: dict | None = None, *, hypothesis_id: str = "GEN", run_id: str = "pre-fix", location: str = "app.py") -> None:
    # #region agent log
    try:
        with open("debug-8f7ea9.log", "a", encoding="utf-8") as _dbg_f:
            _dbg_f.write(json.dumps({
                "sessionId": "8f7ea9",
                "runId": run_id,
                "hypothesisId": hypothesis_id,
                "location": location,
                "message": message,
                "data": data or {},
                "timestamp": int(time.time() * 1000),
            }) + "\n")
    except Exception:
        pass
    # #endregion

# ─── Session State ────────────────────────────────────────────────────────────
if "submissions" not in st.session_state:
    st.session_state.submissions = []
if "next_id" not in st.session_state:
    st.session_state.next_id = 1001
if "flash_msg" not in st.session_state:
    st.session_state.flash_msg = None
if "bulk_upload_row_ids" not in st.session_state:
    st.session_state.bulk_upload_row_ids = [0]
if "bulk_upload_next_row_id" not in st.session_state:
    st.session_state.bulk_upload_next_row_id = 1
if "single_stage_analysis" not in st.session_state:
    st.session_state.single_stage_analysis = None
if "single_stage_analysis_request_sig" not in st.session_state:
    st.session_state.single_stage_analysis_request_sig = None
if "bulk_stage_analysis_cache" not in st.session_state:
    st.session_state.bulk_stage_analysis_cache = {}
if "scoring_mode" not in st.session_state:
    st.session_state.scoring_mode = "Simulated"
if "session_llm_key" not in st.session_state:
    # Load from disk on first run so key survives page refreshes
    st.session_state.session_llm_key = _load_saved_key()

# ─── Shortlist (session-persistent category folders) ─────────────────────────
SHORTLIST_CATEGORIES = [
    "Clothing & Fashion",
    "Food & Beverage / CPG",
    "Automotive Parts",
    "Consumer Hardware",
    "Sustainable Packaging",
    "Furniture & Home Goods",
    "Other",
]

# Visual metadata for folder cards on the Shortlist page
SHORTLIST_CATEGORY_META = {
    "Clothing & Fashion":       {"icon": "👗", "color": "#a78bfa"},
    "Food & Beverage / CPG":    {"icon": "🥗", "color": "#22c55e"},
    "Automotive Parts":         {"icon": "🚗", "color": "#f59e0b"},
    "Consumer Hardware":        {"icon": "🔌", "color": "#3b82f6"},
    "Sustainable Packaging":    {"icon": "📦", "color": "#10b981"},
    "Furniture & Home Goods":   {"icon": "🪑", "color": "#d97706"},
    "Other":                    {"icon": "📁", "color": "#8b949e"},
}

if "shortlist" not in st.session_state:
    # category name → ordered list of submission IDs
    if "favorites" in st.session_state:
        # Migrate legacy session keys from earlier builds
        st.session_state.shortlist = st.session_state.favorites
        del st.session_state["favorites"]
    else:
        st.session_state.shortlist = {cat: [] for cat in SHORTLIST_CATEGORIES}
if "shortlist_pick_sub_id" not in st.session_state:
    # When set, the category-picker dialog is shown for this submission
    st.session_state.shortlist_pick_sub_id = None
    if "fav_pick_sub_id" in st.session_state and st.session_state.fav_pick_sub_id:
        st.session_state.shortlist_pick_sub_id = st.session_state.fav_pick_sub_id
        del st.session_state["fav_pick_sub_id"]
if "shortlist_view_category" not in st.session_state:
    # Drill-down folder on the Shortlist page (None = grid view)
    st.session_state.shortlist_view_category = None
    if "favorites_view_category" in st.session_state:
        st.session_state.shortlist_view_category = st.session_state.favorites_view_category
        del st.session_state["favorites_view_category"]

# Investment memos: submission_id → {markdown, generated_at, source}
if "investment_memos" not in st.session_state:
    st.session_state.investment_memos = {}
if "memo_sub_id" not in st.session_state:
    st.session_state.memo_sub_id = None
if "memo_needs_generate" not in st.session_state:
    st.session_state.memo_needs_generate = False
if "memo_force_regenerate" not in st.session_state:
    st.session_state.memo_force_regenerate = False

# Criterion deep-dive cache: "{sub_id}::{criterion}" → {markdown, source, warning}
# Active modal: criterion_deep_dive_active = "deep_dive_{sub_id}_{criterion_slug}"
if "criterion_detail_cache" not in st.session_state:
    st.session_state.criterion_detail_cache = {}
if "criterion_deep_dive_active" not in st.session_state:
    st.session_state.criterion_deep_dive_active = None
if "criterion_deep_dive_sub_id" not in st.session_state:
    st.session_state.criterion_deep_dive_sub_id = None
if "criterion_deep_dive_criterion" not in st.session_state:
    st.session_state.criterion_deep_dive_criterion = None
if "criterion_detail_force" not in st.session_state:
    st.session_state.criterion_detail_force = False
if "_tracked_nav_page" not in st.session_state:
    st.session_state._tracked_nav_page = None
if "_tracked_sl_view" not in st.session_state:
    st.session_state._tracked_sl_view = None
if "pipeline_detail_sub_id" not in st.session_state:
    st.session_state.pipeline_detail_sub_id = None

# Landing / navigation (Home is default entry point)
NAV_OPTIONS = ["Home", "Dashboard", "Submissions", "Shortlist", "Pipeline", "Rubric Settings"]
if "nav_page" not in st.session_state:
    st.session_state.nav_page = "Home"


def _apply_nav_from_query_params() -> None:
    """Honor deep-link query params before sidebar navigation renders."""
    page_q = st.query_params.get("page")
    if page_q in NAV_OPTIONS:
        st.session_state.nav_page = page_q
    pipe_q = st.query_params.get("pipeline")
    detail_q = st.query_params.get("submission")
    if pipe_q == "detail" and detail_q:
        st.session_state.nav_page = "Pipeline"
        st.session_state.pipeline_detail_sub_id = detail_q


_apply_nav_from_query_params()

# Landing CTA link (?dashboard=1) → enter app
if st.query_params.get("dashboard") == "1":
    # #region agent log
    _debug_log(
        "dashboard query param redirect triggered",
        {"dashboard": st.query_params.get("dashboard"), "nav_page_before": st.session_state.get("nav_page")},
        hypothesis_id="H4",
        location="app.py:1696",
    )
    # #endregion
    st.session_state.nav_page = "Dashboard"
    st.query_params.clear()
    st.rerun()


def _ensure_shortlist_buckets():
    """Guarantee every category key exists (e.g. after adding new categories)."""
    for cat in SHORTLIST_CATEGORIES:
        st.session_state.shortlist.setdefault(cat, [])


def _shortlist_category_for(sub_id: str):
    """Return the category name if submission is on the shortlist, else None."""
    _ensure_shortlist_buckets()
    for cat, ids in st.session_state.shortlist.items():
        if sub_id in ids:
            return cat
    return None


def _is_shortlisted(sub_id: str) -> bool:
    return _shortlist_category_for(sub_id) is not None


def _add_to_shortlist(sub_id: str, category: str):
    """Move submission into a category folder (removes from any prior folder)."""
    _ensure_shortlist_buckets()
    if category not in SHORTLIST_CATEGORIES:
        category = "Other"
    _remove_from_shortlist(sub_id)
    ids = st.session_state.shortlist[category]
    if sub_id not in ids:
        ids.append(sub_id)


def _remove_from_shortlist(sub_id: str):
    """Remove submission from all shortlist folders."""
    _ensure_shortlist_buckets()
    for cat in SHORTLIST_CATEGORIES:
        ids = st.session_state.shortlist.get(cat, [])
        if sub_id in ids:
            ids.remove(sub_id)


def _shortlist_total_count() -> int:
    _ensure_shortlist_buckets()
    return sum(len(st.session_state.shortlist.get(cat, [])) for cat in SHORTLIST_CATEGORIES)


def _submissions_in_shortlist_category(category: str) -> list:
    """Return full submission dicts for IDs saved in a folder, in save order."""
    _ensure_shortlist_buckets()
    id_list = st.session_state.shortlist.get(category, [])
    by_id = {s["id"]: s for s in st.session_state.submissions}
    return [by_id[sid] for sid in id_list if sid in by_id]


def _purge_shortlist_orphans():
    """Drop shortlist IDs whose submissions were deleted."""
    _ensure_shortlist_buckets()
    live = {s["id"] for s in st.session_state.submissions}
    for cat in SHORTLIST_CATEGORIES:
        st.session_state.shortlist[cat] = [
            sid for sid in st.session_state.shortlist.get(cat, []) if sid in live
        ]


def _shortlist_category_dialog_body():
    """Shared body for the shortlist category picker (modal or inline fallback)."""
    sub_id = st.session_state.get("shortlist_pick_sub_id")
    if not sub_id:
        return

    sub = next((s for s in st.session_state.submissions if s["id"] == sub_id), None)
    if not sub:
        st.warning("Submission not found — it may have been deleted.")
        if st.button("Close", use_container_width=True):
            st.session_state.shortlist_pick_sub_id = None
            st.rerun()
        return

    st.markdown(
        f'<div style="font-size:15px;font-weight:600;color:#e6edf3;margin-bottom:4px;">'
        f'{_esc(sub["name"])}</div>'
        f'<div style="font-size:11px;color:#8b949e;margin-bottom:16px;">{_esc(sub_id)}</div>',
        unsafe_allow_html=True,
    )

    current = _shortlist_category_for(sub_id)
    default_idx = (
        SHORTLIST_CATEGORIES.index(current) if current in SHORTLIST_CATEGORIES else 0
    )

    category = st.selectbox(
        "Category / folder",
        SHORTLIST_CATEGORIES,
        index=default_idx,
        help="Organize saved ideas into industry folders for quick review.",
    )

    meta = SHORTLIST_CATEGORY_META.get(category, SHORTLIST_CATEGORY_META["Other"])
    st.markdown(
        f'<div style="font-size:12px;color:#8b949e;margin:8px 0 16px 0;">'
        f'{meta["icon"]} Saved ideas appear under <strong style="color:{meta["color"]};">'
        f'{_esc(category)}</strong> on the Shortlist page.</div>',
        unsafe_allow_html=True,
    )

    btn_save, btn_remove, btn_cancel = st.columns(3)
    with btn_save:
        if st.button("Save", type="primary", use_container_width=True):
            _add_to_shortlist(sub_id, category)
            st.session_state.shortlist_pick_sub_id = None
            st.session_state.flash_msg = (
                "success",
                f"Added '{sub['name']}' to Shortlist · {category}",
            )
            st.rerun()
    with btn_remove:
        if st.button(
            "Remove",
            use_container_width=True,
            disabled=not current,
            help="Remove from all shortlist folders",
        ):
            _remove_from_shortlist(sub_id)
            st.session_state.shortlist_pick_sub_id = None
            st.session_state.flash_msg = (
                "info",
                f"Removed '{sub['name']}' from Shortlist",
            )
            st.rerun()
    with btn_cancel:
        if st.button("Cancel", use_container_width=True):
            st.session_state.shortlist_pick_sub_id = None
            st.rerun()


if hasattr(st, "dialog"):
    @st.dialog("Shortlist")
    def _shortlist_category_dialog():
        """Modal: pick a category folder for the submission in shortlist_pick_sub_id."""
        _shortlist_category_dialog_body()
else:
    def _shortlist_category_dialog():
        """Inline fallback when st.dialog is unavailable."""
        st.markdown("#### Shortlist")
        _shortlist_category_dialog_body()


def _maybe_open_shortlist_dialog():
    """Open the category picker when shortlist_pick_sub_id is set."""
    if not st.session_state.get("shortlist_pick_sub_id"):
        return
    if hasattr(st, "dialog"):
        _shortlist_category_dialog()
    else:
        with st.container(border=True):
            _shortlist_category_dialog_body()


STAGES = rubric.get("pipeline_stages", [
    {"id": 1, "name": "Intake",       "color": "#6e40c9"},
    {"id": 2, "name": "Concept",      "color": "#1f6feb"},
    {"id": 3, "name": "Validation",   "color": "#0ea5e9"},
    {"id": 4, "name": "Prototyping",  "color": "#238636"},
    {"id": 5, "name": "Market Test",  "color": "#9e6a03"},
    {"id": 6, "name": "Scaling",      "color": "#d18000"},
    {"id": 7, "name": "Monitoring",   "color": "#8b949e"},
])
STAGE_NAMES = [s["name"] for s in STAGES]
THRESHOLDS  = rubric.get("scoring_thresholds", {"green": 70, "yellow": 50})
STAGE_DISPLAY_OVERRIDES = {
    "Concept": "Concept Refinement",
}

# ─── Helpers ──────────────────────────────────────────────────────────────────
def _stage_display_name(stage_name: str) -> str:
    """Friendly stage label for UI surfaces without changing stored values."""
    return STAGE_DISPLAY_OVERRIDES.get(stage_name, stage_name)


def _normalize_stage_name(stage_name: str) -> str:
    """Map model or heuristic output back to an internal pipeline stage."""
    raw = (stage_name or "").strip().lower()
    if raw in ("concept refinement", "concept", "concept review"):
        return "Concept"
    for stage in STAGE_NAMES:
        if raw == stage.lower():
            return stage
    return "Intake"


def _stage_color(stage_name: str) -> str:
    """Return configured stage color for cards and badges."""
    stage_norm = _normalize_stage_name(stage_name)
    for stage in STAGES:
        if stage["name"] == stage_norm:
            return stage.get("color", "#8b949e")
    return "#8b949e"


def score_badge_class(score):
    if score >= THRESHOLDS["green"]:  return "badge-green"
    if score >= THRESHOLDS["yellow"]: return "badge-yellow"
    return "badge-red"

def score_hex(score):
    if score >= THRESHOLDS["green"]:  return "#3fb950"
    if score >= THRESHOLDS["yellow"]: return "#d29922"
    return "#f85149"

def pill_class(status):
    return {
        "New":      "pill-new",
        "Scored":   "pill-scored",
        "In Review":"pill-review",
        "Approved": "pill-approved",
        "Rejected": "pill-rejected",
    }.get(status, "pill-new")

def forge_badge(sub):
    """
    Return (label, text_color, bg_color) for the ForgeOS AI quality badge.
    Derived from gating flags and overall score.
    Returns (None, None, None) when the submission is unscored.
    """
    if sub.get("auto_reject"):
        return "Auto-Reject", "#f85149", "#2b0f0f"
    if sub.get("high_risk"):
        return "High Risk",   "#d29922", "#2b1f05"
    score = sub.get("overall", 0)
    if score >= THRESHOLDS["green"]:
        return "Strong",      "#3fb950", "#0d2b1a"
    if score >= THRESHOLDS["yellow"]:
        return "Promising",   "#58a6ff", "#0c1e35"
    if score > 0:
        return "Needs Work",  "#8b949e", "#1c2128"
    return None, None, None

def _esc(text) -> str:
    """Escape text for safe embedding in HTML markdown."""
    return _html.escape(str(text)) if text is not None else ""


def _score_cell_html(val) -> str:
    if val and float(val) > 0:
        bc = score_badge_class(float(val))
        display = int(val) if float(val) == int(float(val)) else val
        return f'<span class="badge-score {bc}">{display}</span>'
    return '<span class="sub-cell-empty">—</span>'


def _status_pill_html(status: str) -> str:
    pc = pill_class(status)
    return f'<span class="pill {pc}">{_esc(status)}</span>'


def _ai_badge_html(sub) -> str:
    blabel, bcolor, bbg = forge_badge(sub)
    if not blabel:
        return '<span class="sub-cell-empty">—</span>'
    return (
        f'<span style="font-size:10px;font-weight:700;color:{bcolor};'
        f'background:{bbg};border:1px solid {bcolor}44;'
        f'border-radius:4px;padding:2px 7px;white-space:nowrap;">'
        f'{_esc(blabel)}</span>'
    )


def _stage_label_html(stage_name: str) -> str:
    sc = next((s["color"] for s in STAGES if s["name"] == stage_name), "#8b949e")
    return (
        f'<span style="font-size:11px;font-weight:600;color:{sc};">'
        f'● {_esc(stage_name)}</span>'
    )


_SUB_TABLE_COLS = [1.0, 2.3, 0.85, 0.85, 0.85, 1.15, 1.2, 1.2, 5.0]
_SUB_TABLE_HEADERS = [
    "ID", "Idea Name", "Overall", "Innov.", "Feas.",
    "Status", "AI Badge", "Stage", "Actions",
]


def _render_submissions_table_header():
    hd = st.columns(_SUB_TABLE_COLS)
    for col, label in zip(hd, _SUB_TABLE_HEADERS):
        col.markdown(
            f'<div class="sub-table-th">{label}</div>',
            unsafe_allow_html=True,
        )


def _render_submission_table_row(
    sub,
    rubric_data,
    *,
    key_prefix: str = "sub",
    show_shortlist: bool = True,
    show_remove_shortlist: bool = False,
    show_advance: bool = False,
):
    """
    One bordered row: data cells, action buttons, optional score breakdown.

    key_prefix — unique Streamlit widget prefix (submissions vs shortlist views).
    show_shortlist — show the ⭐ Shortlist button (Submissions page).
    show_remove_shortlist — show Remove from Shortlist (Shortlist category view).
    """
    sid = sub["id"]
    shortlist_cat = _shortlist_category_for(sid)

    with st.container(border=True):
        row = st.columns(_SUB_TABLE_COLS)

        with row[0]:
            st.markdown(f'<span class="forge-id">{_esc(sid)}</span>', unsafe_allow_html=True)

        with row[1]:
            shortlist_badge = ""
            if shortlist_cat:
                meta = SHORTLIST_CATEGORY_META.get(shortlist_cat, SHORTLIST_CATEGORY_META["Other"])
                shortlist_badge = (
                    f'<br><span style="font-size:10px;font-weight:600;color:{meta["color"]};">'
                    f'⭐ Shortlist · {_esc(shortlist_cat)}</span>'
                )
            st.markdown(
                f'<span style="font-size:13px;color:var(--text-1);font-weight:500;">'
                f'{_esc(sub["name"])}</span>'
                f'<br><span style="font-size:11px;color:var(--text-3);">'
                f'{_esc(sub.get("file_type", ""))} · {_esc(sub.get("submitted_at", ""))}</span>'
                f'{shortlist_badge}',
                unsafe_allow_html=True,
            )

        for col, field in zip(row[2:5], ("overall", "innovation", "feasibility")):
            col.markdown(_score_cell_html(sub.get(field, 0)), unsafe_allow_html=True)

        with row[5]:
            st.markdown(_status_pill_html(sub["status"]), unsafe_allow_html=True)

        with row[6]:
            st.markdown(_ai_badge_html(sub), unsafe_allow_html=True)

        with row[7]:
            st.markdown(_stage_label_html(sub.get("stage", "")), unsafe_allow_html=True)

        with row[8]:
            is_scored = bool(sub.get("categories"))
            n_act = (5 if is_scored else 4) - (0 if show_advance else 1)
            act_cols = st.columns(n_act, gap="small")
            ai = 0
            with act_cols[ai]:
                if st.button("Score", key=f"{key_prefix}_sc_{sid}", use_container_width=True):
                    _close_criterion_detail_dialog()
                    mode_label = st.session_state.get("scoring_mode", "Simulated")
                    spinner_txt = f"Scoring using ForgeOS Extensive Rubric v2 ({mode_label})…"
                    with st.spinner(spinner_txt):
                        if mode_label == "Simulated":
                            time.sleep(0.8)
                        sc2, warn2 = route_scoring(sub, rubric_data)
                        idx = next(
                            i for i, s in enumerate(st.session_state.submissions)
                            if s["id"] == sub["id"]
                        )
                        st.session_state.submissions[idx].update({
                            "overall":     sc2["overall"],
                            "innovation":  sc2["innovation"],
                            "feasibility": sc2["feasibility"],
                            "categories":  sc2["categories"],
                            "auto_reject": sc2["auto_reject"],
                            "high_risk":   sc2["high_risk"],
                            "scored_at":   sc2["scored_at"],
                            "status":      "Scored",
                        })
                        st.session_state.investment_memos.pop(sid, None)
                        _clear_criterion_detail_cache(sid)
                        _close_criterion_deep_dives_for_submission(sid)
                        gate_note = (
                            " · Auto-Reject gate triggered" if sc2["auto_reject"]
                            else (" · High-Risk flag raised" if sc2["high_risk"] else "")
                        )
                        fallback_note = f" · ⚠ {warn2}" if warn2 else ""
                        st.session_state.flash_msg = (
                            "success",
                            f"Scored '{sub['name']}' — Overall: {sc2['overall']}/100"
                            f"{gate_note}{fallback_note}",
                        )
                        st.rerun()
            ai += 1
            if show_advance:
                with act_cols[ai]:
                    cur_stage_idx = (
                        STAGE_NAMES.index(sub["stage"]) if sub["stage"] in STAGE_NAMES else -1
                    )
                    at_last = cur_stage_idx >= len(STAGE_NAMES) - 1
                    if st.button(
                        "Advance",
                        key=f"{key_prefix}_adv_{sid}",
                        disabled=at_last,
                        use_container_width=True,
                    ):
                        new_stage = STAGE_NAMES[cur_stage_idx + 1]
                        with st.spinner(f"Advancing '{sub['name']}' to {new_stage}…"):
                            time.sleep(0.5)
                            idx = next(
                                i for i, s in enumerate(st.session_state.submissions)
                                if s["id"] == sub["id"]
                            )
                            summary = generate_stage_summary(
                                st.session_state.submissions[idx], new_stage
                            )
                            hist = st.session_state.submissions[idx].get("stage_history", [])
                            hist.append({
                                "stage": new_stage,
                                "moved_at": datetime.now().strftime("%Y-%m-%d"),
                            })
                            st.session_state.submissions[idx]["stage"] = new_stage
                            st.session_state.submissions[idx]["stage_summary"] = summary
                            st.session_state.submissions[idx]["stage_history"] = hist
                            st.session_state.flash_msg = (
                                "info",
                                f"'{sub['name']}' advanced to {new_stage} — AI stage brief generated",
                            )
                            st.rerun()
                ai += 1
            if is_scored:
                with act_cols[ai]:
                    if st.button(
                        "📄 Memo",
                        key=f"{key_prefix}_memo_{sid}",
                        use_container_width=True,
                        help="Generate investor-ready executive summary",
                    ):
                        _close_criterion_detail_dialog()
                        st.session_state.memo_sub_id = sid
                        if sid not in st.session_state.investment_memos:
                            st.session_state.memo_needs_generate = True
                        st.rerun()
                ai += 1
            with act_cols[ai]:
                if show_remove_shortlist:
                    if st.button(
                        "Remove from Shortlist",
                        key=f"{key_prefix}_rmsl_{sid}",
                        use_container_width=True,
                        help="Remove this idea from the Shortlist",
                    ):
                        _remove_from_shortlist(sid)
                        st.session_state.flash_msg = (
                            "info",
                            f"Removed '{sub['name']}' from Shortlist",
                        )
                        st.rerun()
                elif show_shortlist:
                    sl_btn_label = "⭐ Shortlisted" if shortlist_cat else "⭐ Shortlist"
                    sl_help = (
                        f"On Shortlist under {shortlist_cat}. Click to change category or remove."
                        if shortlist_cat
                        else "Add this idea to a Shortlist category folder"
                    )
                    if st.button(
                        sl_btn_label,
                        key=f"{key_prefix}_sl_{sid}",
                        use_container_width=True,
                        help=sl_help,
                    ):
                        st.session_state.shortlist_pick_sub_id = sid
                        st.rerun()
            ai += 1
            with act_cols[ai]:
                if st.button("Delete", key=f"{key_prefix}_del_{sid}", use_container_width=True):
                    _remove_from_shortlist(sid)
                    st.session_state.investment_memos.pop(sid, None)
                    _clear_criterion_detail_cache(sid)
                    _close_criterion_deep_dives_for_submission(sid)
                    st.session_state.submissions = [
                        s for s in st.session_state.submissions if s["id"] != sid
                    ]
                    st.rerun()

    if sub.get("categories"):
        with st.expander(
            f"AI Score Breakdown — {sub['name']}",
            expanded=False,
        ):
            _render_score_breakdown_panel(
                sub,
                breakdown_page=st.session_state.get("nav_page"),
            )


def _render_score_breakdown_panel(sub, *, breakdown_page: str | None = None):
    """Score breakdown content (gating, files, gauges, criteria) for one submission."""
    page_ctx = breakdown_page or st.session_state.get("nav_page", "")
    ar = sub.get("auto_reject", [])
    hr = sub.get("high_risk", [])
    if ar:
        ar_items = "".join(f"<li>{_esc(r)}</li>" for r in ar)
        st.markdown(f"""
        <div style="background:#2b0f0f;border:1px solid #6e1818;border-radius:8px;
                    padding:12px 16px;margin-bottom:12px;">
          <div style="font-size:11px;font-weight:700;color:#f85149;
                      text-transform:uppercase;letter-spacing:0.08em;margin-bottom:6px;">
            Auto-Reject Gate Triggered</div>
          <ul style="margin:0;padding-left:16px;font-size:12px;color:#f85149;">{ar_items}</ul>
        </div>""", unsafe_allow_html=True)
    if hr:
        hr_items = "".join(f"<li>{_esc(r)}</li>" for r in hr)
        st.markdown(f"""
        <div style="background:#2b1f05;border:1px solid #9e6a03;border-radius:8px;
                    padding:12px 16px;margin-bottom:12px;">
          <div style="font-size:11px;font-weight:700;color:#d29922;
                      text-transform:uppercase;letter-spacing:0.08em;margin-bottom:6px;">
            High-Risk Flag</div>
          <ul style="margin:0;padding-left:16px;font-size:12px;color:#d29922;">{hr_items}</ul>
        </div>""", unsafe_allow_html=True)

    file_sums = sub.get("file_summaries", [])
    if file_sums:
        _ftype_icons = {
            "pdf": "📄", "image": "🖼", "video": "🎬",
            "text": "📝", "other": "📎",
        }
        _method_labels = {
            "PyMuPDF-blocks": "PyMuPDF (blocks)", "PyMuPDF-text": "PyMuPDF (text)",
            "PyMuPDF-dict": "PyMuPDF (dict)", "PyMuPDF-words": "PyMuPDF (words)",
            "PyMuPDF-html": "PyMuPDF (html)", "PyMuPDF-raw": "PyMuPDF (raw)",
            "PyMuPDF-text-flags": "PyMuPDF (text+flags)",
            "pypdf": "pypdf", "tesseract-ocr": "OCR (Tesseract)",
            "pdf-vision-llm": "Vision fallback",
            "pypdf+vision-llm": "pypdf + Vision",
            "PyMuPDF-text+vision-llm": "PyMuPDF + Vision",
            "image-pdf-limited": "Image PDF (limited text)",
            "vision_llm+pillow": "LLM Vision", "pillow_metadata": "Pillow",
            "cv2+vision": "cv2 + Vision", "cv2": "cv2",
            "metadata": "metadata", "utf8": "text", "python-docx": "docx",
            "empty": "empty",
        }
        ext_note = sub.get("extraction_note") or _build_extraction_note(file_sums)
        if ext_note:
            st.markdown(
                f'<div style="font-size:11px;color:#8aa4c0;background:rgba(59,130,246,0.08);'
                f'border:1px solid rgba(59,130,246,0.2);border-radius:6px;'
                f'padding:8px 12px;margin-bottom:10px;">📎 {_esc(ext_note)}</div>',
                unsafe_allow_html=True,
            )
        st.markdown(
            '<div style="font-size:10px;font-weight:700;color:#8b949e;'
            'text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;">'
            'Uploaded Files</div>',
            unsafe_allow_html=True,
        )
        fs_cols = st.columns(min(len(file_sums), 3))
        for fi, fs in enumerate(file_sums):
            with fs_cols[fi % 3]:
                ftype_ic = _ftype_icons.get(fs.get("file_type", "other"), "📎")
                method_lbl = _method_labels.get(fs.get("extraction_method", ""))
                if not method_lbl:
                    raw_m = fs.get("extraction_method", "")
                    if "+vision-llm" in raw_m:
                        base = raw_m.replace("+vision-llm", "")
                        base_lbl = _method_labels.get(base, base)
                        method_lbl = f"{base_lbl} + Vision"
                    else:
                        method_lbl = raw_m or "—"
                extra = ""
                if fs.get("pages"):
                    extra = f'{fs["pages"]} pages · '
                if fs.get("dimensions"):
                    extra = f'{fs["dimensions"]} · '
                chars_lbl = f'{fs["chars"]:,} chars' if fs.get("chars") else "—"
                size_lbl = f'{fs["file_size_kb"]} KB' if fs.get("file_size_kb") else ""
                fname = _esc(fs.get("name", "")[:30])
                diag_line = ""
                if fs.get("vision_error"):
                    diag_line = (
                        f'<div style="color:#d29922;font-size:10px;margin-top:3px;">'
                        f'Vision: {_esc(fs["vision_error"][:80])}</div>'
                    )
                elif fs.get("extraction_diagnostics"):
                    diag_line = (
                        f'<div style="color:#6e7681;font-size:10px;margin-top:3px;">'
                        f'{_esc(fs["extraction_diagnostics"][:100])}</div>'
                    )
                preview_snip = (fs.get("preview") or "").strip()
                preview_line = ""
                if preview_snip and fs.get("file_type") == "pdf":
                    preview_line = (
                        f'<div style="color:#484f58;font-size:9px;margin-top:4px;'
                        f'font-family:monospace;line-height:1.3;max-height:2.6em;'
                        f'overflow:hidden;">{_esc(preview_snip[:120])}…</div>'
                    )

                if fs.get("thumbnail_b64"):
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{fs["thumbnail_b64"]}" '
                        f'style="width:100%;border-radius:4px;margin-bottom:4px;" alt="">',
                        unsafe_allow_html=True,
                    )
                st.markdown(
                    f'<div style="background:#0d1117;border:1px solid #21262d;'
                    f'border-radius:6px;padding:8px 10px;margin-bottom:8px;font-size:11px;">'
                    f'<div style="font-weight:600;color:#e6edf3;margin-bottom:3px;">'
                    f'{ftype_ic} {fname}</div>'
                    f'<div style="color:#8b949e;">{extra}{chars_lbl} · {size_lbl}</div>'
                    f'<div style="color:#6e7681;font-size:10px;margin-top:2px;">'
                    f'via {_esc(method_lbl)}</div>'
                    f'{diag_line}{preview_line}</div>',
                    unsafe_allow_html=True,
                )

        raw_text = sub.get("extracted_text", "")
        preview_500 = sub.get("extraction_preview") or raw_text[:PDF_EXTRACTION_PREVIEW_CHARS]
        if preview_500.strip():
            with st.expander(
                f"Extracted Content Preview ({len(raw_text):,} chars total)",
                expanded=False,
            ):
                st.caption("First 500 characters:")
                st.code(
                    preview_500 + ("…" if len(raw_text) > PDF_EXTRACTION_PREVIEW_CHARS else ""),
                    language=None,
                )
                if len(raw_text) > PDF_EXTRACTION_PREVIEW_CHARS:
                    st.caption("Full extracted text (truncated at 8,000 chars):")
                    st.code(
                        raw_text[:8000] + ("…" if len(raw_text) > 8000 else ""),
                        language=None,
                    )

    hist = sub.get("stage_history", [])
    if hist:
        dots = ""
        for i, entry in enumerate(hist):
            is_current = entry["stage"] == sub["stage"]
            sc_clr = next(
                (s["color"] for s in STAGES if s["name"] == entry["stage"]), "#8b949e"
            )
            dot_clr = sc_clr if is_current else "#30363d"
            txt_clr = sc_clr if is_current else "#6e7681"
            fw = "700" if is_current else "500"
            sep = (
                '<span style="color:#30363d;margin:0 6px;">|</span>'
                if i < len(hist) - 1 else ""
            )
            dots += (
                f'<span style="font-size:11px;font-weight:{fw};color:{txt_clr};">'
                f'<span style="color:{dot_clr};">●</span> {_esc(entry["stage"])}'
                f'<span style="font-size:9px;color:#6e7681;margin-left:4px;">'
                f'{_esc(entry.get("moved_at", ""))}</span></span>{sep}'
            )
        st.markdown(f"""
        <div style="background:#0d1117;border:1px solid #21262d;border-radius:8px;
                    padding:10px 16px;margin-bottom:12px;overflow-x:auto;white-space:nowrap;">
          <div style="font-size:10px;font-weight:700;color:#8b949e;
                      text-transform:uppercase;letter-spacing:0.08em;margin-bottom:6px;">
            Stage History</div>
          <div>{dots}</div>
        </div>""", unsafe_allow_html=True)

    summ = sub.get("stage_summary", "")
    if summ:
        st.markdown(f"""
        <div style="background:#0c1e35;border:1px solid #1f6feb44;border-radius:8px;
                    padding:12px 16px;margin-bottom:16px;">
          <div style="font-size:10px;font-weight:700;color:#58a6ff;
                      text-transform:uppercase;letter-spacing:0.08em;margin-bottom:4px;">
            AI Stage Note — {_esc(sub.get("stage", ""))}</div>
          <div style="font-size:12px;color:#b0b8c4;line-height:1.6;">{summ}</div>
        </div>""", unsafe_allow_html=True)

    top4 = list(sub["categories"].items())[:4]
    gauge_cols = st.columns(len(top4))
    for i, (cid, cd) in enumerate(top4):
        with gauge_cols[i]:
            st.plotly_chart(
                make_gauge(cd["score"], cd["name"]),
                use_container_width=True,
                key=f"g_{sub['id']}_{i}",
            )

    st.markdown(
        '<div class="section-hd" style="margin-top:12px;">All Criteria</div>',
        unsafe_allow_html=True,
    )

    anchor_colors = {"1-3": "#f85149", "4-6": "#d29922", "7-10": "#3fb950"}
    ev_colors = {"Sufficient": "#3fb950", "Partial": "#d29922", "Insufficient": "#f85149"}

    for cid, cd in sub["categories"].items():
        c = score_hex(cd["score"])
        ev_c = ev_colors.get(cd.get("evidence", "Partial"), "#8b949e")
        s10 = cd.get("score_10", round(cd["score"] / 10, 1))
        just_txt = _esc(cd.get("justification", ""))
        ev_lbl = _esc(cd.get("evidence", "Partial"))
        rf_hits = cd.get("red_flags", [])
        wt = cd.get("weight", "—")
        crit_slug = _crit_key_slug(cid)
        sid = sub["id"]
        panel_key = _criterion_panel_key(sid, cid)

        rf_html = ""
        if rf_hits:
            rf_html = "".join(
                f'<span style="font-size:10px;color:#f85149;margin-right:8px;">'
                f'⚑ {_esc(rf)}</span>'
                for rf in rf_hits
            )
            rf_html = f'<div style="margin-top:4px;">{rf_html}</div>'

        crit_col, expand_col = st.columns([5.5, 1])
        with crit_col:
            st.markdown(f"""
            <div class="crit-breakdown-row">
              <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;">
                <span style="font-size:13px;font-weight:600;color:#e6edf3;">{_esc(cid)}</span>
                <div style="display:flex;align-items:center;gap:8px;">
                  <span style="font-size:11px;color:#8b949e;">Weight: {wt}%</span>
                  <span style="font-size:14px;font-weight:800;color:{c};">{s10}/10</span>
                  <span style="font-size:10px;font-weight:600;color:{ev_c};
                               background:{ev_c}18;border:1px solid {ev_c}44;
                               border-radius:4px;padding:1px 6px;">{ev_lbl}</span>
                </div>
              </div>
              <div style="background:#21262d;border-radius:2px;height:4px;margin-bottom:8px;">
                <div style="width:{cd['score']}%;background:{c};height:100%;border-radius:2px;"></div>
              </div>
              <div style="font-size:12px;color:#8b949e;line-height:1.5;">{just_txt}</div>
              {rf_html}
            </div>""", unsafe_allow_html=True)
        with expand_col:
            st.markdown('<div style="height:8px;"></div>', unsafe_allow_html=True)
            if st.button(
                "▸ Expand",
                key=f"crit_exp_{page_ctx}_{panel_key}",
                use_container_width=True,
                help=f"Deep-dive analysis for {cid}",
            ):
                _open_criterion_deep_dive(sid, cid)
                st.rerun()

    scored_at = sub.get("scored_at", "")
    if scored_at:
        st.markdown(
            f'<div style="font-size:11px;color:#8b949e;margin-top:8px;text-align:right;">'
            f'Scored at {_esc(scored_at)}</div>',
            unsafe_allow_html=True,
        )

def _render_shortlist_folder_grid():
    """Category folder cards on the Shortlist landing view."""
    _purge_shortlist_orphans()
    _ensure_shortlist_buckets()

    st.markdown(
        '<div class="section-hd">Browse by Category</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p style="font-size:13px;color:var(--text-3);margin:-6px 0 18px 0;">'
        "Saved submissions organized by industry. Click a folder to open its ideas.</p>",
        unsafe_allow_html=True,
    )

    cols = st.columns(3)
    for i, cat in enumerate(SHORTLIST_CATEGORIES):
        meta = SHORTLIST_CATEGORY_META.get(cat, SHORTLIST_CATEGORY_META["Other"])
        count = len(st.session_state.shortlist.get(cat, []))
        with cols[i % 3]:
            st.markdown(
                f'<div class="shortlist-folder-card" style="--shortlist-accent:{meta["color"]};">'
                f'<span class="shortlist-folder-icon">{meta["icon"]}</span>'
                f'<div class="shortlist-folder-name">{_esc(cat)}</div>'
                f'<div class="shortlist-folder-count">{count}</div>'
                f'<div class="shortlist-folder-count-label">'
                f'{"idea" if count == 1 else "ideas"} saved</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if st.button(
                "Open folder",
                key=f"sl_open_{i}",
                use_container_width=True,
                disabled=count == 0,
            ):
                st.session_state.shortlist_view_category = cat
                st.rerun()


def _render_shortlist_category_detail(category: str, rubric_data):
    """Shortlist category view — table of saved submissions in one folder."""
    meta = SHORTLIST_CATEGORY_META.get(category, SHORTLIST_CATEGORY_META["Other"])
    items = _submissions_in_shortlist_category(category)

    if st.button("← All folders", key="sl_back_all"):
        st.session_state.shortlist_view_category = None
        _close_criterion_detail_dialog()
        st.rerun()

    st.markdown(
        f'<div style="display:flex;align-items:center;gap:12px;margin:8px 0 20px 0;">'
        f'<span style="font-size:32px;">{meta["icon"]}</span>'
        f'<div>'
        f'<div style="font-size:18px;font-weight:700;color:var(--text-1);">'
        f'{_esc(category)}</div>'
        f'<div style="font-size:12px;color:var(--text-3);">'
        f'{len(items)} saved submission{"s" if len(items) != 1 else ""}</div>'
        f'</div></div>',
        unsafe_allow_html=True,
    )

    if not items:
        st.markdown("""
        <div class="empty-state" style="padding:48px 24px;">
            <div class="empty-icon">⭐</div>
            <div class="empty-title">This folder is empty</div>
            <div class="empty-sub">Use <strong>⭐ Shortlist</strong> on the Submissions page<br>
            to add ideas to this category.</div>
        </div>""", unsafe_allow_html=True)
        return

    cat_idx = SHORTLIST_CATEGORIES.index(category)
    _render_submissions_table_header()
    for sub in items:
        _render_submission_table_row(
            sub,
            rubric_data,
            key_prefix=f"sl_{cat_idx}",
            show_shortlist=False,
            show_remove_shortlist=True,
            show_advance=True,
        )


def _open_pipeline_detail(sub_id: str) -> None:
    """Open the full-page pipeline detail profile for one submission."""
    st.session_state.pipeline_detail_sub_id = sub_id
    st.session_state.nav_page = "Pipeline"
    # #region agent log
    _debug_log(
        "open pipeline detail",
        {"sub_id": sub_id, "nav_page": st.session_state.nav_page},
        hypothesis_id="H1",
        run_id="post-fix",
        location="app.py:_open_pipeline_detail",
    )
    # #endregion


def _close_pipeline_detail() -> None:
    """Return from the pipeline detail profile to the kanban board."""
    st.session_state.pipeline_detail_sub_id = None


def _sync_pipeline_detail_from_query() -> None:
    """Keep pipeline detail route in sync with URL query params."""
    pipe_q = st.query_params.get("pipeline")
    detail_q = st.query_params.get("submission")
    # #region agent log
    _debug_log(
        "sync pipeline detail from query",
        {
            "pipe_q": pipe_q,
            "detail_q": detail_q,
            "nav_page": st.session_state.get("nav_page"),
            "pipeline_detail_sub_id_before": st.session_state.get("pipeline_detail_sub_id"),
        },
        hypothesis_id="H2",
        location="app.py:2630",
    )
    # #endregion
    if pipe_q == "detail" and detail_q:
        st.session_state.pipeline_detail_sub_id = detail_q
        st.session_state.nav_page = "Pipeline"
    # #region agent log
    _debug_log(
        "sync pipeline detail from query complete",
        {"pipeline_detail_sub_id_after": st.session_state.get("pipeline_detail_sub_id")},
        hypothesis_id="H2",
        run_id="post-fix",
        location="app.py:_sync_pipeline_detail_from_query",
    )
    # #endregion


def _render_pipeline_stage_timeline(sub: dict) -> None:
    """Professional vertical timeline for pipeline detail pages."""
    hist = sub.get("stage_history", []) or []
    if not hist:
        st.markdown(
            '<div class="timeline-card"><div style="font-size:12px;color:var(--text-3);">No stage history yet.</div></div>',
            unsafe_allow_html=True,
        )
        return

    steps_html = ""
    for entry in hist:
        stage_name = entry.get("stage", "")
        stage_color = _stage_color(stage_name)
        steps_html += (
            f'<div class="timeline-step">'
            f'<div class="timeline-dot" style="background:{stage_color};"></div>'
            f'<div>'
            f'<div class="timeline-step-title">{_esc(_stage_display_name(stage_name))}</div>'
            f'<div class="timeline-step-date">{_esc(entry.get("moved_at", ""))}</div>'
            f'</div>'
            f'</div>'
        )
    st.markdown(f'<div class="timeline-card">{steps_html}</div>', unsafe_allow_html=True)


def _render_submission_evidence_panel(sub: dict) -> None:
    """Focused evidence panel for full-page submission profiles."""
    file_sums = sub.get("file_summaries", []) or []
    if file_sums:
        _ftype_icons = {
            "pdf": "📄", "image": "🖼", "video": "🎬",
            "text": "📝", "other": "📎",
        }
        ext_note = sub.get("extraction_note") or _build_extraction_note(file_sums)
        if ext_note:
            st.markdown(
                f'<div style="font-size:11px;color:#8aa4c0;background:rgba(59,130,246,0.08);'
                f'border:1px solid rgba(59,130,246,0.2);border-radius:6px;'
                f'padding:8px 12px;margin-bottom:10px;">📎 {_esc(ext_note)}</div>',
                unsafe_allow_html=True,
            )
        cols = st.columns(min(len(file_sums), 3))
        for idx, fs in enumerate(file_sums):
            with cols[idx % 3]:
                ftype_ic = _ftype_icons.get(fs.get("file_type", "other"), "📎")
                chars_lbl = f'{fs.get("chars", 0):,} chars' if fs.get("chars") else "—"
                size_lbl = f'{fs.get("file_size_kb")} KB' if fs.get("file_size_kb") else ""
                if fs.get("thumbnail_b64"):
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{fs["thumbnail_b64"]}" '
                        f'style="width:100%;border-radius:8px;margin-bottom:6px;" alt="">',
                        unsafe_allow_html=True,
                    )
                st.markdown(
                    f'<div class="profile-stat-card" style="min-height:0;">'
                    f'<div class="profile-stat-label">{ftype_ic} {_esc(fs.get("name", ""))}</div>'
                    f'<div class="profile-stat-sub" style="margin-top:8px;">'
                    f'{_esc(fs.get("extraction_method", "none"))} · {chars_lbl}'
                    f'{(" · " + size_lbl) if size_lbl else ""}</div>'
                    f'<div class="profile-stat-sub">{_esc((fs.get("preview") or "No preview available.")[:180])}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

    raw_text = (sub.get("extracted_text") or "").strip()
    if raw_text:
        with st.expander(f"Extracted Content ({len(raw_text):,} chars)", expanded=False):
            st.code(raw_text[:8000] + ("…" if len(raw_text) > 8000 else ""), language=None)
    elif not file_sums:
        st.markdown(
            '<div class="profile-stat-card"><div class="profile-stat-sub">No multimodal evidence is available for this submission yet.</div></div>',
            unsafe_allow_html=True,
        )


def _render_pipeline_detail_page(sub: dict, rubric_data: dict) -> None:
    """Full-page company / idea profile opened from the Pipeline kanban."""
    sid = sub["id"]
    overall = float(sub.get("overall", 0) or 0)
    stage_name = sub.get("stage", "")
    stage_color = _stage_color(stage_name)
    shortlist_cat = _shortlist_category_for(sid)
    memo_cached = sid in st.session_state.investment_memos

    st.markdown(f"""
    <div class="forge-topbar">
      <div class="forge-topbar-left">
        <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Pipeline</span> <span class="forge-sep">/</span> <span>{_esc(sub['name'])}</span></div>
        <div class="forge-page-tag">Submission Profile</div>
      </div>
      <div class="forge-topbar-status">
        <div class="forge-status-dot"></div>
        {_esc(sid)}
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="page-content">', unsafe_allow_html=True)

    top_left, top_mid, top_right = st.columns([1.1, 1.3, 1.2], gap="medium")
    with top_left:
        if st.button("← Back to Pipeline", key="pipe_detail_back", use_container_width=True):
            _close_pipeline_detail()
            st.session_state.nav_page = "Pipeline"
            st.rerun()
    with top_mid:
        memo_label = "📄 Open Investment Memo" if memo_cached else "📄 Generate Investment Memo"
        if st.button(
            memo_label,
            key=f"pipe_detail_memo_{sid}",
            use_container_width=True,
            disabled=not bool(sub.get("categories")),
        ):
            _close_criterion_detail_dialog()
            st.session_state.memo_sub_id = sid
            if sid not in st.session_state.investment_memos:
                st.session_state.memo_needs_generate = True
            st.rerun()
    with top_right:
        sl_label = f"⭐ Shortlisted · {shortlist_cat}" if shortlist_cat else "⭐ Add to Shortlist"
        if st.button(sl_label, key=f"pipe_detail_shortlist_{sid}", use_container_width=True):
            st.session_state.shortlist_pick_sub_id = sid
            st.rerun()

    cur_idx = STAGE_NAMES.index(stage_name) if stage_name in STAGE_NAMES else -1
    at_last = cur_idx >= len(STAGE_NAMES) - 1
    act_score, act_adv, act_del = st.columns(3, gap="small")
    with act_score:
        if st.button("Score", key=f"pipe_detail_sc_{sid}", use_container_width=True):
            mode_label = st.session_state.get("scoring_mode", "Simulated")
            with st.spinner(f"Scoring using ForgeOS Extensive Rubric v2 ({mode_label})…"):
                if mode_label == "Simulated":
                    time.sleep(0.8)
                sc2, warn2 = route_scoring(sub, rubric_data)
                idx = next(i for i, s in enumerate(st.session_state.submissions) if s["id"] == sid)
                st.session_state.submissions[idx].update({
                    "overall": sc2["overall"],
                    "innovation": sc2["innovation"],
                    "feasibility": sc2["feasibility"],
                    "categories": sc2["categories"],
                    "auto_reject": sc2["auto_reject"],
                    "high_risk": sc2["high_risk"],
                    "scored_at": sc2["scored_at"],
                    "status": "Scored",
                })
                st.session_state.investment_memos.pop(sid, None)
                _clear_criterion_detail_cache(sid)
                gate_note = " · Auto-Reject gate triggered" if sc2["auto_reject"] else (" · High-Risk flag raised" if sc2["high_risk"] else "")
                fallback_note = f" · ⚠ {warn2}" if warn2 else ""
                st.session_state.flash_msg = (
                    "success",
                    f"Scored '{sub['name']}' — Overall: {sc2['overall']}/100{gate_note}{fallback_note}",
                )
                st.rerun()
    with act_adv:
        if st.button("Advance", key=f"pipe_detail_adv_{sid}", disabled=at_last, use_container_width=True):
            new_stage = STAGE_NAMES[cur_idx + 1]
            with st.spinner(f"Advancing '{sub['name']}' to {new_stage}…"):
                time.sleep(0.5)
                idx = next(i for i, s in enumerate(st.session_state.submissions) if s["id"] == sid)
                summary = generate_stage_summary(st.session_state.submissions[idx], new_stage)
                hist = st.session_state.submissions[idx].get("stage_history", [])
                hist.append({"stage": new_stage, "moved_at": datetime.now().strftime("%Y-%m-%d")})
                st.session_state.submissions[idx]["stage"] = new_stage
                st.session_state.submissions[idx]["stage_summary"] = summary
                st.session_state.submissions[idx]["stage_history"] = hist
                st.session_state.flash_msg = (
                    "info",
                    f"'{sub['name']}' advanced to {new_stage} — AI stage brief generated",
                )
                st.rerun()
    with act_del:
        if st.button("Delete", key=f"pipe_detail_del_{sid}", use_container_width=True):
            _remove_from_shortlist(sid)
            st.session_state.investment_memos.pop(sid, None)
            _clear_criterion_detail_cache(sid)
            _close_criterion_deep_dives_for_submission(sid)
            _close_pipeline_detail()
            st.session_state.submissions = [
                s for s in st.session_state.submissions if s["id"] != sid
            ]
            st.session_state.flash_msg = ("info", f"Deleted '{sub['name']}'")
            st.rerun()

    st.markdown(
        f'<div class="pipeline-detail-shell">'
        f'<div style="display:flex;justify-content:space-between;gap:18px;align-items:flex-start;flex-wrap:wrap;">'
        f'<div style="flex:1;min-width:320px;">'
        f'<div style="font-size:11px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:0.08em;">Submission Profile</div>'
        f'<div style="font-size:30px;font-weight:800;color:var(--text-1);margin-top:8px;line-height:1.15;">{_esc(sub["name"])}</div>'
        f'<div class="pipeline-detail-meta">'
        f'<span class="pipeline-detail-pill">{_esc(sid)}</span>'
        f'<span class="pipeline-detail-pill">{_esc(sub.get("file_type", "—"))}</span>'
        f'<span class="pipeline-detail-pill" style="color:{stage_color};border-color:{stage_color}55;">● {_esc(_stage_display_name(stage_name))}</span>'
        f'<span class="pipeline-detail-pill">{_esc(sub.get("status", "New"))}</span>'
        f'</div>'
        f'<div style="margin-top:14px;font-size:13px;line-height:1.7;color:var(--text-2);max-width:860px;">'
        f'{_esc((sub.get("notes") or "No submission notes provided.").strip()[:420])}'
        f'</div>'
        f'</div>'
        f'<div style="min-width:260px;max-width:320px;width:100%;background:#0b1220;border:1px solid #1f2937;border-radius:18px;padding:10px 12px;">'
        f'<div style="font-size:11px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:4px;">Overall ForgeOS Score</div>'
        f'</div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    gauge_col, meta_col1, meta_col2 = st.columns([1.25, 1, 1], gap="medium")
    with gauge_col:
        st.plotly_chart(
            make_gauge(overall, "Overall Score", height=250),
            use_container_width=True,
            key=f"pipe_profile_overall_{sid}",
        )
    with meta_col1:
        st.markdown(
            f'<div class="profile-stat-card">'
            f'<div class="profile-stat-label">Current Stage</div>'
            f'<div class="profile-stat-value" style="color:{stage_color};">{_esc(_stage_display_name(stage_name))}</div>'
            f'<div class="profile-stat-sub">This idea is currently positioned in the { _esc(_stage_display_name(stage_name)) } segment of the ForgeOS pipeline.</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    with meta_col2:
        stage_rec = sub.get("intake_stage_recommendation") or {}
        rec_stage = _stage_display_name(_normalize_stage_name(stage_rec.get("recommended_stage", "Intake")))
        rec_conf = stage_rec.get("confidence", "—")
        st.markdown(
            f'<div class="profile-stat-card">'
            f'<div class="profile-stat-label">Auto-Assigned Intake Stage</div>'
            f'<div class="profile-stat-value">{_esc(rec_stage)}</div>'
            f'<div class="profile-stat-sub">Confidence: {_esc(rec_conf)}% · Manual override: {"Yes" if sub.get("intake_stage_manual_override") else "No"}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown('<div class="section-hd" style="margin-top:10px;">Rubric Breakdown</div>', unsafe_allow_html=True)
    if sub.get("categories"):
        _render_score_breakdown_panel(sub, breakdown_page="PipelineDetail")
    else:
        st.markdown(
            '<div class="empty-state" style="padding:36px 24px;">'
            '<div class="empty-icon">📊</div>'
            '<div class="empty-title">This submission has not been scored yet</div>'
            '<div class="empty-sub">Run scoring from the Pipeline board to populate the full 8-criteria rubric breakdown and memo tools.</div>'
            '</div>',
            unsafe_allow_html=True,
        )

    preview_col, timeline_col = st.columns([1.5, 1], gap="medium")
    with preview_col:
        st.markdown('<div class="section-hd" style="margin-top:18px;">Multimodal Evidence</div>', unsafe_allow_html=True)
        _render_submission_evidence_panel(sub)
    with timeline_col:
        st.markdown('<div class="section-hd" style="margin-top:18px;">Stage Timeline</div>', unsafe_allow_html=True)
        _render_pipeline_stage_timeline(sub)

    st.markdown('</div>', unsafe_allow_html=True)


def make_gauge(score, title="", height=150):
    color = score_hex(score)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        number={"font": {"size": 24, "color": color, "family": "Inter"}, "suffix": ""},
        title={"text": title, "font": {"size": 10, "color": "#8b949e", "family": "Inter"}},
        gauge={
            "axis": {"range": [0, 100], "tickcolor": "#21262d",
                     "tickfont": {"color": "#6e7681", "size": 8}},
            "bar": {"color": color, "thickness": 0.25},
            "bgcolor": "#0d1117",
            "bordercolor": "#21262d",
            "borderwidth": 1,
            "steps": [
                {"range": [0,  40],  "color": "#3d1515"},
                {"range": [40, 70],  "color": "#3d2e08"},
                {"range": [70, 100], "color": "#0d2b1a"},
            ],
        }
    ))
    fig.update_layout(
        paper_bgcolor="#0f1e30", plot_bgcolor="#0f1e30",
        height=height, margin=dict(l=10, r=10, t=24, b=4),
        font={"family": "Inter", "color": "#4e6680"},
    )
    return fig

def _criterion_justification(key, band, display_name, score_10, anchor_txt, rng):
    """
    Return a criterion-specific, idea-aware justification that references the
    rubric anchor wording and the submission name. Two variants per (criterion,
    band) so the seeded RNG creates variety across ideas.
    """
    n = display_name  # already title-cased by caller
    s = score_10
    a = anchor_txt

    CRIT_JUST = {
        "Innovation & Novelty": {
            "1-3": [
                f"'{n}' scores {s}/10 on innovation. Rubric anchor: \"{a}\". No novel material, process, or IP pathway is evident — the concept mirrors existing market solutions without a defensible differentiation vector. Prior-art analysis is absent.",
                f"Innovation claim for '{n}' is unsupported ({s}/10). Anchor: \"{a}\". The submission describes incremental improvements at best; no breakthrough in materials, design IP, or manufacturing process could be identified against the rubric sub-factors.",
            ],
            "4-6": [
                f"'{n}' achieves moderate innovation ({s}/10). Anchor: \"{a}\". Some unique elements are present, but IP defensibility is thin and head-to-head competitors are not clearly outpaced. A stronger prior-art map and IP landscape review would materially lift this score.",
                f"Innovation for '{n}' is promising but incomplete ({s}/10). Anchor: \"{a}\". Differentiation is real but could erode quickly without clearer IP strategy or a harder-to-replicate manufacturing process advantage.",
            ],
            "7-10": [
                f"'{n}' demonstrates strong innovation ({s}/10). Anchor: \"{a}\". Breakthrough positioning across material, design, or process is identifiable, and an IP moat is credible. This submission outperforms peer ideas on novelty and competitive defensibility.",
                f"Exceptional innovation trajectory for '{n}' ({s}/10). Anchor: \"{a}\". The concept exhibits genuine differentiation — whether in novel materials, proprietary manufacturing steps, or design IP — that creates a meaningful barrier to fast-follow competition.",
            ],
        },
        "Market Potential & Saturation": {
            "1-3": [
                f"Market case for '{n}' is weak ({s}/10). Anchor: \"{a}\". Market data is absent or implausible; the saturation level is not addressed and no validated demand signal is present. TAM/SAM/SOM claims are either missing or wildly optimistic.",
                f"'{n}' fails to establish a compelling market thesis ({s}/10). Anchor: \"{a}\". The target segment is either too small, too saturated, or the submission conflates total market size with reachable opportunity.",
            ],
            "4-6": [
                f"'{n}' targets a viable but competitive market ({s}/10). Anchor: \"{a}\". Demand signals exist but TAM/SAM/SOM projections need substantiation, and competitive positioning — especially against incumbents — requires much sharper articulation.",
                f"Market potential for '{n}' is moderate ({s}/10). Anchor: \"{a}\". Growth trends are broadly supportive, but the submission does not convincingly show how the product carves out sustainable share in a crowded field.",
            ],
            "7-10": [
                f"'{n}' addresses a large and validated market ({s}/10). Anchor: \"{a}\". Market sizing is credible, the unmet need is clearly articulated with customer evidence, and growth trends strongly support the timing of entry.",
                f"Strong market case for '{n}' ({s}/10). Anchor: \"{a}\". The submission identifies a sizeable, growing segment with documented demand gaps, and positions the product to capture meaningful share without relying on overly optimistic projections.",
            ],
        },
        "Technical & Manufacturing Feasibility": {
            "1-3": [
                f"'{n}' is a pure concept at this stage ({s}/10). Anchor: \"{a}\". No manufacturing path, BOM estimate, or supplier context is provided. Scalability from prototype to production is entirely unaddressed — a critical gap at any pipeline stage.",
                f"Manufacturing feasibility for '{n}' is critically under-evidenced ({s}/10). Anchor: \"{a}\". The submission does not engage with cost structure, supply chain complexity, or process constraints that would allow a feasibility judgement.",
            ],
            "4-6": [
                f"'{n}' shows basic technical viability but scaling is unclear ({s}/10). Anchor: \"{a}\". A prototype pathway exists, but the BOM, supplier network, and cost-at-volume picture are incomplete. Significant engineering and supply chain work remains before production is realistic.",
                f"Technical feasibility for '{n}' is partially established ({s}/10). Anchor: \"{a}\". Core mechanics are plausible, but the submission lacks detail on manufacturing process, tooling costs, and supply chain resilience that would de-risk a production commitment.",
            ],
            "7-10": [
                f"'{n}' demonstrates a clear path to production ({s}/10). Anchor: \"{a}\". BOM realism, supplier references, and scalable process design are all credibly addressed. Cost structure at volume is defensible and the supply chain risk profile is manageable.",
                f"Excellent manufacturing feasibility for '{n}' ({s}/10). Anchor: \"{a}\". The submission provides a convincing factory-floor-to-shelf narrative with realistic cost models, identified supply partners, and a scalable process that survives scrutiny.",
            ],
        },
        "Sustainability & Circularity": {
            "1-3": [
                f"'{n}' shows negligible sustainability thinking ({s}/10). Anchor: \"{a}\". No lifecycle consideration, material certification, or circular design element is present. The submission risks greenwashing exposure without measurable environmental metrics.",
                f"Sustainability score for '{n}' is critically low ({s}/10). Anchor: \"{a}\". Material choices and end-of-life pathways are not addressed. Without LCA evidence or recyclability data, the product cannot credibly claim environmental positioning.",
            ],
            "4-6": [
                f"'{n}' makes basic sustainability efforts but gaps remain ({s}/10). Anchor: \"{a}\". Some eco-conscious language is present, but measurable metrics, certifications, and a genuine circular design strategy are missing. Greenwashing risk is moderate.",
                f"Sustainability approach for '{n}' is partial ({s}/10). Anchor: \"{a}\". Intentions are visible but the submission does not quantify carbon footprint, recyclability, or ethical sourcing in a way that stands up to third-party scrutiny.",
            ],
            "7-10": [
                f"'{n}' demonstrates strong circular design thinking ({s}/10). Anchor: \"{a}\". Material choices, end-of-life pathways, and measurable impact reduction are all credibly addressed. Certifications or LCA references make environmental claims verifiable.",
                f"Sustainability profile for '{n}' is compelling ({s}/10). Anchor: \"{a}\". The submission embeds circularity into the product architecture — not as a marketing overlay but as a structural design principle — with quantifiable environmental commitments.",
            ],
        },
        "Regulatory Compliance & Risk": {
            "1-3": [
                f"'{n}' does not address regulatory requirements ({s}/10). Anchor: \"{a}\". Key standards, certifications, and compliance pathways are absent. This is a critical blocker — without a regulatory roadmap the product cannot reach market in most jurisdictions.",
                f"Regulatory risk for '{n}' is high ({s}/10). Anchor: \"{a}\". No evidence of awareness of relevant directives, safety testing requirements, or certification timelines. This gap could significantly delay or prevent market entry.",
            ],
            "4-6": [
                f"'{n}' has basic regulatory awareness but an incomplete compliance plan ({s}/10). Anchor: \"{a}\". Relevant regulations are acknowledged, but the roadmap to certification — including timelines, costs, and responsible parties — is underdeveloped.",
                f"Compliance approach for '{n}' is nascent ({s}/10). Anchor: \"{a}\". The submission identifies the regulatory landscape at a surface level but does not demonstrate a structured path through testing, approval, and ongoing compliance management.",
            ],
            "7-10": [
                f"'{n}' presents a clear regulatory roadmap ({s}/10). Anchor: \"{a}\". Relevant standards are identified, certification pathways are mapped, and risk mitigation measures are credible. Compliance is treated as a strategic asset, not an afterthought.",
                f"Strong compliance posture for '{n}' ({s}/10). Anchor: \"{a}\". The submission demonstrates proactive regulatory engagement — standards are named, testing protocols are outlined, and the timeline to market approval is realistic and risk-adjusted.",
            ],
        },
        "Team & Execution Capability": {
            "1-3": [
                f"Team behind '{n}' lacks the credentials to execute ({s}/10). Anchor: \"{a}\". No relevant physical-goods experience, manufacturing expertise, or supply chain relationships are demonstrated. Key functional roles appear unfilled with no hiring plan.",
                f"Execution capability for '{n}' is critically under-evidenced ({s}/10). Anchor: \"{a}\". The submission does not establish credibility in the domain — no track record in bringing comparable physical products to market is presented.",
            ],
            "4-6": [
                f"'{n}' has a partially capable team ({s}/10). Anchor: \"{a}\". Some relevant expertise is present, but critical capability gaps exist — particularly in manufacturing operations, regulatory affairs, or commercial channels. An explicit hiring roadmap would strengthen confidence.",
                f"Team for '{n}' shows promise with identified gaps ({s}/10). Anchor: \"{a}\". Founding competencies are present in some domains but the submission acknowledges — or ignores — functional holes that will create execution risk at scale.",
            ],
            "7-10": [
                f"'{n}' is backed by a strong, balanced execution team ({s}/10). Anchor: \"{a}\". Demonstrated track record in physical goods, manufacturing partnerships, and commercial channels is clearly evidenced. The team has credibly done this before.",
                f"Excellent team capability for '{n}' ({s}/10). Anchor: \"{a}\". The founders bring domain-specific depth — engineering, operations, and commercial roles are all credibly covered, with a track record that validates their ability to navigate production and go-to-market.",
            ],
        },
        "Business Model & Commercial Viability": {
            "1-3": [
                f"Commercial case for '{n}' is underdeveloped ({s}/10). Anchor: \"{a}\". No clear revenue model, pricing rationale, or unit economics are presented. The path from concept to profitable operation is entirely hypothetical.",
                f"'{n}' lacks a viable business model ({s}/10). Anchor: \"{a}\". Revenue streams are vague, margin structure is absent, and the go-to-market plan does not establish how the product reaches paying customers at commercially viable volumes.",
            ],
            "4-6": [
                f"'{n}' has a basic but unproven commercial model ({s}/10). Anchor: \"{a}\". Revenue logic is present, but unit economics are not fully worked through and the go-to-market approach lacks channel specificity. Early traction or letters of intent would materially strengthen this.",
                f"Business model for '{n}' is directionally sound but thin ({s}/10). Anchor: \"{a}\". Pricing assumptions are broadly reasonable, but margin analysis at target volumes and channel cost structure are not sufficiently evidenced to build high commercial confidence.",
            ],
            "7-10": [
                f"'{n}' presents a clear path to profitable scaling ({s}/10). Anchor: \"{a}\". Unit economics are realistic, pricing is validated against customer willingness to pay, and the go-to-market plan is specific enough to be credible with well-defined channel partners.",
                f"Strong commercial viability for '{n}' ({s}/10). Anchor: \"{a}\". The revenue model is coherent, margins are defensible at target volumes, and early market signals — through pilots, LOIs, or channel relationships — reduce execution risk significantly.",
            ],
        },
        "Evidence Quality & Realism": {
            "1-3": [
                f"'{n}' is heavy on aspiration, light on evidence ({s}/10). Anchor: \"{a}\". Claims throughout the submission are unsupported by data, references, or validated tests. The overall picture is optimistic but not credible against rigorous rubric scrutiny.",
                f"Evidence quality for '{n}' is critically low ({s}/10). Anchor: \"{a}\". Contradictory statements and unsubstantiated projections undermine confidence across all other criteria. The submission reads as a pitch narrative rather than a grounded innovation case.",
            ],
            "4-6": [
                f"'{n}' provides partial evidence with notable gaps ({s}/10). Anchor: \"{a}\". Some claims are substantiated, but key assertions — particularly around market size, manufacturing costs, and technical performance — rest on assumptions rather than data.",
                f"Realism for '{n}' is moderate ({s}/10). Anchor: \"{a}\". The submission is internally consistent in places but uneven — well-evidenced sections sit alongside claims that are speculative or that rely on best-case scenario projections.",
            ],
            "7-10": [
                f"'{n}' is well-supported, consistent, and realistic ({s}/10). Anchor: \"{a}\". Evidence is cited throughout — market data, technical references, team credentials, and financial assumptions are all grounded and internally consistent. Confidence in this submission is high.",
                f"Excellent evidence quality for '{n}' ({s}/10). Anchor: \"{a}\". The submission demonstrates disciplined realism — projections are conservative, assumptions are made explicit, and supporting data sources are traceable. This is the standard against which peer submissions should be benchmarked.",
            ],
        },
    }

    # Fall back to generic if criterion name not found (future rubric additions)
    generic = {
        "1-3": [f"'{n}' scores {s}/10 — below threshold. Rubric anchor: \"{a}\". Critical gaps identified; immediate remediation required before advancing."],
        "4-6": [f"'{n}' scores {s}/10 — moderate. Rubric anchor: \"{a}\". Partial evidence present; deeper validation recommended across key sub-factors."],
        "7-10": [f"'{n}' scores {s}/10 — strong. Rubric anchor: \"{a}\". Submission exceeds sector benchmarks; clear validation pathway demonstrated."],
    }
    options = CRIT_JUST.get(key, generic).get(band, generic[band])
    return rng.choice(options)


def ai_score_submission(submission, rubric_data):
    """
    Chain-of-thought AI scoring engine.

    Reads all 8 criteria from rubric.json, derives context signals from the
    submission name/notes using keyword matching, then produces per-criterion
    scores (1-10) with idea-specific, rubric-anchored justifications.
    Applies gating rules and returns a fully structured breakdown.

    → Swap the RNG internals for a real LLM call when the API is connected.
    """
    name         = submission.get("name",  "").lower()
    notes        = submission.get("notes", "").lower()
    text         = name + " " + notes
    display_name = submission.get("name", "This submission").title()

    # Seed RNG on submission name → reproducible scores for the same idea
    seed = int(hashlib.md5(name.encode()).hexdigest()[:8], 16)
    rng  = random.Random(seed)

    criteria = rubric_data.get("criteria", [])

    # ── Keyword signals per criterion (positive boost / negative penalty) ─────
    SIGNALS = {
        "Innovation & Novelty": (
            ["breakthrough", "novel", "patent", "unique", "first", "bio", "mycelium",
             "self-heal", "nano", "smart", "micro", "carbon", "polymer", "proprietary",
             "advanced", "new material", "new process"],
            ["me-too", "copy", "basic", "simple", "existing", "commodity", "standard"],
        ),
        "Market Potential & Saturation": (
            ["market", "demand", "customer", "growth", "billion", "segment",
             "untapped", "commercial", "b2b", "validated", "gap", "underserved"],
            ["saturated", "small market", "niche only", "declining", "crowded"],
        ),
        "Technical & Manufacturing Feasibility": (
            ["prototype", "manufacturing", "supply chain", "bom", "scalable",
             "production", "material", "motor", "thermal", "compression",
             "insulation", "packaging", "exoskeleton", "foam", "coating", "tooling"],
            ["concept only", "theoretical", "unproven at scale", "unclear process"],
        ),
        "Sustainability & Circularity": (
            ["biodegradable", "sustainable", "circular", "recyclable", "eco",
             "mycelium", "bio", "carbon", "lca", "ethical", "renewable", "compostable"],
            ["virgin plastic", "toxic", "greenwash", "no certification", "landfill"],
        ),
        "Regulatory Compliance & Risk": (
            ["compliance", "certified", "fda", "ce", "regulatory", "standard",
             "iso", "approval", "testing", "safety", "directive", "reach"],
            ["unregulated", "no compliance plan", "unapproved"],
        ),
        "Team & Execution Capability": (
            ["team", "experience", "expert", "founder", "engineer",
             "track record", "proven", "background", "led", "built", "shipped"],
            ["no team", "inexperienced", "first time", "no relevant experience"],
        ),
        "Business Model & Commercial Viability": (
            ["revenue", "profit", "margin", "pricing", "unit economics",
             "commercial", "traction", "contract", "letters of intent", "b2b", "channel"],
            ["no revenue model", "unclear model", "free only", "donated"],
        ),
        "Evidence Quality & Realism": (
            ["data", "research", "study", "validated", "tested", "evidence",
             "pilot", "prototype", "results", "measured", "demonstrated", "cited"],
            ["vague", "hype", "we believe", "speculative", "unsubstantiated"],
        ),
    }

    scored_criteria = {}

    for crit in criteria:
        key          = crit["criterion"]
        weight       = crit.get("weight", 10)
        anchors      = crit.get("scoring_anchors", {})
        rf_list      = crit.get("red_flags", [])
        sub_facs     = crit.get("sub_factors", [])
        evidence_req = crit.get("evidence_required", "")

        # Base score biased toward mid-range; seeded for reproducibility
        base = rng.randint(48, 76)

        # Keyword signal adjustment
        pos_words, neg_words = SIGNALS.get(key, ([], []))
        boost   = min(sum(3 for w in pos_words if w in text), 21)
        penalty = min(sum(5 for w in neg_words if w in text), 24)
        raw     = max(10, min(95, base + boost - penalty))

        # 1-10 scale
        score_10 = round(raw / 10.0, 1)
        score_10 = max(1.0, min(10.0, score_10))

        # Anchor band
        if score_10 <= 3:
            band       = "1-3"
            anchor_txt = anchors.get("1-3", "Below threshold")
        elif score_10 <= 6:
            band       = "4-6"
            anchor_txt = anchors.get("4-6", "Moderate")
        else:
            band       = "7-10"
            anchor_txt = anchors.get("7-10", "Strong")

        justification  = _criterion_justification(key, band, display_name, score_10, anchor_txt, rng)
        evidence_level = "Sufficient" if score_10 >= 6 else ("Partial" if score_10 >= 4 else "Insufficient")

        # Detect red flags triggered by submission text
        triggered_flags = [
            rf for rf in rf_list
            if any(w in text for w in rf.lower().split() if len(w) > 4)
        ]

        scored_criteria[key] = {
            "name":          key,
            "score_10":      score_10,
            "score":         round(score_10 * 10),  # 0-100 for gauges/progress bars
            "weight":        weight,
            "weight_frac":   weight / 100.0,
            "anchor_band":   band,
            "anchor_text":   anchor_txt,
            "justification": justification,
            "evidence":      evidence_level,
            "evidence_req":  evidence_req,
            "red_flags":     triggered_flags,
            "sub_factors":   sub_facs,
        }

    # ── Weighted overall score (0-100) ────────────────────────────────────────
    total_w = sum(v["weight_frac"] for v in scored_criteria.values())
    overall = round(
        sum(v["score"] * v["weight_frac"] for v in scored_criteria.values()) / max(total_w, 0.01),
        1,
    )
    overall = min(overall, 100.0)

    # ── Gating rules (parsed from rubric) ────────────────────────────────────
    GATE_MAP = {
        "Innovation & Novelty":                  (6.0, "auto-reject"),
        "Technical & Manufacturing Feasibility": (6.0, "auto-reject"),
        "Sustainability & Circularity":          (5.0, "high-risk"),
        "Evidence Quality & Realism":            (4.0, "auto-reject"),
    }
    auto_reject_flags, high_risk_flags = [], []
    for crit_name, (threshold, action) in GATE_MAP.items():
        if crit_name in scored_criteria:
            s = scored_criteria[crit_name]["score_10"]
            if s < threshold:
                msg = f"{crit_name}: {s:.1f}/10 — below gate threshold of {threshold}/10"
                (auto_reject_flags if action == "auto-reject" else high_risk_flags).append(msg)

    return {
        "overall":     overall,
        "innovation":  scored_criteria.get("Innovation & Novelty", {}).get("score", 0),
        "feasibility": scored_criteria.get("Technical & Manufacturing Feasibility", {}).get("score", 0),
        "categories":  scored_criteria,
        "auto_reject": auto_reject_flags,
        "high_risk":   high_risk_flags,
        "scored_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


PDF_MIN_TEXT_CHARS = 300
PDF_EXTRACTION_PREVIEW_CHARS = 500
PDF_VISION_MAX_PAGES = 8
PDF_OCR_MAX_PAGES = 15
EXTRACTED_TEXT_LLM_LIMIT = 10000

_FITZ_OK: bool | None = None
_FITZ_ERR = ""


def _fitz_available() -> tuple[bool, str]:
    """Cached check — PyMuPDF can fail at import on some Windows/Python builds."""
    global _FITZ_OK, _FITZ_ERR
    if _FITZ_OK is None:
        try:
            import fitz  # noqa: F401
            _FITZ_OK = True
        except Exception as exc:
            _FITZ_OK = False
            _FITZ_ERR = str(exc)
    return _FITZ_OK, _FITZ_ERR


def _pdf_page_count(file_bytes: bytes) -> int:
    """Page count via PyMuPDF, pypdf, or pypdfium2."""
    if _fitz_available()[0]:
        try:
            doc = _open_pdf_fitz(file_bytes)
            n = doc.page_count
            doc.close()
            return n
        except Exception:
            pass
    import io as _io

    try:
        import pypdf

        return len(pypdf.PdfReader(_io.BytesIO(file_bytes), strict=False).pages)
    except Exception:
        pass
    try:
        import pypdfium2 as pdfium

        return len(pdfium.PdfDocument(file_bytes))
    except Exception:
        return 0


def _render_pdf_page_png_bytes(file_bytes: bytes, page_index: int) -> bytes | None:
    """Render one PDF page to PNG — PyMuPDF first, then pypdfium2 if DLL unavailable."""
    if _fitz_available()[0]:
        try:
            doc = _open_pdf_fitz(file_bytes)
            png = _pdf_page_to_png_bytes(doc.load_page(page_index))
            doc.close()
            if png:
                return png
        except Exception:
            pass
    try:
        import io as _io

        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(file_bytes)
        if page_index < 0 or page_index >= len(pdf):
            return None
        pil = pdf[page_index].render(scale=2.0).to_pil().convert("RGB")
        buf = _io.BytesIO()
        pil.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _meaningful_text_len(text: str) -> int:
    """Length of text after stripping page markers, noise, and excess whitespace."""
    if not text:
        return 0
    cleaned = re.sub(r"\[Page \d+[^\]]*\]", "", text)
    cleaned = re.sub(r"\[PDF:[^\]]+\]", "", cleaned)
    cleaned = re.sub(r"\[AI Pitch Deck Vision Analysis\]", "", cleaned)
    cleaned = re.sub(r"\[Image-based PDF[^\]]*\]", "", cleaned)
    cleaned = re.sub(r"\(Vision fallback failed:[^)]+\)", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\(No extractable text[^)]*\)", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\[Note:[^\]]+\]", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return len(cleaned)


def _strip_html_tags(html: str) -> str:
    """Rough HTML → plain text for PyMuPDF html/xhtml modes."""
    import html as _html

    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
    text = re.sub(r"(?is)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</p>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = _html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _pymupdf_page_blocks(pg) -> str:
    blocks = pg.get_text("blocks")
    return "\n".join(
        b[4].strip()
        for b in sorted(blocks, key=lambda b: (b[1], b[0]))
        if len(b) > 4 and b[4].strip()
    )


def _pymupdf_page_dict(pg) -> str:
    data = pg.get_text("dict")
    lines: list[str] = []
    for block in data.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            line_text = "".join(s.get("text", "") for s in line.get("spans", [])).strip()
            if line_text:
                lines.append(line_text)
    return "\n".join(lines)


def _pymupdf_page_words(pg) -> str:
    words = pg.get_text("words")
    if not words:
        return ""
    words.sort(key=lambda w: (w[5], w[6], w[7], w[1], w[0]))
    return " ".join(w[4].strip() for w in words if len(w) > 4 and w[4].strip())


def _pymupdf_page_html(pg) -> str:
    try:
        return _strip_html_tags(pg.get_text("html"))
    except Exception:
        return ""


def _open_pdf_fitz(file_bytes: bytes):
    """Open PDF with PyMuPDF; try stream variants."""
    import io as _io

    import fitz

    last_err: Exception | None = None
    for stream in (file_bytes, _io.BytesIO(file_bytes)):
        try:
            return fitz.open(stream=stream, filetype="pdf")
        except Exception as exc:
            last_err = exc
    raise last_err or RuntimeError("Cannot open PDF")


def _pdf_page_to_png_bytes(pg, scales: tuple = (2.0, 1.5, 1.0)) -> bytes | None:
    """Render a PDF page to PNG bytes with several fallback strategies."""
    import io as _io

    import fitz
    from PIL import Image

    for scale in scales:
        for dpi in (None, int(120 * scale)):
            try:
                if dpi:
                    pix = pg.get_pixmap(dpi=dpi, alpha=False)
                else:
                    pix = pg.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                return pix.tobytes("png")
            except Exception:
                pass
    try:
        pix = pg.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _extract_pdf_text_pymupdf(file_bytes: bytes) -> tuple[str, int, str, str | None]:
    """Try several PyMuPDF text modes; return (text, page_count, method, error)."""
    ok, fitz_err = _fitz_available()
    if not ok:
        return "", 0, "none", f"PyMuPDF unavailable: {fitz_err}"

    open_error: str | None = None
    doc = None
    try:
        doc = _open_pdf_fitz(file_bytes)
    except Exception as exc:
        open_error = str(exc)

    if doc is None:
        return "", 0, "none", open_error

    page_count = doc.page_count
    page_extractors = [
        ("PyMuPDF-blocks", _pymupdf_page_blocks),
        ("PyMuPDF-text", lambda pg: pg.get_text("text", sort=True).strip()),
        ("PyMuPDF-dict", _pymupdf_page_dict),
        ("PyMuPDF-words", _pymupdf_page_words),
        ("PyMuPDF-html", _pymupdf_page_html),
        ("PyMuPDF-raw", lambda pg: pg.get_text().strip()),
    ]
    try:
        page_extractors.append(
            ("PyMuPDF-text-flags", lambda pg: pg.get_text(
                "text", sort=True, flags=fitz.TEXT_PRESERVE_WHITESPACE,
            ).strip()),
        )
    except Exception:
        pass

    best_text, best_len, best_method = "", 0, "none"
    for method_name, extractor in page_extractors:
        pages: list[str] = []
        for i in range(page_count):
            pg_text = ""
            for loader in (
                lambda idx=i: doc.load_page(idx),
                lambda idx=i: doc[idx],
            ):
                try:
                    pg = loader()
                    pg_text = extractor(pg)
                    if pg_text.strip():
                        break
                except Exception:
                    pg_text = ""
            if pg_text.strip():
                pages.append(f"[Page {i + 1}]\n{pg_text.strip()}")
        combined = "\n\n".join(pages)
        n = _meaningful_text_len(combined)
        if n > best_len:
            best_text, best_len, best_method = combined, n, method_name
    doc.close()
    return best_text, page_count, best_method, open_error


def _extract_pdf_text_pypdf(file_bytes: bytes) -> tuple[str, int, str | None]:
    """Fallback text extraction via pypdf (default + layout modes)."""
    import io as _io

    import pypdf

    err: str | None = None
    try:
        reader = pypdf.PdfReader(_io.BytesIO(file_bytes), strict=False)
    except Exception as exc:
        return "", 0, str(exc)

    pages: list[str] = []
    for i, pg in enumerate(reader.pages):
        candidates: list[str] = []
        for mode in (None, "layout"):
            try:
                if mode:
                    t = pg.extract_text(extraction_mode=mode) or ""
                else:
                    t = pg.extract_text() or ""
                if t.strip():
                    candidates.append(t.strip())
            except Exception as exc:
                err = str(exc)
        if candidates:
            best = max(candidates, key=_meaningful_text_len)
            pages.append(f"[Page {i + 1}]\n{best}")
    return "\n\n".join(pages), len(reader.pages), err


def _ocr_available() -> tuple[bool, str]:
    """Check whether Tesseract OCR is usable (avoids DLL import crashes)."""
    try:
        import pytesseract

        pytesseract.get_tesseract_version()
        return True, ""
    except Exception as exc:
        return False, str(exc)[:160]


def _extract_pdf_text_ocr(
    file_bytes: bytes, max_pages: int = PDF_OCR_MAX_PAGES,
) -> tuple[str, int, str | None]:
    """OCR scanned/image PDF pages with pytesseract (optional; never raises)."""
    ok, ocr_err = _ocr_available()
    if not ok:
        return "", 0, ocr_err or "OCR unavailable"

    try:
        import io as _io

        import pytesseract
        from PIL import Image
    except Exception as exc:
        return "", 0, str(exc)[:160]

    page_total = _pdf_page_count(file_bytes)
    if not page_total:
        return "", 0, "Cannot open PDF for OCR"

    pages: list[str] = []
    page_errors: list[str] = []
    limit = min(page_total, max_pages)
    for i in range(limit):
        try:
            png_bytes = _render_pdf_page_png_bytes(file_bytes, i)
            if not png_bytes:
                page_errors.append(f"page {i + 1}: render failed")
                continue
            img = Image.open(_io.BytesIO(png_bytes)).convert("RGB")
            ocr_text = pytesseract.image_to_string(img).strip()
            if ocr_text:
                pages.append(f"[Page {i + 1} OCR]\n{ocr_text}")
        except Exception as exc:
            page_errors.append(f"page {i + 1}: {exc}")
    err = "; ".join(page_errors[:3]) if page_errors and not pages else None
    return "\n\n".join(pages), limit, err


def _build_extraction_note(file_summaries: list) -> str:
    """Human-readable extraction summary for score breakdown UI."""
    _labels = {
        "PyMuPDF-blocks": "PyMuPDF (blocks)", "PyMuPDF-text": "PyMuPDF (text)",
        "PyMuPDF-dict": "PyMuPDF (dict)", "PyMuPDF-words": "PyMuPDF (words)",
        "PyMuPDF-html": "PyMuPDF (html)", "PyMuPDF-raw": "PyMuPDF (raw)",
        "PyMuPDF-text-flags": "PyMuPDF (text+flags)",
        "pypdf": "pypdf", "tesseract-ocr": "OCR (Tesseract)",
        "pdf-vision-llm": "Vision fallback", "image-pdf-limited": "Image PDF (limited)",
        "empty": "empty",
    }
    pdf_sums = [s for s in file_summaries if s.get("file_type") == "pdf"]
    if not pdf_sums:
        total = sum(s.get("chars", 0) for s in file_summaries)
        if file_summaries and total:
            return f"Extracted {total:,} characters total"
        return ""
    fs = pdf_sums[0]
    method = fs.get("extraction_method", "")
    method_lbl = _labels.get(method, method.replace("+", " + "))
    chars = fs.get("chars", 0)
    before = fs.get("text_chars_before_vision", 0)
    parts: list[str] = []
    if "vision" in method.lower():
        parts.append(
            f"Extracted {before:,} chars (text) · Vision fallback used → {chars:,} chars total"
        )
    elif method == "image-pdf-limited":
        parts.append(f"Image-based PDF — {chars:,} chars (limited text extraction)")
    elif "ocr" in method.lower() or "tesseract" in method.lower():
        parts.append(f"Extracted {chars:,} chars via OCR")
    else:
        parts.append(f"Extracted {chars:,} chars via {method_lbl}")
    if fs.get("vision_error"):
        parts.append(f"Vision note: {fs['vision_error'][:100]}")
    if fs.get("extraction_diagnostics"):
        parts.append(fs["extraction_diagnostics"][:120])
    return " · ".join(parts)


def _vision_describe_image(
    img_bytes: bytes,
    filename: str,
    api_key: str,
    base_url: str,
    model: str,
    prompt: str | None = None,
    max_tokens: int = 400,
) -> str:
    """
    Send an image to a vision-capable LLM and return a structured description.
    Raises RuntimeError on failure — caller handles fallback.
    """
    import base64 as _b64
    ext  = filename.rsplit(".", 1)[-1].lower() if "." in filename else "jpg"
    mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
            "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp"}.get(ext, "image/jpeg")
    data_url = f"data:{mime};base64,{_b64.b64encode(img_bytes).decode()}"

    prompt = prompt or (
        "You are analyzing an uploaded file for a physical goods product innovation platform. "
        "Describe what you see in structured form covering: "
        "(1) Product or concept shown, "
        "(2) Key design or technical features visible, "
        "(3) Apparent materials or manufacturing process, "
        "(4) Market positioning signals. "
        "Be specific and technical. Max 200 words."
    )
    payload = json.dumps({
        "model": model,
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": data_url, "detail": "high"}},
        ]}],
        "max_tokens": max_tokens,
        "temperature": 0.3,
    }).encode("utf-8")
    req = urllib.request.Request(
        url=f"{base_url}/chat/completions", data=payload, method="POST",
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode())["choices"][0]["message"]["content"].strip()
    except Exception as exc:
        raise RuntimeError(f"Vision API: {exc}") from exc


_PDF_VISION_PROMPT = (
    "You are analyzing a product innovation pitch deck (PDF page image) for a physical-goods "
    "innovation screening platform. Extract ALL readable content from this slide. Include:\n"
    "- Company / brand and product name\n"
    "- Headline value proposition and problem/solution\n"
    "- Product details, materials, form factor, packaging\n"
    "- Market size, traction, retail/distribution, pricing if shown\n"
    "- Team, timeline, funding ask, manufacturing notes\n"
    "Be thorough and factual — this replaces failed text extraction. Max 250 words per page."
)


def _vision_describe_pdf(
    file_bytes: bytes,
    filename: str,
    api_key: str,
    base_url: str,
    model: str,
    max_pages: int = PDF_VISION_MAX_PAGES,
) -> tuple[str, int, str | None]:
    """
    Render PDF pages to images and describe via vision LLM (pitch-deck fallback).
    Returns (combined_text, pages_analyzed, error_note). Never raises.
    """
    total_pages = _pdf_page_count(file_bytes)
    if not total_pages:
        fitz_ok, fitz_err = _fitz_available()
        hint = fitz_err[:80] if not fitz_ok else "unknown"
        return "", 0, f"Cannot open PDF for vision ({hint})"

    sections: list[str] = []
    page_errors: list[str] = []
    limit = min(total_pages, max_pages)
    pages_ok = 0
    for i in range(limit):
        try:
            png_bytes = _render_pdf_page_png_bytes(file_bytes, i)
            if not png_bytes:
                page_errors.append(f"page {i + 1}: could not render slide image")
                continue
            desc = _vision_describe_image(
                png_bytes,
                f"{filename}_page{i + 1}.png",
                api_key,
                base_url,
                model,
                prompt=_PDF_VISION_PROMPT,
                max_tokens=650,
            )
            if desc.strip():
                sections.append(f"[Page {i + 1} — Vision Analysis]\n{desc.strip()}")
                pages_ok += 1
        except Exception as exc:
            page_errors.append(f"page {i + 1}: {exc}")

    if total_pages > limit and sections:
        sections.append(
            f"[Note: PDF has {total_pages} pages; vision analysis covers first {limit} pages.]"
        )

    err_note = "; ".join(page_errors[:4]) if page_errors else None
    if not sections and err_note:
        return "", 0, err_note
    return "\n\n".join(sections), pages_ok, err_note


def extract_file_text(
    uploaded_files,
    status_slot=None,
    progress_bar=None,
    progress_label=None,
    use_llm_vision: bool = False,
    api_key: str = "",
    base_url: str = "",
    model: str = "",
):
    """
    Rich multimodal content extraction.
    Returns (combined_text: str, summaries: list[dict]).

    Each summary dict keys:
      name, file_type, extraction_method, chars, pages (PDF),
      dimensions (image), thumbnail_b64 (image/video), preview, file_size_kb
    """
    import io as _io
    if not uploaded_files:
        return "", []

    parts: list = []
    summaries: list = []
    total_files = len(uploaded_files)

    for idx, f in enumerate(uploaded_files, start=1):
        fname      = f.name
        flower     = fname.lower()
        ftype_mime = f.type or ""
        file_bytes = f.read()
        file_size_kb = round(len(file_bytes) / 1024, 1)

        if progress_bar:
            progress_bar.progress((idx - 1) / max(total_files, 1))
        if progress_label:
            progress_label.markdown(
                f'<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                f'Analyzing file {idx} of {total_files}: '
                f'<strong style="color:#e6edf3;">{_esc(fname)}</strong></div>',
                unsafe_allow_html=True,
            )

        summary: dict = {
            "name": fname, "file_type": "other", "extraction_method": "none",
            "chars": 0, "pages": None, "dimensions": None,
            "thumbnail_b64": None, "preview": "", "file_size_kb": file_size_kb,
        }

        if status_slot:
            status_slot.markdown(
                f'<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                f'⚙ Parsing <strong style="color:#e6edf3;">{fname}</strong> '
                f'<span style="color:#6e7681;">({file_size_kb} KB)</span></div>',
                unsafe_allow_html=True,
            )

        try:
            # ── PDF ──────────────────────────────────────────────────────────
            if flower.endswith(".pdf") or "pdf" in ftype_mime:
                summary["file_type"] = "pdf"
                page_count = 0
                method = "none"
                combined_pdf = ""
                text_before_vision = 0
                diagnostics: list[str] = []
                pymupdf_error: str | None = None
                ocr_error: str | None = None
                vision_error: str | None = None
                fitz_ok, fitz_err = _fitz_available()
                if not fitz_ok:
                    diagnostics.append(
                        f"PyMuPDF unavailable — using pypdf + pypdfium2 ({fitz_err[:60]})"
                    )

                # 1) PyMuPDF — multiple text modes, pick richest result
                try:
                    combined_pdf, page_count, method, pymupdf_error = (
                        _extract_pdf_text_pymupdf(file_bytes)
                    )
                    if pymupdf_error and not combined_pdf.strip():
                        diagnostics.append(f"PyMuPDF: {pymupdf_error[:90]}")
                except Exception as exc:
                    combined_pdf = ""
                    pymupdf_error = str(exc)
                    diagnostics.append(f"PyMuPDF: {pymupdf_error[:90]}")

                if not page_count:
                    page_count = _pdf_page_count(file_bytes)

                # 2) pypdf fallback if PyMuPDF empty or weak
                if _meaningful_text_len(combined_pdf) < PDF_MIN_TEXT_CHARS:
                    try:
                        pypdf_text, pypdf_pages, pypdf_err = _extract_pdf_text_pypdf(file_bytes)
                        if pypdf_err and not pypdf_text.strip():
                            diagnostics.append(f"pypdf: {pypdf_err[:90]}")
                        if _meaningful_text_len(pypdf_text) > _meaningful_text_len(combined_pdf):
                            combined_pdf = pypdf_text
                            page_count = pypdf_pages or page_count
                            method = "pypdf"
                    except ImportError:
                        diagnostics.append("pypdf not installed")
                    except Exception as exc:
                        diagnostics.append(f"pypdf: {str(exc)[:90]}")

                text_before_vision = _meaningful_text_len(combined_pdf)

                # 3) OCR if still too short (skipped gracefully when Tesseract/DLL unavailable)
                if text_before_vision < PDF_MIN_TEXT_CHARS:
                    ocr_text, ocr_pages, ocr_error = _extract_pdf_text_ocr(file_bytes)
                    if ocr_error and not ocr_text.strip():
                        diagnostics.append(f"OCR skipped: {ocr_error[:90]}")
                    if _meaningful_text_len(ocr_text) > text_before_vision:
                        combined_pdf = ocr_text
                        if not page_count:
                            page_count = ocr_pages
                        method = "tesseract-ocr"
                        text_before_vision = _meaningful_text_len(combined_pdf)
                        if status_slot:
                            status_slot.markdown(
                                f'<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                                f'📷 OCR fallback for <strong>{_esc(fname)}</strong> '
                                f'({text_before_vision:,} chars)</div>',
                                unsafe_allow_html=True,
                            )

                # 4) Vision LLM fallback for image-heavy pitch decks
                vision_block = ""
                vision_pages_ok = 0
                if text_before_vision < PDF_MIN_TEXT_CHARS and api_key:
                    if status_slot:
                        status_slot.markdown(
                            f'<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                            f'👁 Vision analysis for <strong>{_esc(fname)}</strong> '
                            f'(only {text_before_vision:,} chars from text/OCR)…</div>',
                            unsafe_allow_html=True,
                        )
                    vision_block, vision_pages_ok, vision_error = _vision_describe_pdf(
                        file_bytes, fname, api_key, base_url, model,
                    )
                    if vision_error and not vision_block.strip():
                        diagnostics.append(f"Vision: {vision_error[:100]}")
                    elif vision_error and vision_block.strip():
                        diagnostics.append(f"Vision partial: {vision_error[:80]}")

                    if vision_block.strip():
                        base_method = method if method != "none" else "pdf-vision-llm"
                        method = (
                            "pdf-vision-llm"
                            if base_method == "pdf-vision-llm"
                            else f"{base_method}+vision-llm"
                        )
                    elif vision_error:
                        vision_error = vision_error[:200]

                if vision_block.strip():
                    prefix = combined_pdf.strip()
                    if prefix:
                        combined_pdf = (
                            f"{prefix}\n\n[AI Pitch Deck Vision Analysis]\n{vision_block}"
                        )
                    else:
                        combined_pdf = f"[AI Pitch Deck Vision Analysis]\n{vision_block}"
                elif not combined_pdf.strip():
                    if api_key and vision_error:
                        short_err = vision_error.split(";")[0][:80]
                        combined_pdf = (
                            "[Image-based PDF — limited text extracted. "
                            f"Vision analysis unavailable ({short_err}). "
                            "Add submission notes or retry with a text-based export.]"
                        )
                        method = "image-pdf-limited"
                    elif api_key:
                        combined_pdf = (
                            "[Image-based PDF — limited text extracted. "
                            "No embedded text found; add submission notes for richer scoring.]"
                        )
                        method = "image-pdf-limited"
                    else:
                        combined_pdf = (
                            "[Image-based PDF — limited text extracted. "
                            "Enable Real LLM mode with an API key for AI slide analysis, "
                            "or add submission notes.]"
                        )
                        method = "image-pdf-limited"
                elif text_before_vision < PDF_MIN_TEXT_CHARS:
                    combined_pdf = (
                        f"{combined_pdf.strip()}\n\n"
                        "[Note: Partial text only — this PDF appears image-heavy. "
                        "Scoring uses available embedded text.]"
                    )

                preview_raw = combined_pdf.replace("\n", " ")
                _pdf_method_labels = {
                    "PyMuPDF-blocks": "PyMuPDF", "PyMuPDF-text": "PyMuPDF",
                    "PyMuPDF-dict": "PyMuPDF", "PyMuPDF-words": "PyMuPDF",
                    "PyMuPDF-html": "PyMuPDF", "PyMuPDF-raw": "PyMuPDF",
                    "PyMuPDF-text-flags": "PyMuPDF", "pypdf": "pypdf",
                    "tesseract-ocr": "OCR", "pdf-vision-llm": "Vision",
                    "image-pdf-limited": "Image PDF (limited)",
                }
                method_display = _pdf_method_labels.get(method, method)
                label = (
                    f"[PDF: {fname} · {page_count or '?'} page(s) · {method_display}]"
                )
                parts.append(f"{label}\n{combined_pdf}")
                summary.update({
                    "extraction_method": method,
                    "chars": len(combined_pdf),
                    "text_chars_before_vision": text_before_vision,
                    "vision_pages": vision_pages_ok,
                    "pages": page_count or None,
                    "preview": preview_raw[:PDF_EXTRACTION_PREVIEW_CHARS],
                    "vision_error": vision_error or "",
                    "extraction_diagnostics": " · ".join(diagnostics)[:240],
                })
                status_detail = f"{summary['chars']:,} chars"
                if text_before_vision and "vision" in method.lower():
                    status_detail = (
                        f"{text_before_vision:,} text + vision → {summary['chars']:,} chars"
                    )
                elif method == "image-pdf-limited":
                    status_detail = f"{summary['chars']:,} chars (limited / image PDF)"
                if status_slot:
                    status_slot.markdown(
                        f'<div style="font-size:12px;color:#6e7681;padding:2px 0;">'
                        f'✓ PDF <strong>{_esc(fname)}</strong>: {status_detail} '
                        f'via {_esc(method_display)}</div>',
                        unsafe_allow_html=True,
                    )

            # ── Image ─────────────────────────────────────────────────────────
            elif any(flower.endswith(x) for x in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
                summary["file_type"] = "image"
                from PIL import Image as _PILImage
                import base64 as _b64

                img = _PILImage.open(_io.BytesIO(file_bytes))
                w, h = img.size
                fmt  = img.format or flower.rsplit(".", 1)[-1].upper()

                # EXIF
                exif_lines: list = []
                try:
                    raw_exif = img._getexif()
                    if raw_exif:
                        from PIL.ExifTags import TAGS as _TAGS
                        wanted = {"Make", "Model", "DateTime", "Software"}
                        for tid, val in raw_exif.items():
                            tname = _TAGS.get(tid, "")
                            if tname in wanted:
                                exif_lines.append(f"{tname}: {str(val)[:80]}")
                except Exception:
                    pass

                # Thumbnail (for UI display)
                thumb = img.copy()
                thumb.thumbnail((240, 180), _PILImage.LANCZOS)
                if thumb.mode not in ("RGB", "L"):
                    thumb = thumb.convert("RGB")
                buf = _io.BytesIO()
                thumb.save(buf, format="JPEG", quality=75)
                thumb_b64 = _b64.b64encode(buf.getvalue()).decode()

                exif_str   = "; ".join(exif_lines) if exif_lines else "none"
                meta_text  = (
                    f"Image: {fname}\n"
                    f"Dimensions: {w}×{h} px  Format: {fmt}\n"
                    f"File size: {file_size_kb} KB\n"
                    f"EXIF: {exif_str}"
                )

                # LLM vision analysis
                vision_text = ""
                if use_llm_vision and api_key:
                    try:
                        vision_text = _vision_describe_image(file_bytes, fname, api_key, base_url, model)
                        method = "vision_llm+pillow"
                    except Exception as ve:
                        vision_text = f"(Vision analysis unavailable: {ve})"
                        method = "pillow_metadata"
                else:
                    method = "pillow_metadata"

                full_img_text = meta_text + (f"\n\nAI Visual Analysis:\n{vision_text}" if vision_text else "")
                parts.append(f"[Image: {fname}]\n{full_img_text}")
                summary.update({
                    "extraction_method": method,
                    "chars": len(full_img_text),
                    "dimensions": f"{w}×{h}",
                    "thumbnail_b64": thumb_b64,
                    "preview": (vision_text[:300] if vision_text and "unavailable" not in vision_text
                                else meta_text[:200]),
                })

            # ── Video ─────────────────────────────────────────────────────────
            elif any(flower.endswith(x) for x in (".mp4", ".mov", ".avi", ".webm", ".mkv")):
                summary["file_type"] = "video"
                ext_label = flower.rsplit(".", 1)[-1].upper()
                meta_text = (
                    f"Video file: {fname}\n"
                    f"Format: {ext_label}  File size: {file_size_kb} KB\n"
                    "Note: Full transcription not available. Score this submission based on "
                    "the idea name, notes, and any accompanying documents."
                )
                frame_note = ""
                thumb_b64  = None

                # Try cv2 for first-frame extraction (optional dependency)
                try:
                    import cv2 as _cv2
                    import tempfile, os as _os, base64 as _b64
                    from PIL import Image as _PILImage
                    with tempfile.NamedTemporaryFile(
                        suffix=f".{flower.rsplit('.', 1)[-1]}", delete=False
                    ) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    cap = _cv2.VideoCapture(tmp_path)
                    fps         = cap.get(_cv2.CAP_PROP_FPS) or 25
                    total_fr    = int(cap.get(_cv2.CAP_PROP_FRAME_COUNT))
                    duration_s  = round(total_fr / fps, 1) if total_fr else 0
                    vid_w       = int(cap.get(_cv2.CAP_PROP_FRAME_WIDTH))
                    vid_h       = int(cap.get(_cv2.CAP_PROP_FRAME_HEIGHT))
                    ret, frame  = cap.read()
                    cap.release()
                    _os.unlink(tmp_path)
                    meta_text += f"\nResolution: {vid_w}×{vid_h}  Duration: {duration_s}s  FPS: {round(fps,1)}"
                    if ret:
                        rgb   = _cv2.cvtColor(frame, _cv2.COLOR_BGR2RGB)
                        pil_f = _PILImage.fromarray(rgb)
                        pil_f.thumbnail((240, 180), _PILImage.LANCZOS)
                        buf_f = _io.BytesIO()
                        pil_f.save(buf_f, format="JPEG", quality=75)
                        thumb_b64 = _b64.b64encode(buf_f.getvalue()).decode()
                        if use_llm_vision and api_key:
                            try:
                                buf_vis = _io.BytesIO()
                                pil_f.save(buf_vis, format="JPEG", quality=85)
                                frame_note = "\n\nFirst-frame visual analysis:\n" + _vision_describe_image(
                                    buf_vis.getvalue(), f"{fname}_frame.jpg", api_key, base_url, model
                                )
                            except Exception:
                                frame_note = "\n(First-frame vision analysis unavailable)"
                except ImportError:
                    pass
                except Exception:
                    pass

                full_vid_text = meta_text + frame_note
                parts.append(f"[Video: {fname}]\n{full_vid_text}")
                summary.update({
                    "extraction_method": "cv2+vision" if frame_note else ("cv2" if thumb_b64 else "metadata"),
                    "chars": len(full_vid_text),
                    "thumbnail_b64": thumb_b64,
                    "preview": full_vid_text[:200],
                })

            # ── Text / document ──────────────────────────────────────────────
            elif any(flower.endswith(x) for x in (".txt", ".md", ".csv", ".docx", ".doc")):
                summary["file_type"] = "text"
                if flower.endswith(".docx"):
                    try:
                        import docx as _docx
                        doc  = _docx.Document(_io.BytesIO(file_bytes))
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        parts.append(f"[Document: {fname}]\n{text[:6000]}")
                        summary.update({
                            "extraction_method": "python-docx",
                            "chars": len(text),
                            "preview": text[:300],
                        })
                    except ImportError:
                        decoded = file_bytes.decode("utf-8", errors="replace")[:4000]
                        parts.append(f"[Document: {fname}]\n{decoded}")
                        summary.update({"extraction_method": "utf8", "chars": len(decoded), "preview": decoded[:300]})
                else:
                    decoded = file_bytes.decode("utf-8", errors="replace")[:6000]
                    parts.append(f"[Text: {fname}]\n{decoded}")
                    summary.update({"extraction_method": "utf8", "chars": len(decoded), "preview": decoded[:300]})

            # ── Fallback ─────────────────────────────────────────────────────
            else:
                try:
                    decoded = file_bytes.decode("utf-8", errors="replace")[:2000]
                    parts.append(f"[File: {fname}]\n{decoded}")
                    summary.update({"extraction_method": "utf8", "chars": len(decoded), "preview": decoded[:200]})
                except Exception:
                    parts.append(f"[File: {fname}] (binary — cannot extract text)")

        except Exception as outer_err:
            parts.append(f"[File: {fname}] (error: {outer_err})")
            summary["preview"] = f"Extraction error: {outer_err}"

        summaries.append(summary)

        if progress_bar:
            progress_bar.progress(idx / max(total_files, 1))

    return "\n\n".join(parts), summaries


_STAGE_RECOMMENDER_SYSTEM = """You are ForgeOS, an expert intake analyst for physical-goods innovation submissions.
Recommend the best starting stage in this 7-stage pipeline:
- Intake
- Concept (this is shown to users as "Concept Refinement")
- Validation
- Prototyping
- Market Test
- Scaling
- Monitoring

Return ONLY valid JSON with this exact schema:
{
  "recommended_stage": "<one stage from the list above>",
  "confidence": <integer 50-95>,
  "explanation": "<1-2 sentence rationale grounded in the supplied evidence>"
}

Rules:
- Be conservative and prefer an earlier stage when evidence is mixed.
- Base the decision only on the supplied notes, extracted content, and file summaries.
- Use "Concept" for concept-stage recommendations, not "Concept Refinement".
- Do not invent traction, prototypes, regulatory progress, or launch status."""


def _uploaded_files_signature(uploaded_files, *, use_vis: bool = False) -> str:
    """Stable signature for uploaded files to avoid reparsing unchanged drafts."""
    if not uploaded_files:
        return "no-files"
    parts: list[str] = []
    for file_obj in uploaded_files:
        try:
            file_bytes = file_obj.getvalue()
        except Exception:
            file_obj.seek(0)
            file_bytes = file_obj.read()
            file_obj.seek(0)
        digest = hashlib.md5(file_bytes[:65536]).hexdigest()
        size = getattr(file_obj, "size", len(file_bytes))
        parts.append(f"{file_obj.name}:{size}:{file_obj.type}:{digest}")
    payload = f"{use_vis}|{'|'.join(parts)}".encode("utf-8")
    return hashlib.md5(payload).hexdigest()


def _build_stage_recommendation_context(
    idea_name: str,
    notes_txt: str,
    extracted_text: str,
    file_summaries: list[dict] | None = None,
) -> str:
    """Compact context payload for the intake-stage recommender."""
    file_summaries = file_summaries or []
    summary_lines = []
    for item in file_summaries[:6]:
        summary_lines.append(
            f"- {item.get('name', 'file')} | {item.get('file_type', 'other')} | "
            f"method={item.get('extraction_method', 'none')} | chars={item.get('chars', 0)}"
        )
    lines = [
        f"IDEA NAME: {idea_name.strip()}",
        f"NOTES: {notes_txt.strip()[:1200]}",
        "",
        "FILES:",
        "\n".join(summary_lines) if summary_lines else "- none",
        "",
        "EXTRACTED CONTENT:",
        extracted_text[:4000] if extracted_text else "(none)",
    ]
    return "\n".join(lines)


def _recommend_starting_stage_llm(
    idea_name: str,
    notes_txt: str,
    extracted_text: str,
    file_summaries: list[dict] | None = None,
) -> dict:
    """LLM-based stage recommendation with compact JSON output."""
    raw = _llm_chat_text(
        _STAGE_RECOMMENDER_SYSTEM,
        _build_stage_recommendation_context(idea_name, notes_txt, extracted_text, file_summaries),
        _effective_llm_key(),
        _FORGE_LLM_BASE_URL,
        _FORGE_LLM_MODEL,
        temperature=0.2,
        max_tokens=220,
        timeout=22,
    )
    parsed = json.loads(raw)
    stage_name = _normalize_stage_name(parsed.get("recommended_stage", "Intake"))
    confidence = int(parsed.get("confidence", 65))
    confidence = max(50, min(95, confidence))
    explanation = str(parsed.get("explanation", "")).strip()
    if not explanation:
        explanation = "The submission appears most aligned with this stage based on the evidence provided."
    return {
        "recommended_stage": stage_name,
        "confidence": confidence,
        "explanation": explanation,
        "source": "llm",
    }


def _recommend_starting_stage_local(
    idea_name: str,
    notes_txt: str,
    extracted_text: str,
    file_summaries: list[dict] | None = None,
) -> dict:
    """Fast heuristic fallback for intake-stage recommendation."""
    text = f"{idea_name}\n{notes_txt}\n{extracted_text}".lower()
    text = re.sub(r"\s+", " ", text).strip()
    file_summaries = file_summaries or []
    total_chars = len(extracted_text or "") + len(notes_txt or "")
    if not text:
        return {
            "recommended_stage": "Intake",
            "confidence": 58,
            "explanation": (
                "Very limited information is available, so Intake is the safest default until "
                "supporting material or clearer development evidence is provided."
            ),
            "source": "heuristic",
        }

    stage_patterns: dict[str, list[tuple[str, int]]] = {
        "Monitoring": [
            ("post-launch", 6), ("post launch", 6), ("live customers", 6), ("retention", 5),
            ("repeat purchase", 5), ("nps", 4), ("customer success", 4), ("installed base", 5),
        ],
        "Scaling": [
            ("mass production", 6), ("scale-up", 5), ("scaling", 5), ("retail rollout", 6),
            ("distribution", 4), ("purchase order", 5), ("manufacturing line", 5),
            ("capacity expansion", 6), ("co-manufacturer", 4),
        ],
        "Market Test": [
            ("pilot", 5), ("beta", 4), ("pre-order", 5), ("preorder", 5), ("waitlist", 4),
            ("traction", 5), ("paid trial", 5), ("sell-through", 5), ("customer interviews", 3),
            ("early customers", 5), ("test market", 5), ("loi", 4), ("letters of intent", 4),
        ],
        "Prototyping": [
            ("prototype", 6), ("bom", 5), ("bill of materials", 6), ("cad", 5), ("3d printed", 5),
            ("sample build", 5), ("tooling", 5), ("supplier quote", 5), ("alpha build", 5),
            ("bench test", 4), ("manufacturing drawing", 5),
        ],
        "Validation": [
            ("proof of concept", 5), ("poc", 4), ("feasibility", 4), ("market research", 4),
            ("survey", 3), ("validated", 4), ("lab test", 4), ("regulatory review", 4),
            ("customer discovery", 4), ("unit economics", 3),
        ],
        "Concept": [
            ("concept", 3), ("idea", 2), ("vision", 2), ("hypothesis", 3),
            ("problem statement", 3), ("opportunity", 2), ("white space", 3), ("roadmap", 2),
        ],
    }
    scores = {stage: 0 for stage in STAGE_NAMES}
    matches: dict[str, list[str]] = {stage: [] for stage in STAGE_NAMES}

    for stage_name, patterns in stage_patterns.items():
        for keyword, weight in patterns:
            if keyword in text:
                scores[stage_name] += weight
                matches[stage_name].append(keyword)

    if file_summaries:
        has_pdf = any(item.get("file_type") == "pdf" for item in file_summaries)
        has_media = any(item.get("file_type") in ("image", "video") for item in file_summaries)
        if has_pdf:
            scores["Concept"] += 2
        if has_media:
            scores["Prototyping"] += 2

    if total_chars > 1200:
        scores["Validation"] += 1
        scores["Concept"] += 1
    if total_chars > 2800:
        scores["Prototyping"] += 1
        scores["Market Test"] += 1

    ordered = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    top_stage, top_score = ordered[0]
    runner_up_score = ordered[1][1] if len(ordered) > 1 else 0

    if top_score <= 0:
        top_stage = "Concept" if total_chars >= 120 else "Intake"
        top_score = 1 if top_stage == "Concept" else 0

    if top_stage == "Monitoring" and top_score < 6:
        top_stage = "Scaling"
    if top_stage == "Scaling" and top_score < 5:
        top_stage = "Market Test"
    if top_stage == "Market Test" and not any(
        kw in text for kw in ("pilot", "beta", "pre-order", "preorder", "traction", "early customers", "sell-through")
    ):
        top_stage = "Validation" if scores["Validation"] >= scores["Prototyping"] else "Prototyping"

    evidence_terms = matches.get(top_stage, [])[:3]
    confidence = 62 + min(24, top_score * 3 + max(0, top_score - runner_up_score) * 2)
    confidence = max(55, min(92, confidence))

    if top_stage == "Intake":
        explanation = (
            "The submission is still light on development evidence, so Intake is the most prudent "
            "starting point before moving into formal concept work."
        )
    elif top_stage == "Concept":
        explanation = (
            "The material reads as an early-stage concept with a defined opportunity, but it does not "
            "yet show enough validation or prototype-ready detail to start further down the pipeline."
        )
    elif top_stage == "Validation":
        explanation = (
            "The submission goes beyond a raw idea and shows early validation activity, but it still "
            "needs stronger proof before prototype execution or market testing."
        )
    elif top_stage == "Prototyping":
        explanation = (
            "The materials reference prototype-level development signals such as BOMs, samples, CAD, "
            "or supplier/tooling work, which indicates readiness for hands-on build iteration."
        )
    elif top_stage == "Market Test":
        explanation = (
            "The idea appears to have progressed into customer or channel testing, with evidence of pilots, "
            "pre-orders, early customers, or traction signals that justify a market-test entry point."
        )
    elif top_stage == "Scaling":
        explanation = (
            "The submission suggests the product is beyond testing and is focused on production scale, "
            "distribution growth, or commercialization readiness."
        )
    else:
        explanation = (
            "The business appears post-launch and operational, with language consistent with active "
            "performance tracking rather than initial development work."
        )

    if evidence_terms:
        explanation += f" Key signals detected: {', '.join(evidence_terms)}."

    return {
        "recommended_stage": top_stage,
        "confidence": int(confidence),
        "explanation": explanation,
        "source": "heuristic",
    }


def recommend_starting_stage(
    idea_name: str,
    notes_txt: str,
    extracted_text: str,
    file_summaries: list[dict] | None = None,
    *,
    allow_llm: bool = True,
) -> tuple[dict, str | None]:
    """Recommend an intake stage using LLM when helpful, else local heuristics."""
    if allow_llm and _llm_ready() and (extracted_text or notes_txt):
        try:
            return _recommend_starting_stage_llm(idea_name, notes_txt, extracted_text, file_summaries), None
        except Exception as exc:
            return _recommend_starting_stage_local(idea_name, notes_txt, extracted_text, file_summaries), (
                "Stage recommendation used ForgeOS heuristic analysis because live AI stage triage was unavailable "
                f"({ _classify_llm_error(exc) })."
            )
    return _recommend_starting_stage_local(idea_name, notes_txt, extracted_text, file_summaries), None


def _stage_analysis_cache_key(
    idea_name: str,
    notes_txt: str,
    uploaded,
    use_vis: bool,
    *,
    cache_scope: str = "single",
) -> str:
    idea_name = (idea_name or "").strip()
    notes_txt = (notes_txt or "").strip()
    if uploaded:
        file_sig = _uploaded_files_signature(uploaded, use_vis=use_vis)
        return f"{cache_scope}|{idea_name}|{notes_txt}|{file_sig}"
    return f"{cache_scope}|{idea_name}|{notes_txt}|notes-only"


def _get_stage_analysis(
    idea_name: str,
    notes_txt: str,
    uploaded,
    use_vis: bool,
    *,
    cache_scope: str = "single",
    show_spinner: bool = True,
    require_explicit_run: bool = False,
) -> dict | None:
    """Analyze draft content and return a cached stage recommendation."""
    idea_name = (idea_name or "").strip()
    notes_txt = (notes_txt or "").strip()
    has_files = bool(uploaded)
    if not has_files and not idea_name and not notes_txt:
        if cache_scope == "single":
            st.session_state.single_stage_analysis = None
        return None

    cache_key = _stage_analysis_cache_key(
        idea_name, notes_txt, uploaded, use_vis, cache_scope=cache_scope,
    )
    if require_explicit_run:
        if cache_scope == "single":
            requested_key = st.session_state.get("single_stage_analysis_request_sig")
            if requested_key != cache_key:
                st.session_state.single_stage_analysis = None
                return None
        else:
            requested_keys = st.session_state.bulk_stage_analysis_cache.get("__requested__", set())
            if cache_key not in requested_keys:
                st.session_state.bulk_stage_analysis_cache.pop(cache_key, None)
                return None
    if cache_scope == "single":
        cached = st.session_state.get("single_stage_analysis")
        if cached and cached.get("cache_key") == cache_key:
            return cached
    else:
        cached = st.session_state.bulk_stage_analysis_cache.get(cache_key)
        if cached:
            return cached

    spinner_label = (
        "Analyzing uploaded material and recommending a starting stage…"
        if has_files
        else "Analyzing idea details and recommending a starting stage…"
    )
    spinner_ctx = st.spinner(spinner_label) if show_spinner else nullcontext()
    with spinner_ctx:
        progress_bar = None
        progress_label = None
        if has_files and show_spinner:
            progress_bar = st.progress(0.0)
            progress_label = st.empty()
        if has_files:
            file_sig = _uploaded_files_signature(uploaded, use_vis=use_vis)
            extracted_text, file_summaries = extract_file_text(
                uploaded,
                progress_bar=progress_bar,
                progress_label=progress_label,
                use_llm_vision=(use_vis and _llm_ready() and st.session_state.get("scoring_mode", "Simulated AI") == "Real LLM"),
                api_key=_effective_llm_key(),
                base_url=_FORGE_LLM_BASE_URL,
                model=_FORGE_LLM_MODEL,
            )
            if progress_bar:
                progress_bar.progress(1.0)
            if progress_label:
                progress_label.markdown(
                    '<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                    'Finalizing recommended stage…</div>',
                    unsafe_allow_html=True,
                )
            recommendation, warning = recommend_starting_stage(
                idea_name,
                notes_txt,
                extracted_text,
                file_summaries,
                allow_llm=True,
            )
            if progress_label:
                progress_label.empty()
            result = {
                "cache_key": cache_key,
                "file_signature": file_sig,
                "extracted_text": extracted_text,
                "file_summaries": file_summaries,
                "recommendation": recommendation,
                "warning": warning,
                "file_count": len(uploaded or []),
                "total_chars": sum(item.get("chars", 0) for item in file_summaries),
            }
        else:
            recommendation, warning = recommend_starting_stage(
                idea_name,
                notes_txt,
                "",
                [],
                allow_llm=False,
            )
            result = {
                "cache_key": cache_key,
                "file_signature": "notes-only",
                "extracted_text": "",
                "file_summaries": [],
                "recommendation": recommendation,
                "warning": warning,
                "file_count": 0,
                "total_chars": len(notes_txt),
            }

    if cache_scope == "single":
        st.session_state.single_stage_analysis = result
    else:
        st.session_state.bulk_stage_analysis_cache[cache_key] = result
    return result


def _get_single_stage_analysis(
    idea_name: str,
    notes_txt: str,
    uploaded,
    use_vis: bool,
) -> dict | None:
    return _get_stage_analysis(
        idea_name,
        notes_txt,
        uploaded,
        use_vis,
        cache_scope="single",
        require_explicit_run=bool(uploaded),
    )


def _render_stage_recommendation_card(stage_analysis: dict) -> None:
    rec = stage_analysis.get("recommendation", {})
    rec_stage = _normalize_stage_name(rec.get("recommended_stage", "Intake"))
    rec_display = _stage_display_name(rec_stage)
    rec_conf = int(rec.get("confidence", 0))
    rec_expl = rec.get("explanation", "")
    rec_color = _stage_color(rec_stage)
    rec_source = stage_analysis.get("warning") or (
        "Recommendation based on extracted submission evidence."
        if stage_analysis.get("file_count")
        else "Recommendation based on the idea name and notes provided."
    )
    st.markdown(
        f"""
        <div style="margin-top:10px;padding:14px 16px;border-radius:14px;
                    border:1px solid {rec_color}33;background:linear-gradient(180deg,#111826 0%,#0f1724 100%);
                    box-shadow:0 10px 24px rgba(0,0,0,0.18);">
          <div>
            <div style="font-size:11px;font-weight:700;letter-spacing:0.08em;text-transform:uppercase;color:#8b949e;">
              Auto-Assign Stage
            </div>
            <div style="margin-top:6px;font-size:20px;font-weight:800;color:{rec_color};">
              Recommended Starting Stage: { _esc(rec_display) } ({rec_conf}% confidence)
            </div>
            <div style="margin-top:8px;font-size:13px;line-height:1.6;color:#c9d1d9;max-width:880px;">
              { _esc(rec_expl) }
            </div>
          </div>
          <div style="margin-top:10px;font-size:12px;color:#8b949e;">
            { _esc(rec_source) }
          </div>
          <div style="margin-top:8px;font-size:12px;color:#8b949e;">
            This stage will be assigned automatically when you submit.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _append_submission_from_upload(
    idea_name: str,
    notes_txt: str,
    uploaded,
    use_vis: bool,
    *,
    status_slot=None,
    precomputed_extracted: str | None = None,
    precomputed_summaries: list[dict] | None = None,
    stage_recommendation: dict | None = None,
    manual_stage_override: bool = False,
    fallback_stage: str = "Intake",
) -> tuple[str, int, int]:
    """Create and append one submission. Returns (sid, file_count, chars_extracted)."""
    ftypes = list({f.type.split("/")[-1].upper() for f in (uploaded or [])}) or ["—"]
    extracted = precomputed_extracted if precomputed_extracted is not None else ""
    file_summaries: list = list(precomputed_summaries or [])
    if uploaded and precomputed_extracted is None and precomputed_summaries is None:
        _is_llm_mode = st.session_state.get("scoring_mode", "Simulated AI") == "Real LLM"
        extracted, file_summaries = extract_file_text(
            uploaded,
            status_slot=status_slot,
            use_llm_vision=(use_vis and _llm_ready() and _is_llm_mode),
            api_key=_effective_llm_key(),
            base_url=_FORGE_LLM_BASE_URL,
            model=_FORGE_LLM_MODEL,
        )
    if stage_recommendation is None:
        stage_recommendation, _ = recommend_starting_stage(
            idea_name,
            notes_txt,
            extracted,
            file_summaries,
            allow_llm=True,
        )
    assigned_stage = _normalize_stage_name(fallback_stage)
    if stage_recommendation and not manual_stage_override:
        assigned_stage = _normalize_stage_name(stage_recommendation.get("recommended_stage", assigned_stage))
    sid = f"FOS-{st.session_state.next_id}"
    st.session_state.next_id += 1
    total_chars = sum(s.get("chars", 0) for s in file_summaries)
    st.session_state.submissions.append({
        "id":             sid,
        "name":           idea_name.strip(),
        "file_type":      ", ".join(ftypes),
        "status":         "New",
        "stage":          assigned_stage,
        "overall":        0.0,
        "innovation":     0.0,
        "feasibility":    0.0,
        "categories":     {},
        "auto_reject":    [],
        "high_risk":      [],
        "scored_at":      "",
        "stage_summary":  "",
        "stage_history":  [{"stage": assigned_stage, "moved_at": datetime.now().strftime("%Y-%m-%d")}],
        "submitted_at":   datetime.now().strftime("%Y-%m-%d"),
        "notes":          notes_txt,
        "extracted_text": extracted,
        "file_summaries": file_summaries,
        "extraction_note": _build_extraction_note(file_summaries),
        "extraction_preview": extracted[:PDF_EXTRACTION_PREVIEW_CHARS] if extracted else "",
        "intake_stage_recommendation": stage_recommendation or {},
        "intake_stage_manual_override": manual_stage_override,
    })
    return sid, len(uploaded or []), total_chars


def _render_upload_file_previews(uploaded, max_previews: int = 5) -> None:
    """Thumbnail / icon previews for files selected before submit."""
    if not uploaded:
        return
    import io as _preview_io
    from PIL import Image as _PreviewImage

    n_prev = min(len(uploaded), max_previews)
    prev_cols = st.columns(n_prev)
    _type_icons = {
        "pdf": "📄", "mp4": "🎬", "mov": "🎬", "avi": "🎬",
        "webm": "🎬", "txt": "📝", "md": "📝", "csv": "📊", "docx": "📝",
    }
    for pi, pf in enumerate(uploaded[:n_prev]):
        with prev_cols[pi]:
            pn = pf.name.lower()
            pf.seek(0)
            if any(pn.endswith(x) for x in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
                try:
                    raw = pf.read()
                    img_prev = _PreviewImage.open(_preview_io.BytesIO(raw))
                    st.image(img_prev, caption=pf.name[:22], use_container_width=True)
                except Exception:
                    st.markdown(f"🖼 `{pf.name[:20]}`")
            else:
                ext_p = pn.rsplit(".", 1)[-1] if "." in pn else "file"
                icon_p = _type_icons.get(ext_p, "📎")
                sz_p = round(pf.size / 1024, 1) if hasattr(pf, "size") else "?"
                st.markdown(
                    f'<div style="border:1px solid #21262d;border-radius:6px;'
                    f'padding:8px 10px;text-align:center;font-size:12px;color:#8b949e;">'
                    f'<div style="font-size:22px;margin-bottom:4px;">{icon_p}</div>'
                    f'<div style="color:#e6edf3;font-weight:600;word-break:break-all;">{pf.name[:20]}</div>'
                    f'<div style="font-size:10px;margin-top:2px;">{ext_p.upper()} · {sz_p} KB</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            pf.seek(0)
    if len(uploaded) > n_prev:
        st.caption(f"+ {len(uploaded) - n_prev} more file(s)")


def _reset_bulk_upload_form() -> None:
    """Clear bulk-upload row widgets after a successful batch submit."""
    old_ids = list(st.session_state.bulk_upload_row_ids)
    st.session_state.bulk_upload_row_ids = [0]
    st.session_state.bulk_upload_next_row_id = max(st.session_state.bulk_upload_next_row_id, 1)
    st.session_state.bulk_stage_analysis_cache = {}
    for rid in old_ids:
        for prefix in ("bulk_name_", "bulk_notes_", "bulk_files_"):
            st.session_state.pop(f"{prefix}{rid}", None)


def _ai_score_llm_pure(submission, rubric_data, api_key: str, base_url: str, model: str):
    """
    Pure LLM scoring — no st.session_state or module-level globals accessed.
    All config is passed in explicitly so this is safe to call from worker threads.
    Raises RuntimeError on API/parse failure (caller handles fallback).
    """
    criteria    = rubric_data.get("criteria", [])
    rubric_json = json.dumps(rubric_data, indent=2)

    system_prompt = f"""You are ForgeOS, an expert innovation scoring engine for physical goods companies.
You will score a product idea submission against the ForgeOS Extensive Rubric v2.

RUBRIC JSON:
{rubric_json}

INSTRUCTIONS:
- Score each criterion on a scale of 1.0–10.0 (one decimal place).
- Return ONLY a valid JSON object — no markdown, no extra text.
- The JSON must have this exact schema:
{{
  "criteria": {{
    "<criterion_name>": {{
      "score_10": <float 1.0-10.0>,
      "justification": "<2-3 sentence rubric-anchored justification>",
      "red_flags": [<list of triggered red flag strings, may be empty>],
      "gating_pass": <true if score meets the gating threshold for this criterion, else false>
    }}
  }},
  "weighted_total": <float 0-100>
}}
- Use the exact criterion names from the rubric.
- Weight the overall score as a weighted average using the weights in the rubric.
- Be discerning and critical — physical goods innovation is hard. Do not inflate scores.
- Base your scores on the idea name, notes, and any uploaded file content provided."""

    name           = submission.get("name", "")
    notes          = submission.get("notes", "")
    extracted_text = submission.get("extracted_text", "")

    user_parts = [f"IDEA NAME: {name}"]
    if notes.strip():
        user_parts.append(f"NOTES: {notes}")
    if extracted_text.strip():
        user_parts.append(f"UPLOADED FILE CONTENT:\n{extracted_text[:EXTRACTED_TEXT_LLM_LIMIT]}")
    user_message = "\n\n".join(user_parts)

    payload = json.dumps({
        "model":           model,
        "messages":        [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_message},
        ],
        "response_format": {"type": "json_object"},
        "temperature":     0.2,
    }).encode("utf-8")

    req = urllib.request.Request(
        url     = f"{base_url}/chat/completions",
        data    = payload,
        headers = {
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method  = "POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw_body   = resp.read().decode("utf-8")
            api_result = json.loads(raw_body)
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"LLM API HTTP {e.code}: {err_body[:300]}") from e
    except urllib.error.URLError as e:
        raise RuntimeError(f"LLM API connection error: {e.reason}") from e

    content  = api_result["choices"][0]["message"]["content"]
    llm_json = json.loads(content)

    llm_criteria    = llm_json.get("criteria", {})
    scored_criteria = {}

    GATE_MAP = {
        "Innovation & Novelty":                  (6.0, "auto-reject"),
        "Technical & Manufacturing Feasibility": (6.0, "auto-reject"),
        "Sustainability & Circularity":          (5.0, "high-risk"),
        "Evidence Quality & Realism":            (4.0, "auto-reject"),
    }

    for crit in criteria:
        key          = crit["criterion"]
        weight       = crit.get("weight", 10)
        llm_crit     = llm_criteria.get(key, {})
        score_10     = float(llm_crit.get("score_10", 5.0))
        score_10     = max(1.0, min(10.0, round(score_10, 1)))
        justification = llm_crit.get("justification", "")
        red_flags    = llm_crit.get("red_flags", [])

        if score_10 <= 3:
            band      = "1-3"
            anchor_txt = crit.get("scoring_anchors", {}).get("1-3", "Below threshold")
        elif score_10 <= 6:
            band      = "4-6"
            anchor_txt = crit.get("scoring_anchors", {}).get("4-6", "Moderate")
        else:
            band      = "7-10"
            anchor_txt = crit.get("scoring_anchors", {}).get("7-10", "Strong")

        evidence_level = "Sufficient" if score_10 >= 6 else ("Partial" if score_10 >= 4 else "Insufficient")

        scored_criteria[key] = {
            "name":          key,
            "score_10":      score_10,
            "score":         round(score_10 * 10),
            "weight":        weight,
            "weight_frac":   weight / 100.0,
            "anchor_band":   band,
            "anchor_text":   anchor_txt,
            "justification": justification,
            "evidence":      evidence_level,
            "evidence_req":  crit.get("evidence_required", ""),
            "red_flags":     red_flags,
            "sub_factors":   crit.get("sub_factors", []),
        }

    total_w = sum(v["weight_frac"] for v in scored_criteria.values())
    overall = round(
        sum(v["score"] * v["weight_frac"] for v in scored_criteria.values()) / max(total_w, 0.01),
        1,
    )
    overall = min(overall, 100.0)

    auto_reject_flags, high_risk_flags = [], []
    for crit_name, (threshold, action) in GATE_MAP.items():
        if crit_name in scored_criteria:
            s = scored_criteria[crit_name]["score_10"]
            if s < threshold:
                msg_g = f"{crit_name}: {s:.1f}/10 — below gate threshold of {threshold}/10"
                (auto_reject_flags if action == "auto-reject" else high_risk_flags).append(msg_g)

    return {
        "overall":     overall,
        "innovation":  scored_criteria.get("Innovation & Novelty", {}).get("score", 0),
        "feasibility": scored_criteria.get("Technical & Manufacturing Feasibility", {}).get("score", 0),
        "categories":  scored_criteria,
        "auto_reject": auto_reject_flags,
        "high_risk":   high_risk_flags,
        "scored_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def ai_score_submission_llm(submission, rubric_data):
    """
    Real LLM scoring using an OpenAI-compatible API (default: xAI Grok).
    Sends the full rubric as system context and the submission details as user message.
    Parses the structured JSON response and maps it to the same dict shape as
    ai_score_submission(). Falls back to ai_score_submission() on any error and
    returns a (result, warning_message) tuple.
    """
    criteria    = rubric_data.get("criteria", [])
    rubric_json = json.dumps(rubric_data, indent=2)

    system_prompt = f"""You are ForgeOS, an expert innovation scoring engine for physical goods companies.
You will score a product idea submission against the ForgeOS Extensive Rubric v2.

RUBRIC JSON:
{rubric_json}

INSTRUCTIONS:
- Score each criterion on a scale of 1.0–10.0 (one decimal place).
- Return ONLY a valid JSON object — no markdown, no extra text.
- The JSON must have this exact schema:
{{
  "criteria": {{
    "<criterion_name>": {{
      "score_10": <float 1.0-10.0>,
      "justification": "<2-3 sentence rubric-anchored justification>",
      "red_flags": [<list of triggered red flag strings, may be empty>],
      "gating_pass": <true if score meets the gating threshold for this criterion, else false>
    }}
  }},
  "weighted_total": <float 0-100>
}}
- Use the exact criterion names from the rubric.
- Weight the overall score as a weighted average using the weights in the rubric.
- Be discerning and critical — physical goods innovation is hard. Do not inflate scores.
- Base your scores on the idea name, notes, and any uploaded file content provided."""

    name           = submission.get("name", "")
    notes          = submission.get("notes", "")
    extracted_text = submission.get("extracted_text", "")

    user_parts = [f"IDEA NAME: {name}"]
    if notes.strip():
        user_parts.append(f"NOTES: {notes}")
    if extracted_text.strip():
        user_parts.append(f"UPLOADED FILE CONTENT:\n{extracted_text[:EXTRACTED_TEXT_LLM_LIMIT]}")
    user_message = "\n\n".join(user_parts)

    payload = json.dumps({
        "model":           _FORGE_LLM_MODEL,
        "messages":        [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_message},
        ],
        "response_format": {"type": "json_object"},
        "temperature":     0.2,
    }).encode("utf-8")

    req = urllib.request.Request(
        url     = f"{_FORGE_LLM_BASE_URL}/chat/completions",
        data    = payload,
        headers = {
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {_effective_llm_key()}",
        },
        method  = "POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw_body   = resp.read().decode("utf-8")
            api_result = json.loads(raw_body)
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"LLM API HTTP {e.code}: {err_body[:300]}") from e
    except urllib.error.URLError as e:
        raise RuntimeError(f"LLM API connection error: {e.reason}") from e

    content = api_result["choices"][0]["message"]["content"]
    llm_json = json.loads(content)

    llm_criteria  = llm_json.get("criteria", {})
    scored_criteria = {}

    GATE_MAP = {
        "Innovation & Novelty":                  (6.0, "auto-reject"),
        "Technical & Manufacturing Feasibility": (6.0, "auto-reject"),
        "Sustainability & Circularity":          (5.0, "high-risk"),
        "Evidence Quality & Realism":            (4.0, "auto-reject"),
    }

    for crit in criteria:
        key          = crit["criterion"]
        weight       = crit.get("weight", 10)
        llm_crit     = llm_criteria.get(key, {})
        score_10     = float(llm_crit.get("score_10", 5.0))
        score_10     = max(1.0, min(10.0, round(score_10, 1)))
        justification = llm_crit.get("justification", "")
        red_flags    = llm_crit.get("red_flags", [])

        if score_10 <= 3:
            band = "1-3"
            anchor_txt = crit.get("scoring_anchors", {}).get("1-3", "Below threshold")
        elif score_10 <= 6:
            band = "4-6"
            anchor_txt = crit.get("scoring_anchors", {}).get("4-6", "Moderate")
        else:
            band = "7-10"
            anchor_txt = crit.get("scoring_anchors", {}).get("7-10", "Strong")

        evidence_level = "Sufficient" if score_10 >= 6 else ("Partial" if score_10 >= 4 else "Insufficient")

        scored_criteria[key] = {
            "name":          key,
            "score_10":      score_10,
            "score":         round(score_10 * 10),
            "weight":        weight,
            "weight_frac":   weight / 100.0,
            "anchor_band":   band,
            "anchor_text":   anchor_txt,
            "justification": justification,
            "evidence":      evidence_level,
            "evidence_req":  crit.get("evidence_required", ""),
            "red_flags":     red_flags,
            "sub_factors":   crit.get("sub_factors", []),
        }

    total_w = sum(v["weight_frac"] for v in scored_criteria.values())
    overall = round(
        sum(v["score"] * v["weight_frac"] for v in scored_criteria.values()) / max(total_w, 0.01),
        1,
    )
    overall = min(overall, 100.0)

    auto_reject_flags, high_risk_flags = [], []
    for crit_name, (threshold, action) in GATE_MAP.items():
        if crit_name in scored_criteria:
            s = scored_criteria[crit_name]["score_10"]
            if s < threshold:
                msg = f"{crit_name}: {s:.1f}/10 — below gate threshold of {threshold}/10"
                (auto_reject_flags if action == "auto-reject" else high_risk_flags).append(msg)

    return {
        "overall":     overall,
        "innovation":  scored_criteria.get("Innovation & Novelty", {}).get("score", 0),
        "feasibility": scored_criteria.get("Technical & Manufacturing Feasibility", {}).get("score", 0),
        "categories":  scored_criteria,
        "auto_reject": auto_reject_flags,
        "high_risk":   high_risk_flags,
        "scored_at":   datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def route_scoring(submission, rubric_data):
    """
    Route scoring to either the real LLM or the simulated engine based on
    st.session_state.scoring_mode. Returns (score_dict, warning_str_or_None).
    warning_str is non-None when Real LLM was requested but fell back to Simulated.
    """
    mode = st.session_state.get("scoring_mode", "Simulated")
    if mode == "Real LLM":
        if not _llm_ready():
            return ai_score_submission(submission, rubric_data), "No API key available — fell back to Simulated scoring."
        try:
            result = ai_score_submission_llm(submission, rubric_data)
            return result, None
        except Exception as e:
            return ai_score_submission(submission, rubric_data), f"LLM error ({e}) — fell back to Simulated scoring."
    return ai_score_submission(submission, rubric_data), None


# ─── Criterion deep-dive (Expand) ────────────────────────────────────────────

CRITERION_DETAIL_TEXT_LIMIT = 5000
CRITERION_DETAIL_NOTES_LIMIT = 1500

_CRITERION_DETAIL_SYSTEM = """You are ForgeOS, a senior physical-goods innovation analyst advising investors and innovation directors.
Write an expert deep-dive for ONE rubric criterion only. Markdown only — use exactly these ## headings:

## Summary
Open by restating the assigned score's one-line justification (verbatim or near-verbatim). Add one sentence on what this score means for screening this idea.

## Score Rationale
Write 3–4 thoughtful paragraphs explaining WHY this exact score (1–10) was assigned. Use explicit chain-of-thought reasoning (e.g., "After reviewing the extracted content…", "The submission mentions X but lacks Y, which is why this scores N rather than N+2"). Reference weight %, evidence rating, sub-factors, and what evidence would raise or lower the score. Be analytical, not promotional.

## Rubric Anchor Analysis
For EACH band — **1–3**, **4–6**, and **7–10** — quote the supplied anchor language, then explain in 2–3 sentences whether the submission aligns, partially aligns, or fails to align. Conclude why the assigned score belongs in its band and not an adjacent band.

## Evidence from Submission
4–6 bullets citing specific passages from notes and extracted files. Use > blockquotes for direct quotes. If evidence is thin, state precisely what is missing.

## Flags & Highlights
List triggered red flags, rubric watch items, and standout positives for this criterion only.

Rules: use ONLY supplied data; do not invent revenue, customers, patents, or test results; professional investor-grade tone; 450–650 words total."""


def _classify_llm_error(exc: Exception) -> str:
    """Return 'timeout' or 'error' for user-facing fallback messaging."""
    import socket
    import urllib.error

    if isinstance(exc, socket.timeout):
        return "timeout"
    reason = getattr(exc, "reason", None)
    if isinstance(reason, socket.timeout):
        return "timeout"
    if isinstance(exc, urllib.error.URLError):
        r = str(reason or exc).lower()
        if "timed out" in r or "timeout" in r:
            return "timeout"
    msg = str(exc).lower()
    if "timed out" in msg or "timeout" in msg:
        return "timeout"
    return "error"


def _criterion_detail_fallback_message(exc: Exception | None = None) -> str:
    """Clean, professional notice when structured fallback is shown instead of LLM."""
    if exc is None:
        return (
            "Structured deep-dive generated from submission evidence and rubric anchors. "
            "Add an API key in Rubric Settings for AI-authored analysis."
        )
    if _classify_llm_error(exc) == "timeout":
        return (
            "AI analysis timed out — showing ForgeOS structured deep-dive below. "
            "Try **Regenerate** for a fresh attempt."
        )
    return (
        "AI analysis unavailable — showing ForgeOS structured deep-dive below."
    )


def _crit_key_slug(name: str) -> str:
    """Safe Streamlit widget key fragment from criterion name."""
    return re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")[:48]


def _criterion_detail_cache_key(sub_id: str, criterion_name: str) -> str:
    return f"{sub_id}::{criterion_name}"


def _close_memo_dialog() -> None:
    """Reset memo modal state without touching cached memo content."""
    st.session_state.memo_sub_id = None
    st.session_state.memo_needs_generate = False
    st.session_state.memo_force_regenerate = False


def _close_criterion_detail_dialog() -> None:
    """Close the criterion deep-dive modal."""
    _close_all_criterion_deep_dives()


def _criterion_panel_key(sub_id: str, criterion_name: str) -> str:
    """Unique key for one criterion deep-dive (session + widget identity)."""
    return f"deep_dive_{sub_id}_{_crit_key_slug(criterion_name)}"


def _is_criterion_deep_dive_open(sub_id: str, criterion_name: str) -> bool:
    panel_key = _criterion_panel_key(sub_id, criterion_name)
    return st.session_state.get("criterion_deep_dive_active") == panel_key


def _open_criterion_deep_dive(sub_id: str, criterion_name: str) -> None:
    """Open criterion deep-dive modal; clear memo so features stay isolated."""
    _close_memo_dialog()
    st.session_state.criterion_deep_dive_active = _criterion_panel_key(sub_id, criterion_name)
    st.session_state.criterion_deep_dive_sub_id = sub_id
    st.session_state.criterion_deep_dive_criterion = criterion_name
    st.session_state.criterion_detail_force = True


def _close_all_criterion_deep_dives() -> None:
    st.session_state.criterion_deep_dive_active = None
    st.session_state.criterion_deep_dive_sub_id = None
    st.session_state.criterion_deep_dive_criterion = None
    st.session_state.criterion_detail_force = False


def _close_criterion_deep_dives_for_submission(sub_id: str) -> None:
    if st.session_state.get("criterion_deep_dive_sub_id") == sub_id:
        _close_all_criterion_deep_dives()


def _clear_criterion_detail_cache(sub_id: str) -> None:
    """Drop cached deep-dives when a submission is re-scored or deleted."""
    prefix = f"{sub_id}::"
    for key in list(st.session_state.criterion_detail_cache.keys()):
        if key.startswith(prefix):
            del st.session_state.criterion_detail_cache[key]


def _score_band_from_10(score_10: float) -> str:
    if score_10 <= 3:
        return "1-3"
    if score_10 <= 6:
        return "4-6"
    return "7-10"


def _find_rubric_criterion(criterion_name: str, rubric_data: dict) -> dict:
    for crit in rubric_data.get("criteria", []):
        if crit.get("criterion") == criterion_name:
            return crit
    return {}


def _extract_evidence_snippets(
    text: str,
    keywords: list[str] | None = None,
    max_snippets: int = 6,
    min_len: int = 30,
) -> list[str]:
    """Pull quotable sentences from notes or extracted file content."""
    if not text or not text.strip():
        return []
    cleaned = re.sub(r"\[PDF:[^\]]+\]", "", text)
    cleaned = re.sub(r"\[Page \d+[^\]]*\]", "", cleaned)
    cleaned = re.sub(r"\[Image-based PDF[^\]]*\]", "", cleaned)
    cleaned = re.sub(r"\[AI Pitch Deck Vision Analysis\]", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    parts = re.split(r"(?<=[.!?])\s+", cleaned)
    kw = [k.lower() for k in (keywords or []) if k and len(k) > 2]
    ranked: list[tuple[int, str]] = []
    for part in parts:
        part = part.strip().strip('"')
        if len(part) < min_len:
            continue
        if part.startswith("(") and "failed" in part.lower():
            continue
        if part.startswith("[") and "unavailable" in part.lower():
            continue
        hit = sum(1 for k in kw if k in part.lower()) if kw else 0
        ranked.append((hit, part[:280]))
    ranked.sort(key=lambda x: (-x[0], -len(x[1])))
    seen: set[str] = set()
    snippets: list[str] = []
    for _, part in ranked:
        key = part[:60].lower()
        if key in seen:
            continue
        seen.add(key)
        snippets.append(part)
        if len(snippets) >= max_snippets:
            break
    return snippets


def _build_criterion_detail_user_message(
    submission: dict,
    criterion_name: str,
    cd: dict,
    rubric_crit: dict,
) -> str:
    """Rich LLM context for one criterion deep-dive."""
    s10 = cd.get("score_10", round(cd.get("score", 0) / 10, 1))
    band = _score_band_from_10(float(s10))
    anchors = rubric_crit.get("scoring_anchors", {})
    definition = (
        rubric_crit.get("definition")
        or rubric_crit.get("description")
        or ""
    )
    sub_factors = rubric_crit.get("sub_factors", [])
    rf_rubric = rubric_crit.get("red_flags", [])
    gates = rubric_crit.get("auto_reject_gates") or rubric_crit.get("gating") or ""

    lines = [
        f"CRITERION: {criterion_name}",
        f"ASSIGNED SCORE: {s10}/10",
        f"ASSIGNED BAND: {band}",
        f"WEIGHT: {cd.get('weight', '—')}%",
        f"EVIDENCE RATING: {cd.get('evidence', 'Partial')}",
        f"ONE-LINE JUSTIFICATION: {(cd.get('justification') or '').strip()}",
        f"TRIGGERED RED FLAGS: {', '.join(cd.get('red_flags', [])) or 'none'}",
        "",
        f"CRITERION DEFINITION: {definition[:500]}",
        f"SUB-FACTORS: {', '.join(sub_factors) if sub_factors else '—'}",
        f"RUBRIC RED FLAGS (watch list): {', '.join(rf_rubric) if rf_rubric else 'none'}",
    ]
    if gates:
        lines.append(f"GATING: {str(gates)[:220]}")
    lines.extend(["", "SCORING ANCHORS (use all three in your analysis):"])
    for anchor_band in ("1-3", "4-6", "7-10"):
        anchor_txt = anchors.get(anchor_band, "")
        marker = " ← ASSIGNED" if anchor_band == band else ""
        lines.append(f"  {anchor_band}{marker}: {anchor_txt[:320]}")

    lines.extend([
        "",
        f"IDEA: {submission.get('name', '')}",
        f"SUBMISSION ID: {submission.get('id', '')}",
    ])

    notes = submission.get("notes", "").strip()
    if notes:
        lines.extend(["", "SUBMISSION NOTES:", notes[:CRITERION_DETAIL_NOTES_LIMIT]])

    extracted = submission.get("extracted_text", "").strip()
    if extracted:
        lines.extend(["", "EXTRACTED FILE CONTENT:", extracted[:CRITERION_DETAIL_TEXT_LIMIT]])

    return "\n".join(lines)


def _anchor_alignment_paragraph(
    band: str,
    s10: float,
    anchor_txt: str,
    has_evidence: bool,
) -> str:
    """One paragraph per anchor band for structured fallback."""
    assigned = _score_band_from_10(s10)
    label = f"**{band}**"
    anchor_quote = f'*"{anchor_txt[:200]}{"…" if len(anchor_txt) > 200 else ""}"*' if anchor_txt else "*No anchor text supplied.*"

    if band == assigned:
        return (
            f"{label} {anchor_quote} — **This is the assigned band.** "
            f"The score of **{s10}/10** fits here because the available evidence "
            f"{'supports' if has_evidence else 'only partially supports'} the rubric language above. "
            f"A higher band would require stronger, criterion-specific proof; a lower band would imply "
            f"more severe gaps or red-flag triggers than observed."
        )
    if band == "7-10" and s10 < 7:
        return (
            f"{label} {anchor_quote} — The submission **does not yet meet** this band. "
            f"Missing proof points or unresolved risks prevent a 7+ score despite any partial strengths."
        )
    if band == "1-3" and s10 > 3:
        return (
            f"{label} {anchor_quote} — The submission **clears** this low band; "
            f"fundamental disqualifiers for this criterion were not the primary driver of the score."
        )
    if band == "4-6" and s10 > 6:
        return (
            f"{label} {anchor_quote} — The submission **exceeds** this mid band, "
            f"showing above-average alignment with the criterion relative to a typical 4–6 placement."
        )
    if band == "4-6" and s10 <= 3:
        return (
            f"{label} {anchor_quote} — The submission **falls below** this band, "
            f"indicating material weakness on this criterion."
        )
    return (
        f"{label} {anchor_quote} — Partial alignment only; the assigned score reflects "
        f"mixed signals rather than a clean fit for this anchor band."
    )


def _score_reasoning_paragraphs(
    criterion_name: str,
    name: str,
    s10: float,
    band: str,
    justification: str,
    cd: dict,
    snippets: list[str],
    sub_factors: list[str],
) -> str:
    """Multi-paragraph chain-of-thought for structured fallback."""
    ev = cd.get("evidence", "Partial")
    wt = cd.get("weight", "—")
    rf = cd.get("red_flags", [])

    p1 = (
        f"After reviewing the submission for **{name}**, ForgeOS assigned **{s10}/10** on "
        f"**{criterion_name}** (weight **{wt}%**). The screening justification states: "
        f"*\"{justification}\"*. This one-line summary reflects an evidence rating of **{ev}** "
        f"and places the idea in rubric anchor band **{band}**."
    )

    evidence_clause = (
        "Specific passages in the notes and extracted files speak directly to this criterion."
        if snippets
        else "However, little criterion-specific text was found in notes or extracted files, limiting confidence."
    )
    sf_clause = (
        f" Sub-factors evaluated include: {', '.join(sub_factors[:5])}."
        if sub_factors else ""
    )
    p2 = (
        f"{evidence_clause}{sf_clause} "
        f"Where the submission provides concrete detail, it supports the mid-to-upper range of the assigned band; "
        f"where detail is absent, the score is capped rather than upgraded."
    )

    if s10 >= 7:
        uplift = "To reach 9–10, the team would need unmistakable proof (metrics, validation, or third-party confirmation) tied to each sub-factor."
        lower = "A drop below 7 would follow from triggered red flags or materially insufficient evidence."
    elif s10 >= 4:
        uplift = "Raising this to 7+ would require closing documented gaps with verifiable traction, specs, or customer proof."
        lower = f"A score closer to 3 would imply severe misalignment or red-flag triggers{' — including: ' + ', '.join(rf) if rf else ''}."
    else:
        uplift = "Even moving to the 4–6 band would require addressing the core weaknesses noted above with new evidence."
        lower = "Further deterioration would follow if gating red flags materialize or due diligence reveals critical omissions."

    p3 = (
        f"In chain-of-thought terms: the submission mentions relevant themes for **{criterion_name}**, "
        f"but the combination of evidence quality (**{ev}**) and rubric anchor language for band **{band}** "
        f"is why the score lands at **{s10}** rather than a full point higher or lower. {uplift} {lower}"
    )

    return f"{p1}\n\n{p2}\n\n{p3}"


def _generate_criterion_detail_local(
    submission: dict,
    criterion_name: str,
    rubric_data: dict,
) -> str:
    """Structured deep-dive when LLM is unavailable or times out."""
    cd = submission.get("categories", {}).get(criterion_name, {})
    rubric_crit = _find_rubric_criterion(criterion_name, rubric_data)
    s10 = cd.get("score_10", round(cd.get("score", 0) / 10, 1))
    band = _score_band_from_10(float(s10))
    anchors = rubric_crit.get("scoring_anchors", {})

    combined = (submission.get("notes", "") + "\n" + submission.get("extracted_text", "")).strip()
    sub_factors = rubric_crit.get("sub_factors", [])
    keywords = [criterion_name] + list(sub_factors)
    snippets = _extract_evidence_snippets(combined, keywords=keywords, max_snippets=6)
    name = submission.get("name", "this submission")
    justification = (cd.get("justification") or "").strip()
    has_evidence = bool(snippets)

    summary = (
        f"*{justification or 'No one-line justification recorded.'}* "
        f"For pipeline review, **{s10}/10** on **{criterion_name}** (weight **{cd.get('weight', '—')}%**) "
        f"signals a **{band}**-band placement — "
        f"{'a candidate for advancement with validation' if float(s10) >= 7 else 'mixed signals requiring diligence' if float(s10) >= 4 else 'material concern before shortlist consideration'}."
    )

    rationale = _score_reasoning_paragraphs(
        criterion_name, name, float(s10), band, justification, cd, snippets, sub_factors
    )

    anchor_blocks = [
        _anchor_alignment_paragraph(b, float(s10), anchors.get(b, ""), has_evidence)
        for b in ("1-3", "4-6", "7-10")
    ]
    anchor_section = "\n\n".join(anchor_blocks)

    if snippets:
        quote_block = "\n".join(f"- > \"{s}\"" for s in snippets)
        evidence_note = ""
    else:
        quote_block = "- > *No direct quotes matched this criterion in notes or extracted files.*"
        evidence_note = (
            "\n\n*After reviewing available material, criterion-specific evidence is thin. "
            "Add submission notes or ensure pitch-deck text extraction succeeded for richer Expand views.*"
        )

    rf_rubric = rubric_crit.get("red_flags", [])
    rf_hit = cd.get("red_flags", [])
    pos_highlights: list[str] = []
    if float(s10) >= 7:
        pos_highlights.append(
            f"Score **{s10}/10** aligns with the **{band}** anchor — above typical sector baseline for this criterion."
        )
    if cd.get("evidence") == "Sufficient":
        pos_highlights.append("Evidence rated **Sufficient** — uploaded material supports rubric review.")
    if snippets:
        pos_highlights.append(
            f"**{len(snippets)}** quotable passage(s) in notes/files directly reference this criterion or its sub-factors."
        )
    if not rf_hit and float(s10) >= 5:
        pos_highlights.append("No criterion-specific red flags triggered in automated screening.")

    flag_lines: list[str] = []
    for rf in rf_hit:
        flag_lines.append(f"- **Triggered red flag:** {rf}")
    for rf in rf_rubric:
        if rf not in rf_hit and float(s10) <= 5:
            flag_lines.append(f"- **Watch item:** {rf} — not auto-flagged but score warrants scrutiny")
    if not flag_lines:
        flag_lines.append("- No active red flags for this criterion.")

    return f"""## Summary

{summary}

## Score Rationale

{rationale}

## Rubric Anchor Analysis

{anchor_section}

## Evidence from Submission

{quote_block}{evidence_note}

## Flags & Highlights

**Red flags & watch items:**
{chr(10).join(flag_lines)}

**Standout positives:**
{chr(10).join(f'- {h}' for h in pos_highlights) if pos_highlights else '- No standout positives at this score level.'}
"""


def _generate_criterion_detail_llm(
    submission: dict,
    criterion_name: str,
    rubric_data: dict,
    api_key: str,
    base_url: str,
    model: str,
) -> str:
    cd = submission.get("categories", {}).get(criterion_name, {})
    rubric_crit = _find_rubric_criterion(criterion_name, rubric_data)
    user_message = _build_criterion_detail_user_message(submission, criterion_name, cd, rubric_crit)
    return _llm_chat_text(
        _CRITERION_DETAIL_SYSTEM,
        user_message,
        api_key,
        base_url,
        model,
        temperature=0.35,
        max_tokens=1600,
        timeout=55,
    )


def generate_criterion_detail(submission: dict, criterion_name: str, rubric_data: dict):
    """
    Generate expanded criterion analysis. Returns (markdown, source_label, warning_or_none).
    Uses LLM when an API key is configured; otherwise structured local deep-dive.
    On timeout or LLM error, falls back to structured analysis with a clean user message.
    """
    if _llm_ready():
        try:
            md = _generate_criterion_detail_llm(
                submission,
                criterion_name,
                rubric_data,
                _effective_llm_key(),
                _FORGE_LLM_BASE_URL,
                _FORGE_LLM_MODEL,
            )
            return md, "llm", None
        except Exception as exc:
            md = _generate_criterion_detail_local(submission, criterion_name, rubric_data)
            return md, "structured", _criterion_detail_fallback_message(exc)
    md = _generate_criterion_detail_local(submission, criterion_name, rubric_data)
    return md, "structured", _criterion_detail_fallback_message(None)


def _ensure_criterion_detail_cached(sub_id: str, criterion_name: str, rubric_data: dict) -> None:
    """Generate and cache deep-dive when dialog opens or user clicks Regenerate."""
    cache_key = _criterion_detail_cache_key(sub_id, criterion_name)
    if cache_key in st.session_state.criterion_detail_cache and not st.session_state.criterion_detail_force:
        return
    sub = next((s for s in st.session_state.submissions if s["id"] == sub_id), None)
    if not sub:
        return
    md, source, warning = generate_criterion_detail(sub, criterion_name, rubric_data)
    st.session_state.criterion_detail_cache[cache_key] = {
        "markdown": md,
        "source": source,
        "warning": warning,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    st.session_state.criterion_detail_force = False


def _criterion_detail_dialog_body():
    """Modal body for expanded criterion analysis."""
    sub_id = st.session_state.get("criterion_deep_dive_sub_id")
    crit_name = st.session_state.get("criterion_deep_dive_criterion")
    panel_key = st.session_state.get("criterion_deep_dive_active")
    if not sub_id or not crit_name or not panel_key:
        return

    sub = next((s for s in st.session_state.submissions if s["id"] == sub_id), None)
    if not sub:
        st.error("Submission not found.")
        if st.button("Close", key=f"crit_detail_close_missing_{panel_key}"):
            _close_all_criterion_deep_dives()
            st.rerun()
        return

    cd = sub.get("categories", {}).get(crit_name, {})
    s10 = cd.get("score_10", round(cd.get("score", 0) / 10, 1))
    score_clr = score_hex(cd.get("score", 0))

    st.markdown(
        f'<div style="display:flex;align-items:center;justify-content:space-between;'
        f'margin-bottom:12px;">'
        f'<div><div style="font-size:18px;font-weight:700;color:var(--text-1);">'
        f'{_esc(crit_name)}</div>'
        f'<div style="font-size:12px;color:var(--text-3);margin-top:4px;">'
        f'{_esc(sub.get("name", ""))} · {_esc(sub_id)}</div></div>'
        f'<div style="font-size:22px;font-weight:800;color:{score_clr};">{s10}/10</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    hdr1, hdr2 = st.columns([1, 1])
    with hdr1:
        if st.button("↻ Regenerate", key=f"{panel_key}_regen", use_container_width=True):
            st.session_state.criterion_detail_force = True
            st.session_state.criterion_detail_cache.pop(
                _criterion_detail_cache_key(sub_id, crit_name), None
            )
            st.rerun()
    with hdr2:
        if st.button("Close", key=f"{panel_key}_close", use_container_width=True):
            _close_all_criterion_deep_dives()
            st.rerun()

    with st.spinner("Generating detailed analysis…"):
        _ensure_criterion_detail_cached(sub_id, crit_name, rubric)

    cached = st.session_state.criterion_detail_cache.get(
        _criterion_detail_cache_key(sub_id, crit_name)
    )
    if not cached:
        st.error("Could not generate analysis. Please try **Regenerate**.")
        return

    if cached.get("warning"):
        st.info(cached["warning"])

    src_lbl = "AI deep-dive" if cached.get("source") == "llm" else "ForgeOS structured analysis"
    st.caption(f"Source: {src_lbl} · Generated {cached.get('generated_at', '')}")
    st.markdown(cached["markdown"])


if hasattr(st, "dialog"):
    try:
        @st.dialog("Criterion Deep Dive", width="large")
        def _criterion_detail_dialog():
            _criterion_detail_dialog_body()
    except TypeError:
        @st.dialog("Criterion Deep Dive")
        def _criterion_detail_dialog():
            _criterion_detail_dialog_body()
else:
    def _criterion_detail_dialog():
        st.markdown("#### Criterion Deep Dive")
        _criterion_detail_dialog_body()


def _maybe_open_criterion_deep_dive_dialog():
    """Open criterion deep-dive modal when a criterion Expand is active."""
    if not st.session_state.get("criterion_deep_dive_active"):
        return
    if hasattr(st, "dialog"):
        _criterion_detail_dialog()
    else:
        with st.container(border=True):
            _criterion_detail_dialog_body()


# ─── Investment memo generation ──────────────────────────────────────────────

_MEMO_SYSTEM_PROMPT = """You are ForgeOS, a senior venture analyst and physical-goods innovation strategist.
Write a concise, investor-ready Executive Summary / Investment Memo in polished Markdown.

AUDIENCE: Seed/Series A investors, corporate innovation teams, and manufacturing partners.

TONE: Professional, direct, evidence-based. No hype or unsupported superlatives.

STRUCTURE (use these exact section headings as ## headers):
## Executive Summary
- Idea name as # title, then a compelling one-line tagline
- 2–3 paragraph overview of the opportunity

## Score Overview
- Overall ForgeOS score (/100) and investment posture (Strong / Promising / Needs Work / Auto-Reject)
- Top 3 strengths (highest-scoring rubric criteria with scores)
- Key risks and red flags (gating flags, low criteria, criterion red flags)

## Rubric Breakdown
- All 8 criteria: name, score /10, weight %, and 1–2 sentence justification each (use provided data)

## Pipeline & Stage Progress
- Current stage and status
- Brief synthesis of stage history and prior stage outputs (if provided)

## Evidence & Multimodal Highlights
- Bullet key points from submission notes and extracted file content
- Call out data gaps honestly

## Recommended Next Steps
- 4–6 numbered, actionable next steps for the innovation team
- Include regulatory, manufacturing, or go-to-market items where relevant

## Investment Considerations
- Short paragraph on fit, timing, and what would increase conviction

RULES:
- Output ONLY valid Markdown (no JSON, no code fences wrapping the whole doc).
- Ground every claim in the supplied submission data; do not invent patents, revenue, or customers.
- If data is missing, state what evidence is needed.
- Keep total length roughly 800–1,200 words."""


def _llm_chat_text(
    system_prompt: str,
    user_message: str,
    api_key: str,
    base_url: str,
    model: str,
    *,
    temperature: float = 0.35,
    max_tokens: int = 4500,
    timeout: int = 90,
) -> str:
    """OpenAI-compatible chat completion returning assistant text."""
    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }).encode("utf-8")

    req = urllib.request.Request(
        url=f"{base_url.rstrip('/')}/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        api_result = json.loads(resp.read().decode("utf-8"))
    content = api_result["choices"][0]["message"]["content"].strip()
    # Strip accidental markdown fences
    if content.startswith("```"):
        lines = content.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        content = "\n".join(lines).strip()
    return content


def _build_memo_context_payload(submission: dict, rubric_data: dict) -> str:
    """Assemble structured context for memo generation (LLM or local)."""
    cats = submission.get("categories") or {}
    lines = [
        f"SUBMISSION ID: {submission.get('id', '')}",
        f"IDEA NAME: {submission.get('name', '')}",
        f"STATUS: {submission.get('status', '')}",
        f"PIPELINE STAGE: {submission.get('stage', '')}",
        f"SUBMITTED: {submission.get('submitted_at', '')}",
        f"FILE TYPES: {submission.get('file_type', '')}",
        "",
        f"OVERALL SCORE: {submission.get('overall', 0)}/100",
        f"INNOVATION SCORE: {submission.get('innovation', 0)}/100",
        f"FEASIBILITY SCORE: {submission.get('feasibility', 0)}/100",
        f"SCORED AT: {submission.get('scored_at', 'N/A')}",
        "",
    ]

    if submission.get("auto_reject"):
        lines.append("AUTO-REJECT GATES:")
        lines.extend(f"  - {f}" for f in submission["auto_reject"])
        lines.append("")
    if submission.get("high_risk"):
        lines.append("HIGH-RISK FLAGS:")
        lines.extend(f"  - {f}" for f in submission["high_risk"])
        lines.append("")

    sl_cat = _shortlist_category_for(submission.get("id", ""))
    if sl_cat:
        lines.append(f"SHORTLIST FOLDER: {sl_cat}")
        lines.append("")

    notes = submission.get("notes", "").strip()
    if notes:
        lines.append("SUBMISSION NOTES:")
        lines.append(notes[:8000])
        lines.append("")

    extracted = submission.get("extracted_text", "").strip()
    if extracted:
        lines.append("EXTRACTED FILE CONTENT (multimodal):")
        lines.append(extracted[:EXTRACTED_TEXT_LLM_LIMIT])
        lines.append("")

    for fs in submission.get("file_summaries") or []:
        lines.append(
            f"FILE: {fs.get('name', '?')} | type={fs.get('file_type')} | "
            f"method={fs.get('extraction_method')} | chars={fs.get('chars', 0)}"
        )
        if fs.get("preview"):
            lines.append(f"  Preview: {str(fs['preview'])[:500]}")
    lines.append("")

    lines.append("RUBRIC CRITERIA SCORES:")
    for crit in rubric_data.get("criteria", []):
        key = crit["criterion"]
        cd = cats.get(key, {})
        s10 = cd.get("score_10", "—")
        wt = crit.get("weight", 0)
        just = cd.get("justification", "No justification recorded.")
        rf = cd.get("red_flags") or []
        lines.append(f"- {key} (weight {wt}%): {s10}/10")
        lines.append(f"  Justification: {just}")
        if rf:
            lines.append(f"  Red flags: {', '.join(rf)}")
    lines.append("")

    hist = submission.get("stage_history") or []
    if hist:
        lines.append("STAGE HISTORY:")
        for e in hist:
            lines.append(f"  - {e.get('stage')} ({e.get('moved_at', '')})")
        lines.append("")

    summ = submission.get("stage_summary", "")
    if summ:
        # Strip HTML tags for LLM context
        plain = re.sub(r"<[^>]+>", " ", summ)
        plain = re.sub(r"\s+", " ", plain).strip()
        lines.append("CURRENT STAGE AI SUMMARY (plain text):")
        lines.append(plain[:4000])
        lines.append("")

    lines.append(f"RUBRIC NAME: {rubric_data.get('rubric_name', 'ForgeOS Rubric')}")
    return "\n".join(lines)


def _generate_investment_memo_local(submission: dict, rubric_data: dict) -> str:
    """Deterministic investor memo from scored data (no LLM)."""
    name = submission.get("name", "Untitled Idea")
    overall = submission.get("overall", 0)
    cats = submission.get("categories") or {}
    blabel, _, _ = forge_badge(submission)

    ranked = sorted(
        cats.items(),
        key=lambda x: x[1].get("score_10", 0),
        reverse=True,
    )
    strengths = ranked[:3]
    weaknesses = sorted(cats.items(), key=lambda x: x[1].get("score_10", 0))[:3]

    tagline = (submission.get("notes") or "").strip().split("\n")[0][:160]
    if not tagline:
        tagline = f"Physical-goods innovation evaluated at {overall}/100 on ForgeOS Extensive Rubric v2."

    md = [f"# {name}", "", f"*{tagline}*", "", "---", ""]
    md.append("## Executive Summary")
    md.append("")
    md.append(
        f"**{name}** is a physical-goods innovation currently in the **{submission.get('stage', 'Intake')}** "
        f"pipeline stage with ForgeOS overall score **{overall}/100** ({blabel or 'Unrated'}). "
    )
    notes = submission.get("notes", "").strip()
    if notes:
        md.append(notes[:600] + ("…" if len(notes) > 600 else ""))
    else:
        md.append("Detailed founder notes were not provided; assessment relies on scoring and file extracts.")
    md.append("")

    md.append("## Score Overview")
    md.append("")
    md.append(f"| Metric | Value |")
    md.append(f"|--------|-------|")
    md.append(f"| Overall | **{overall}/100** |")
    md.append(f"| Innovation | {submission.get('innovation', 0)}/100 |")
    md.append(f"| Feasibility | {submission.get('feasibility', 0)}/100 |")
    md.append(f"| Posture | {blabel or '—'} |")
    md.append("")

    md.append("### Top strengths")
    for k, cd in strengths:
        md.append(f"- **{k}** — {cd.get('score_10', '—')}/10: {cd.get('justification', '')[:200]}")
    md.append("")

    md.append("### Risks & red flags")
    if submission.get("auto_reject"):
        for f in submission["auto_reject"]:
            md.append(f"- ⛔ {f}")
    if submission.get("high_risk"):
        for f in submission["high_risk"]:
            md.append(f"- ⚠ {f}")
    for k, cd in weaknesses:
        for rf in cd.get("red_flags") or []:
            md.append(f"- **{k}**: {rf}")
    if not submission.get("auto_reject") and not submission.get("high_risk"):
        for k, cd in weaknesses:
            md.append(f"- **{k}** — {cd.get('score_10', '—')}/10 (relative weakness)")
    md.append("")

    md.append("## Rubric Breakdown")
    md.append("")
    for crit in rubric_data.get("criteria", []):
        key = crit["criterion"]
        cd = cats.get(key, {})
        md.append(f"### {key} ({crit.get('weight', 0)}%)")
        md.append(f"**Score:** {cd.get('score_10', '—')}/10 · Evidence: {cd.get('evidence', '—')}")
        md.append("")
        md.append(cd.get("justification", "—"))
        md.append("")

    md.append("## Pipeline & Stage Progress")
    md.append("")
    md.append(f"- **Current stage:** {submission.get('stage', '—')}")
    md.append(f"- **Status:** {submission.get('status', '—')}")
    for e in submission.get("stage_history") or []:
        md.append(f"- {e.get('stage')} — {e.get('moved_at', '')}")
    md.append("")

    md.append("## Evidence & Multimodal Highlights")
    md.append("")
    for fs in submission.get("file_summaries") or []:
        md.append(f"- **{fs.get('name', 'file')}** ({fs.get('file_type', 'other')}): "
                  f"{fs.get('chars', 0):,} chars via {fs.get('extraction_method', '—')}")
    ext = submission.get("extracted_text", "").strip()
    if ext:
        md.append("")
        md.append("```")
        md.append(ext[:1500] + ("…" if len(ext) > 1500 else ""))
        md.append("```")
    else:
        md.append("- No extracted file content on record.")
    md.append("")

    md.append("## Recommended Next Steps")
    md.append("")
    stage = submission.get("stage", "Intake")
    steps = [
        "Validate top weakness criterion with primary data (customer interviews or lab test).",
        "Confirm manufacturing path and BOM at target MOQ with at least one supplier quote.",
        "Close regulatory/compliance gaps flagged in rubric before channel expansion.",
    ]
    if stage in ("Intake", "Concept"):
        steps.insert(0, "Complete technical feasibility prototype and attach results to submission.")
    if submission.get("auto_reject"):
        steps.insert(0, "Address auto-reject gates before re-submitting for investment review.")
    for i, s in enumerate(steps[:6], 1):
        md.append(f"{i}. {s}")
    md.append("")

    md.append("## Investment Considerations")
    md.append("")
    if overall >= THRESHOLDS["green"] and not submission.get("auto_reject"):
        md.append(
            "Score profile supports continued diligence. Prioritize evidence that de-risks "
            "manufacturing scale-up and confirms differentiated IP or process."
        )
    elif submission.get("auto_reject"):
        md.append(
            "Auto-reject gates indicate the concept does not yet meet minimum innovation, "
            "feasibility, or evidence thresholds for capital deployment."
        )
    else:
        md.append(
            "Promising elements exist but conviction requires stronger evidence on weaker "
            "criteria and clearer unit economics at volume."
        )
    md.append("")
    md.append(f"*Generated by ForgeOS · {datetime.now().strftime('%Y-%m-%d %H:%M')} · Structured export*")

    return "\n".join(md)


def _generate_investment_memo_llm(
    submission: dict,
    rubric_data: dict,
    api_key: str,
    base_url: str,
    model: str,
) -> str:
    """LLM-authored investment memo (Markdown)."""
    context = _build_memo_context_payload(submission, rubric_data)
    user_msg = (
        "Generate the investment memo for this ForgeOS submission. "
        "Use only the data below.\n\n"
        f"{context}"
    )
    return _llm_chat_text(
        _MEMO_SYSTEM_PROMPT,
        user_msg,
        api_key,
        base_url,
        model,
        temperature=0.35,
        max_tokens=4500,
        timeout=90,
    )


def generate_investment_memo(submission: dict, rubric_data: dict):
    """
    Generate memo markdown. Returns (markdown, source_label, warning_or_none).
    source_label is 'llm' or 'structured'.
    """
    if _llm_ready():
        try:
            md = _generate_investment_memo_llm(
                submission,
                rubric_data,
                _effective_llm_key(),
                _FORGE_LLM_BASE_URL,
                _FORGE_LLM_MODEL,
            )
            return md, "llm", None
        except Exception as e:
            md = _generate_investment_memo_local(submission, rubric_data)
            return md, "structured", f"LLM unavailable ({e}) — structured memo generated from score data."
    md = _generate_investment_memo_local(submission, rubric_data)
    return md, "structured", (
        "No API key — structured memo from ForgeOS scores. "
        "Add a key in Rubric Settings for AI-authored memos."
    )


def _memo_to_pdf_bytes(markdown_text: str) -> bytes:
    """Render memo as a simple multi-page PDF via PyMuPDF."""
    import fitz

    doc = fitz.open()
    page_w, page_h = 595, 842  # A4 pt
    margin, fontsize, leading = 50, 10, 14
    max_chars = 95

    def _wrap_line(line: str) -> list[str]:
        line = line.replace("\t", "    ")
        if len(line) <= max_chars:
            return [line] if line else [""]
        words, out, cur = line.split(), [], ""
        for w in words:
            if len(cur) + len(w) + 1 <= max_chars:
                cur = f"{cur} {w}".strip()
            else:
                if cur:
                    out.append(cur)
                cur = w
        if cur:
            out.append(cur)
        return out or [""]

    y = margin
    page = doc.new_page(width=page_w, height=page_h)

    for raw_line in markdown_text.split("\n"):
        for line in _wrap_line(raw_line):
            if y > page_h - margin:
                page = doc.new_page(width=page_w, height=page_h)
                y = margin
            page.insert_text((margin, y), line, fontsize=fontsize, fontname="helv")
            y += leading

    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _ensure_memo_cached(sub_id: str, rubric_data: dict) -> None:
    """Generate and cache memo when memo_needs_generate or memo_force_regenerate is set."""
    if not (
        st.session_state.get("memo_needs_generate")
        or st.session_state.get("memo_force_regenerate")
    ):
        return

    sub = next((s for s in st.session_state.submissions if s["id"] == sub_id), None)
    if not sub or not sub.get("categories"):
        st.session_state.memo_needs_generate = False
        st.session_state.memo_force_regenerate = False
        return

    with st.spinner("Generating investment memo…"):
        md, source, _warn = generate_investment_memo(sub, rubric_data)
        st.session_state.investment_memos[sub_id] = {
            "markdown": md,
            "source": source,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "warning": _warn,
        }
    st.session_state.memo_needs_generate = False
    st.session_state.memo_force_regenerate = False


def _investment_memo_dialog_body():
    """Modal body: memo preview, copy, export, regenerate."""
    sub_id = st.session_state.get("memo_sub_id")
    if not sub_id:
        return

    sub = next((s for s in st.session_state.submissions if s["id"] == sub_id), None)
    if not sub:
        st.warning("Submission not found.")
        if st.button("Close", use_container_width=True):
            st.session_state.memo_sub_id = None
            st.rerun()
        return

    st.markdown(
        f'<div style="font-size:15px;font-weight:600;color:#e6edf3;">{_esc(sub["name"])}</div>'
        f'<div style="font-size:11px;color:#8b949e;margin-bottom:8px;">{_esc(sub_id)} · '
        f'Investment Memo</div>',
        unsafe_allow_html=True,
    )

    _ensure_memo_cached(sub_id, rubric)
    cached = st.session_state.investment_memos.get(sub_id)

    if not cached or not cached.get("markdown"):
        st.error("Could not generate memo — ensure the submission is scored.")
        if st.button("Close", use_container_width=True):
            st.session_state.memo_sub_id = None
            st.rerun()
        return

    if cached.get("warning"):
        st.info(cached["warning"])

    source_lbl = "AI-authored (LLM)" if cached.get("source") == "llm" else "Structured export"
    st.markdown(
        f'<div class="memo-source-tag">{source_lbl} · {cached.get("generated_at", "")}</div>',
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        st.markdown(cached["markdown"])

    # Raw markdown for copy / download (also shown for manual copy fallback)
    with st.expander("View raw Markdown", expanded=False):
        st.text_area(
            "memo_raw",
            value=cached["markdown"],
            height=200,
            label_visibility="collapsed",
        )

    safe_name = re.sub(r"[^\w\-]+", "_", sub.get("name", "memo"))[:40]
    fname_base = f"ForgeOS_Memo_{safe_name}_{sub_id}"

    btn1, btn2, btn3, btn4, btn5 = st.columns(5)
    with btn1:
        if hasattr(st, "copy_button"):
            st.copy_button(
                "Copy Memo",
                cached["markdown"],
                use_container_width=True,
                help="Copy full memo to clipboard",
            )
        else:
            st.download_button(
                "Copy as .md",
                cached["markdown"],
                file_name=f"{fname_base}.md",
                mime="text/markdown",
                use_container_width=True,
            )
    with btn2:
        st.download_button(
            "Export Markdown",
            cached["markdown"],
            file_name=f"{fname_base}.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with btn3:
        try:
            pdf_bytes = _memo_to_pdf_bytes(cached["markdown"])
            st.download_button(
                "Export PDF",
                pdf_bytes,
                file_name=f"{fname_base}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as ex:
            st.button("Export PDF", disabled=True, help=str(ex), use_container_width=True)
    with btn4:
        if st.button("Regenerate", use_container_width=True):
            st.session_state.memo_force_regenerate = True
            st.rerun()
    with btn5:
        if st.button("Close", use_container_width=True):
            st.session_state.memo_sub_id = None
            st.rerun()


if hasattr(st, "dialog"):
    try:
        @st.dialog("Investment Memo", width="large")
        def _investment_memo_dialog():
            _investment_memo_dialog_body()
    except TypeError:
        @st.dialog("Investment Memo")
        def _investment_memo_dialog():
            _investment_memo_dialog_body()
else:
    def _investment_memo_dialog():
        st.markdown("#### Investment Memo")
        _investment_memo_dialog_body()


def _maybe_open_memo_dialog():
    """Open investment memo modal when memo_sub_id is set."""
    if not st.session_state.get("memo_sub_id"):
        return
    if hasattr(st, "dialog"):
        _investment_memo_dialog()
    else:
        with st.container(border=True):
            _investment_memo_dialog_body()


def generate_stage_summary(submission, new_stage):
    """
    Generate a rich, physical-goods-specific HTML stage summary.
    Product type is inferred from the submission name; all content is
    pre-computed so no backslash escapes appear inside f-string expressions.
    """
    name  = submission.get("name", "this idea")
    nl    = name.lower()

    # ── Product-type detection ─────────────────────────────────────────────────
    is_bio     = any(w in nl for w in ["bio", "mycelium", "compost", "organic", "plant", "hemp"])
    is_carbon  = any(w in nl for w in ["carbon", "fibre", "fiber", "graphene"])
    is_polymer = any(w in nl for w in ["polymer", "coating", "composite", "resin", "foam", "panel"])
    is_mech    = any(w in nl for w in ["motor", "drive", "exoskeleton", "frame", "actuator", "gear"])
    is_thermal = any(w in nl for w in ["thermal", "insulation", "heat", "cooling", "regulator"])
    is_pack    = any(w in nl for w in ["packaging", "pack", "wrap", "container", "pouch"])
    is_wear    = any(w in nl for w in ["sleeve", "compression", "garment", "textile", "wearable"])
    is_smart   = any(w in nl for w in ["smart", "sensor", "iot", "connected", "digital", "electronic"])

    # ── Material / commercial context ─────────────────────────────────────────
    if is_bio:
        mat1     = "Mycelium composite (Ecovative-spec substrate)"
        mat2     = "Recycled PLA structural reinforcement"
        mfg      = "Compression moulding with biodegradable release agents"
        bom      = "$4–$9 per unit at 5K MOQ"
        retail   = "$18–$35 DTC / $10–$18 wholesale"
        r_tech   = "Medium"
        r_reg    = "Low (food-contact cert. if applicable)"
        r_mkt    = "Medium — consumer education required on novel materials"
        kpis     = ["Biodegradation rate (days)", "Compressive strength (kPa)", "Customer NPS"]
        sup_list = ["Ecovative Design", "Mogu S.r.l.", "BioFab"]
    elif is_carbon or is_polymer:
        mat1     = "High-performance polymer alloy (PEEK/PA66 blend)"
        mat2     = "UD carbon-fibre prepreg (T300 grade)"
        mfg      = "Autoclave cure + CNC post-processing for tight tolerances"
        bom      = "$18–$38 per unit at 2K MOQ"
        retail   = "$75–$180 B2B / OEM channel"
        r_tech   = "High — thermal cycling and UV degradation testing required"
        r_reg    = "Medium — REACH compliance, RoHS if electronic"
        r_mkt    = "Low — industrial buyers are specification-driven"
        kpis     = ["Tensile strength vs spec", "Surface defect rate (%)", "B2B reorder rate"]
        sup_list = ["Toray Industries", "Solvay", "Hexcel"]
    elif is_mech:
        mat1     = "Aerospace-grade Al 6061-T6 billet"
        mat2     = "Carbon-fibre composite structural panels"
        mfg      = "CNC machining + die-casting; anodised finish"
        bom      = "$25–$55 per unit at 1K MOQ"
        retail   = "$120–$280 direct + distributor"
        r_tech   = "High — dynamic load, fatigue, and drop-test certification required"
        r_reg    = "Medium — CE marking / UL for mechatronic variants"
        r_mkt    = "Medium — OEM partnerships accelerate channel velocity"
        kpis     = ["Mean cycles to failure", "Assembly defect rate (%)", "On-time delivery %"]
        sup_list = ["Protocase", "Xometry", "Fictiv"]
    elif is_thermal:
        mat1     = "Aerogel-reinforced silica blanket (10mm)"
        mat2     = "Vacuum-insulated panel (VIP) core option"
        mfg      = "Lamination + die-cutting with sealed edge processing"
        bom      = "$9–$22 per unit at 3K MOQ"
        retail   = "$40–$95 DTC; $22–$50 wholesale"
        r_tech   = "Medium — moisture ingress and thermal cycling required"
        r_reg    = "Low — confirm UL 94 flammability rating"
        r_mkt    = "Medium — technical sell requires performance data sheets"
        kpis     = ["Thermal resistance (R-value)", "Moisture ingress rate", "Field failure rate"]
        sup_list = ["Cabot Corporation", "Evonik", "Kingspan Group"]
    elif is_pack:
        mat1     = "Cellulose-based moulded fibre (recycled content)"
        mat2     = "PHA bioplastic barrier coating"
        mfg      = "Wet-press moulded fibre + inline PHA spray coating"
        bom      = "$0.45–$1.40 per unit at 50K MOQ"
        retail   = "$2.20–$4.80 wholesale to brand customers"
        r_tech   = "Low — key risk is moisture-barrier performance spec"
        r_reg    = "Low — confirm EN 13432 / ASTM D6400 compostability"
        r_mkt    = "Low — strong regulatory tailwinds (SUP Directive, plastic bans)"
        kpis     = ["WVTR (g/m²/day)", "Unit cost vs virgin plastic", "Brand customer LTV"]
        sup_list = ["Huhtamaki", "Sealed Air", "Smurfit Kappa"]
    elif is_wear:
        mat1     = "Medical-grade 70D nylon/elastane knit (CE Class I)"
        mat2     = "Silicone grip bead strips"
        mfg      = "Cut-and-sew flatlock seaming; ultrasonic bonding option"
        bom      = "$6–$14 per unit at 5K MOQ"
        retail   = "$39–$89 DTC; $22–$45 wholesale"
        r_tech   = "Low — main risk is sizing accuracy across segments"
        r_reg    = "Medium — CE Class I MDR 2017/745 if medical claims made"
        r_mkt    = "Medium — DTC channel requires clinical claims or influencer evidence"
        kpis     = ["Return rate by size band", "NPS", "Reorder rate at 90 days"]
        sup_list = ["Sanfori", "Amann Group", "Coats plc"]
    elif is_smart:
        mat1     = "PCB assembly (4-layer, 35μm copper, FR4 substrate)"
        mat2     = "Polycarbonate IP65 enclosure"
        mfg      = "SMT PCBA + injection-moulded housing + firmware flash and test"
        bom      = "$20–$48 per unit at 2K MOQ"
        retail   = "$89–$199 DTC; $45–$95 OEM"
        r_tech   = "High — EMC (CE/FCC), battery safety (UN38.3) required"
        r_reg    = "High — CE mandatory; FCC ID for US; UKCA post-Brexit"
        r_mkt    = "Medium — SaaS/platform revenue model improves LTV significantly"
        kpis     = ["Device uptime (%)", "MTBF (hours)", "MRR growth"]
        sup_list = ["JLCPCB", "PCBWay", "Foxconn"]
    else:
        mat1     = "Primary material TBD — pending Validation review"
        mat2     = "Secondary reinforcement TBD"
        mfg      = "Manufacturing process to be determined post feasibility review"
        bom      = "TBD — cost modelling scheduled for Validation"
        retail   = "TBD"
        r_tech   = "TBD"
        r_reg    = "TBD"
        r_mkt    = "TBD"
        kpis     = ["Core product KPI", "Quality metric", "Commercial metric"]
        sup_list = ["Suppliers to be sourced at Validation"]

    sup1 = sup_list[0]
    sup2 = sup_list[1] if len(sup_list) > 1 else "TBD"
    kpi1 = kpis[0]
    kpi2 = kpis[1] if len(kpis) > 1 else "Quality metric"
    kpi3 = kpis[2] if len(kpis) > 2 else "Commercial metric"
    mat1_short = mat1.split("(")[0].strip()
    mat2_short = mat2.split("(")[0].strip()

    seed = int(hashlib.md5((name + new_stage).encode()).hexdigest()[:8], 16)
    rng  = random.Random(seed)

    # ── HTML building helpers (inline styles only) ────────────────────────────
    def bul(items, color="#58a6ff"):
        parts = []
        for item in items:
            parts.append(
                f'<div style="font-size:11px;color:#b0b8c4;line-height:1.75;">'
                f'<span style="color:{color};">▸</span> {item}</div>'
            )
        return "".join(parts)

    def sec(title, body):
        return (
            f'<div style="margin-bottom:10px;">'
            f'<div style="font-size:9px;text-transform:uppercase;letter-spacing:0.1em;'
            f'color:#8b949e;font-weight:700;margin-bottom:5px;">{title}</div>'
            f'{body}'
            f'</div>'
        )

    def kv(label, val):
        return (
            f'<div style="font-size:11px;color:#b0b8c4;margin-bottom:3px;">'
            f'<span style="color:#8b949e;">{label}:</span> <b style="color:#e6edf3;">{val}</b></div>'
        )

    footer_style = (
        'style="font-size:10px;color:#6e7681;border-top:1px solid #21262d44;'
        'padding-top:6px;margin-top:6px;"'
    )

    # ── Per-stage templates (2 variants, seeded pick) ─────────────────────────
    CONCEPT = [
        sec("Refined Product Specification v0.1",
            kv("Concept", name)
            + kv("Primary material", mat1)
            + kv("Secondary", mat2)
            + kv("Manufacturing process", mfg)
            + kv("Unit cost target", bom))
        + sec("Key Design Decisions", bul([
            "Lock in differentiated feature set before IP filing window closes",
            "Confirm material supply chain depth with 2+ qualified vendors",
            f"Commission DFM review from {sup1} or equivalent contract manufacturer",
            "Benchmark vs top 3 competitors: pricing, performance, sustainability",
        ]))
        + f'<div {footer_style}>Next gate: Concept review + rubric re-score in 14 days</div>',

        sec("Material & Process Rationale",
            kv("Primary", mat1)
            + kv("Secondary", mat2)
            + kv("Process", mfg)
            + kv("BOM target", bom))
        + sec("Initial Design Direction", bul([
            f"Form factor: compact, production-ready, optimised for {mfg.split('+')[0].strip().lower()}",
            "Finish: premium industrial aesthetic — anodised / matte / textured",
            "Key interaction: single-touch / tool-free / zero-waste assembly principle",
            "IP note: provisional patent application recommended within 30 days",
        ], color="#3fb950"))
        + f'<div {footer_style}>Next gate: Concept review + rubric re-score in 14 days</div>',
    ]

    VALIDATION = [
        sec("Risk Analysis",
            kv("Technical risk", r_tech)
            + kv("Regulatory risk", r_reg)
            + kv("Market risk", r_mkt))
        + sec("Feasibility Summary", bul([
            f"Manufacturing process: {mfg}",
            f"Indicative BOM: {bom}",
            f"Target retail: {retail}",
            "First-article lead time: 8–12 weeks from tooling sign-off",
            f"Primary supplier candidate: {sup1}",
        ]))
        + sec("Validation Sprint — 4 Weeks", bul([
            f"Week 1–2: BOM costing with {sup1} and {sup2} (RFQs issued)",
            "Week 2–3: Customer discovery — 10 B2B decision-maker interviews",
            "Week 3–4: Regulatory pathway mapping + compliance counsel brief",
        ], color="#d29922"))
        + f'<div {footer_style}>Gate criteria: BOM within 15% of model · 8+ positive customer interviews</div>',

        sec("Technical Feasibility Report", bul([
            f"Core material: {mat1}",
            f"Process: {mfg}",
            f"Technical risk: {r_tech}",
            f"First-article cost estimate commissioned from {sup1}",
        ]))
        + sec("Market Validation Plan", bul([
            "Target: 10 structured interviews with procurement / R&D decision-makers",
            f"Pricing hypothesis: {retail}",
            "Key probe: willingness to pay vs incumbent solution price point",
            f"Regulatory path: {r_reg}",
        ]))
        + f'<div {footer_style}>Gate criteria: BOM within 15% of model · 8+ positive customer interviews</div>',
    ]

    PROTOTYPING = [
        sec("Bill of Materials — Draft v0.1", bul([
            f"[01] {mat1_short} — bulk pricing from {sup1}",
            f"[02] {mat2_short} — RFQ issued to {sup2}",
            "[03] Tooling and fixturing — amortised over 5K+ units",
            "[04] Assembly labour — CMO rate TBD (target: below 20% of BOM)",
            f"[05] Packaging — included in BOM target of {bom}",
            "[06] QC testing per unit — est. $0.50–$1.50",
        ]))
        + sec("Prototype Milestones", bul([
            "Week 1–2: CAD finalisation + DFM sign-off with CMO",
            "Week 3–4: Tooling order placed; first-article by Week 6",
            "Week 5–6: Functional prototype build + internal performance test",
            "Week 7–8: User testing cohort (n=12) + design iteration",
            "Week 9: Go/no-go gate review with updated rubric score",
        ], color="#3fb950"))
        + f'<div {footer_style}>Manufacturing partner shortlisted: {sup1}</div>',

        sec("Prototyping Plan", bul([
            f"Process: {mfg}",
            f"CMO shortlist: {sup1} (primary), {sup2} (backup)",
            "Target: 3 functional prototypes + 5 for user testing cohort",
            f"BOM target: {bom}",
            "Critical path: tooling lead time typically 6–10 weeks",
        ]))
        + sec("Testing Protocol", bul([
            "Performance: validate vs rubric scoring anchor specifications",
            "Durability: 500-cycle accelerated life test (IEC 60068 reference)",
            f"Regulatory pre-screen: {r_reg}",
            "User sessions: structured 45-min tests with 12 target customers",
        ], color="#d29922"))
        + f'<div {footer_style}>Gate criteria: prototype meets spec · user NPS at or above 45</div>',
    ]

    MARKET_TEST = [
        sec("Go-to-Market Outline", bul([
            "Channel 1: DTC e-commerce (Shopify) — hero SKU only for focus",
            "Channel 2: 1 regional distributor for B2B / trade pilot",
            "Launch geography: UK + DACH (EU regulatory alignment)",
            "Launch cohort: 50 beta customers from validated discovery pool",
        ]))
        + sec("Pricing Strategy", bul([
            f"Anchor price: {retail}",
            "Approach: value-based — benchmark vs incumbent at 1.3x premium",
            "A/B test: premium vs value positioning in paid acquisition",
            f"Gross margin at {bom}: target 55–65% DTC / 40–50% wholesale",
        ], color="#d29922"))
        + sec("Target Customer Profile", bul([
            "Primary: procurement leads at mid-market manufacturers (50–500 employees)",
            "Secondary: DTC early adopters in sustainability / performance segment",
            "Core pain point: current solutions fail on performance, cost, or sustainability",
        ]))
        + f'<div {footer_style}>Gate: 8-week sell-through at or above 60% · NPS at or above 45 · CAC within model</div>',

        sec("Launch Playbook", bul([
            f"Hero SKU: {name} — single variant for market focus",
            f"Price point: {retail} — tested and validated with beta cohort",
            "Acquisition: paid search + LinkedIn (B2B) + partner referral programme",
            "Week 1–2: soft launch to beta list; Week 3+: paid channels fully live",
        ]))
        + sec("Commercial Assumptions", bul([
            f"BOM at launch: {bom}",
            "Gross margin target: 55–65% DTC / 40–50% wholesale",
            "CAC budget: 15–20% of first-year estimated LTV",
            "Payback period target: 9 months or less",
        ], color="#3fb950"))
        + f'<div {footer_style}>Gate: 8-week sell-through at or above 60% · NPS at or above 45 · CAC within model</div>',
    ]

    SCALING = [
        sec("Supply Chain Plan", bul([
            f"Tier 1 CMO: {sup1} — primary, 70% of volume commitment",
            f"Tier 2 backup: {sup2} — 30% or overflow (dual-source strategy active)",
            "Safety stock: 8 weeks forward cover at rolling forecast volume",
            "Incoterms: DAP — landed cost fully included in margin model",
            "Sea freight lead time: 10–14 weeks farm-to-warehouse",
        ]))
        + sec("Volume Cost Model", bul([
            f"Unit BOM at 10K MOQ: {bom}",
            f"Target retail: {retail}",
            "Logistics and duty: estimated 8–12% of landed cost",
            "Gross margin at scale: 58–68% DTC / 42–52% wholesale",
            "Breakeven volume: modelled at 3.5K units per month",
        ], color="#3fb950"))
        + f'<div {footer_style}>QMS: ISO 9001 scope extension filed; 3 inline QC checkpoints active</div>',

        sec("Manufacturing Scale-Up", bul([
            f"Primary CMO: {sup1} — NDA and quality agreement signed",
            f"Process: {mfg}",
            "Production ramp: 1K to 5K to 15K units over 3 quarters",
            "Tooling investment amortised over 12 months at forecast volume",
        ]))
        + sec("Operational Priorities", bul([
            "QMS: 8 inline inspection points defined and staffed",
            "Compliance: all certifications confirmed before channel expansion",
            "Inventory: demand-driven replenishment model activated",
            "Channel: 3 new distributors contracted for Q2 rollout",
        ], color="#58a6ff"))
        + f'<div {footer_style}>QMS: ISO 9001 scope extension filed; 3 inline QC checkpoints active</div>',
    ]

    MONITORING = [
        sec("KPI Dashboard — Live Tracking", bul([
            f"Product: {kpi1} — weekly vs baseline target",
            f"Quality: {kpi2} — automated alert if threshold breached",
            f"Commercial: {kpi3} — cohort analysis run monthly",
            "Return rate: target below 2.5%; triggers engineering review above 4%",
            "NPS: monthly pulse survey; target score at or above 50",
            "Gross margin: weekly P&L reviewed vs model",
        ]))
        + sec("90-Day Review Framework", bul([
            "Day 30: first commercial report — sell-through, returns, NPS baseline",
            "Day 60: product council review — improvement backlog prioritised",
            "Day 90: full P&L vs model; go/no-go for next-gen variant scoping",
        ], color="#3fb950"))
        + f'<div {footer_style}>V2 concept exploration scoped for month 4; innovation team briefed</div>',

        sec("Post-Launch Improvement Roadmap", bul([
            "Sprint 1 (month 1): resolve top 3 issues from beta customer feedback",
            "Sprint 2 (month 2): packaging cost optimisation — target 8% reduction",
            f"Sprint 3 (month 3): premium variant scoping for {name} V2",
            "Month 4: innovation brief to pipeline team for next-gen concept",
        ]))
        + sec("Live Metrics", bul([
            f"Tracking: {kpi1}, {kpi2}, {kpi3}",
            "Alert thresholds set in ops dashboard",
            "Weekly ops standup: return rate, NPS, gross margin reviewed",
            "Customer feedback loop: support tickets + quarterly user panel",
        ], color="#58a6ff"))
        + f'<div {footer_style}>V2 concept exploration scoped for month 4; innovation team briefed</div>',
    ]

    INTAKE_FALLBACK = [
        f'<div style="font-size:11px;color:#b0b8c4;">'
        f'{name} returned to Intake for re-assessment. '
        f'Assign a stage owner and complete the rubric scoring checklist before re-advancing.</div>'
    ]

    template_map = {
        "Concept":     CONCEPT,
        "Validation":  VALIDATION,
        "Prototyping": PROTOTYPING,
        "Market Test": MARKET_TEST,
        "Scaling":     SCALING,
        "Monitoring":  MONITORING,
    }

    options = template_map.get(new_stage, INTAKE_FALLBACK)
    return rng.choice(options)

DEMO_SUBMISSIONS = [
    # ── High-potential innovations ────────────────────────────────────────────
    (
        "Regenera™ Mycelium-Leather Work Boot",
        "Regenera Materials (Oakland, CA) is commercializing a breakthrough mycelium-composite "
        "upper material that matches full-grain leather tensile strength while cutting embodied "
        "carbon 78% versus bovine leather (third-party LCA cited, ISO 14040). Two provisional "
        "patents cover the cross-linking process. We have a validated pilot with 1,200 pairs "
        "manufactured via contract partner in Portugal; BOM at 10k units/month is $41/pair with "
        "52% gross margin at $129 MSRP. Founder team includes a former Allbirds materials lead "
        "and a supply chain engineer who built Nike's EU factory network. Letters of intent from "
        "two regional safety-footwear distributors ($2.1M annual run-rate). GRS and OEKO-TEX "
        "certifications in progress; REACH compliance documented.",
        "PDF", "Scaling", "Approved",
    ),
    (
        "NovaSeal™ Compostable High-Barrier Food Pouch",
        "NovaSeal replaces multi-layer PE/EVOH laminates with a proprietary cellulose-nano "
        "crystalline (CNC) coating on FSC kraft, achieving OTR < 0.8 cc/m²/day — the first "
        "compostable pouch certified for 12-month ambient shelf life on dry goods. Novel "
        "coating process uses aqueous deposition (no PFAS), scalable on existing flexographic "
        "lines with tooling retrofit estimated at $180k per line. Market: $4.2B sustainable "
        "flexible packaging segment growing 14% CAGR; validated demand via 14 CPG brand "
        "interviews and a paid pilot with a regional snack company (8-week shelf test passed). "
        "Team has FDA food-contact experience; FSC and BPI compostable certifications secured. "
        "Unit economics demonstrated at pilot scale with measured barrier data and customer "
        "renewal contract signed.",
        "PDF", "Market Test", "Approved",
    ),
    (
        "VoltEdge 800V SiC Traction Inverter Module",
        "VoltEdge Power Systems is developing a proprietary 800V silicon-carbide inverter module "
        "for Tier-2 EV OEMs, delivering 98.4% peak efficiency and 40% volume reduction versus "
        " incumbent IGBT designs. Patent pending on integrated liquid-cooling manifold. Working "
        "prototype tested at 200kW continuous; BOM modeled at volume (50k units/yr) shows "
        "18% cost advantage. Supply chain secured with Wolfspeed die sourcing and domestic "
        "assembly partner. Automotive ISO 26262 functional safety roadmap in place; CE and "
        "UN/ECE R100 testing scheduled Q3. Founding team: ex-Tesla powertrain engineer (led "
        "Model 3 inverter), ex-BorgWarner manufacturing director. LOI from mid-size EV startup "
        "for 15k units/year. Traction data and thermal test results documented.",
        "PDF", "Validation", "In Review",
    ),
    (
        "LoopFrame Open-Repair Smart Speaker",
        "LoopFrame Audio is building a consumer smart speaker designed for 10-year service life: "
        "proprietary modular driver assembly, standard fastener layout, published repair manuals, "
        "and replaceable compute module. Breakthrough industrial design with magnetic faceplate — "
        "patent filed on the novel snap-fit architecture. Differentiated in a crowded market via "
        "circular design and EU Right-to-Repair compliance built-in from day one. Functional "
        "prototype demonstrated at CES; 340-unit beta sold out in 72 hours with 4.6★ reviews. "
        "Manufacturing partner in Shenzhen with audited ethical sourcing (SA8000); scalable "
        "production line qualified. BOM $38 at 25k units; $149 retail with 44% margin. Founder "
        "previously shipped 200k units at Sonos competitor. Recyclable aluminum enclosure; LCA "
        "shows 62% lower e-waste versus sealed competitors. Measured repair-cycle data cited.",
        "Video", "Prototyping", "In Review",
    ),
    (
        "TerraCrisp Fermented Plant-Protein Crisp",
        "TerraCrisp Foods uses a novel solid-state fermentation process to convert upcycled "
        "brewer's spent grain into a crunchy, shelf-stable protein crisp with 22g complete "
        "protein per 100g and neutral flavor profile — a first in the ambient CPG aisle. "
        "Process breakthrough reduces water use 90% versus wet extrusion. Prototype batches "
        "validated in university sensory study (n=186, 78% preference vs. leading soy crisp). "
        "Co-manufacturing agreement with regional food processor; scalable production line "
        "mapped. SAM $890M (better-for-you savory snacks); early traction: 3 retail LOIs, "
        "Amazon launch pilot at 400 units/week sell-through. Team includes fermentation PhD "
        "and former Kind Snacks VP Supply Chain. FDA GRAS pathway scoped; allergen testing "
        "complete with measured nutritional data cited.",
        "PDF", "Validation", "Scored",
    ),
    # ── Solid but average submissions ─────────────────────────────────────────
    (
        "ThreadMark Circular Workwear Platform",
        "ThreadMark embeds washable RFID tags in industrial workwear to track garment lifecycle, "
        "enabling corporate clients to reduce replacement spend and meet Scope 3 reporting "
        "requirements. Breakthrough first-of-kind integration of RFID hardware with a proprietary "
        "rental-and-repair routing platform — patent pending on the novel unique lifecycle "
        "tracking system and advanced circular workflow. Prototype vests tested with one "
        "municipal fleet (180 garments, 6-month pilot with validated results). Manufacturing "
        "via existing textile partner in Vietnam; supply chain established with scalable "
        "production to 5k units/month. Market segment is commercial workwear ($12B globally) "
        "with growing sustainability demand. Team has textile sourcing experience and one "
        "engineer who built RFID systems at a logistics startup. Revenue model: SaaS + "
        "per-garment fee; pricing validated with two customer contracts.",
        "PDF", "Concept", "Scored",
    ),
    (
        "AeroForge 3D-Printed Titanium Brake Caliper",
        "AeroForge Motorsport produces limited-run titanium brake calipers via laser powder-bed "
        "fusion, targeting track-day and club-racing enthusiasts. Proprietary lattice topology "
        "delivers breakthrough mass reduction — 35% lighter versus OEM cast iron with novel "
        "thermal management channels. Prototype dyno-tested on 10 vehicles; performance data "
        "measured and cited. Manufacturing is feasible at low volume (500 sets/yr) via qualified "
        "additive partner; tooling and supply chain for titanium powder secured. Niche segment "
        "with passionate customer base; commercial traction via forum pre-orders (42 sets). "
        "Founder is aerospace metallurgist with 3D printing background and proven track record. "
        "DOT compliance path for aftermarket automotive parts identified; safety testing "
        "scheduled.",
        "Image", "Prototyping", "Scored",
    ),
    (
        "BlueHarvest Algae Omega-3 Beverage",
        "BlueHarvest cultivates a proprietary microalgae strain to produce a neutral-tasting "
        "omega-3 RTD beverage — a novel fermentation breakthrough positioning against fish-oil "
        "supplements. Pilot bioreactor producing 200L/week; product formulation tested in focus "
        "groups with validated preference data. Market for algae omega-3 is growing; our unique "
        "strain offers differentiated DHA profile. Manufacturing path mapped: industrial "
        "bioreactor partner identified for scalable production from pilot scale. Team includes "
        "fermentation PhD and former beverage ops lead. Sustainability story is strong (no fish, "
        "lower carbon) with LCA draft completed. FDA GRAS pathway scoped; early B2B commercial "
        "interest from two supplement brands.",
        "PDF", "Concept", "Scored",
    ),
    (
        "ThermaWeave Phase-Change Athletic Base Layer",
        "ThermaWeave integrates micro-encapsulated phase-change material (PCM) into a merino "
        "blend base layer via a proprietary knitting pattern that improves drape and wash "
        "durability — a novel textile process validated in lab testing. Prototype garments "
        "manufactured (500 units) via Portuguese knitwear partner with scalable production "
        "capacity. Market: premium athletic apparel ($28B segment) with validated customer "
        "demand via DTC pre-order campaign ($34k raised). Supply chain for PCM capsules secured "
        "from German supplier; BOM and margin analysis complete. Team is two former Under Armour "
        "designers with manufacturing partner relationships. Recyclable packaging; sustainable "
        "merino sourcing with ethical certification.",
        "PDF", "Validation", "Scored",
    ),
    (
        "PackRight Reusable Shipper Insert System",
        "PackRight replaces single-use foam and bubble wrap with a modular corrugated insert "
        "system designed for 20+ reuse cycles in e-commerce fulfillment. Recyclable, "
        "cost-neutral at 8+ cycles per internal model. Prototype tested with one mid-size "
        "3PL (4,000 shipments); damage rate comparable to foam. Innovation is incremental — "
        "reusable packaging concepts exist (Limeloop, RePack) — but our fold-flat design "
        "reduces return-shipping cost 40%. Manufacturing via standard die-cut corrugated "
        "partner; scalable. Market validated through 3PL pilot but broader demand unproven. "
        "Founder has packaging engineering background; early B2B pricing model drafted.",
        "PDF", "Market Test", "Scored",
    ),
    (
        "Heritage Mill Ancient-Grain Pasta Co-Pack",
        "Heritage Mill partners with regional grain farmers to produce bronze-die pasta from "
        "einkorn and emmer under a unique co-brand retail model — a novel farm-to-shelf "
        "traceability breakthrough via QR-linked harvest data and first-of-kind grower revenue "
        "share. Prototype production runs complete at existing Italian co-packer with scalable "
        "manufacturing capacity. Supply chain for specialty grains is seasonal but contracts "
        "secured with three farms for 2026 harvest. Market is moderate-growth premium pasta "
        "($1.1B US) with validated demand from two regional grocers. Team includes a "
        "chef-founder with proven CPG launch experience and operations lead. Sustainable and "
        "ethical sourcing documented; compostable packaging pilot underway.",
        "PDF", "Concept", "Scored",
    ),
    (
        "SilKote Bio-Based Marine Anti-Fouling Coating",
        "SilKote Marine develops a silicone foul-release coating using bio-based resin "
        "feedstock, reducing copper biocide content 60% versus conventional antifouling "
        "paints. Novel polymer chemistry with one provisional patent. Lab and static panel "
        "testing show 18-month performance; dynamic hull testing underway. Manufacturing "
        "requires specialty batch reactor — partner identified in Netherlands. Regulatory "
        "path includes IMO biocide review and regional VOC compliance. Niche B2B market "
        "($800M antifouling coatings) with long sales cycles. Team: polymer chemist + marine "
        "industry sales veteran with proven track record at Hempel.",
        "PDF", "Scaling", "In Review",
    ),
    (
        "CarbonCork Negative-Emission Vehicle Interior Trim",
        "CarbonCork produces injection-molded interior trim panels from a proprietary cork "
        "composite blended with captured-CO₂ mineral filler — a breakthrough material and novel "
        "manufacturing process achieving net-negative embodied carbon for EV OEMs. Patent "
        "pending on the composite formulation; first automotive application of this bio-based "
        "carbon sequestration approach. Material prototype meets OEM flammability screening "
        "(FMVSS 302 pending full test). Manufacturing via existing cork processor and twin-screw "
        "compounding; scalable to 100k panels/month with identified Tier-1 integration path. "
        "Market pull from automotive sustainability mandates (EU fleet regulations). Team "
        "includes former Faurecia materials engineer with proven track record. LOI from "
        "European EV startup for concept vehicle fitment. LCA demonstrates 1.8 kg CO₂e "
        "sequestered per panel.",
        "PDF", "Monitoring", "Approved",
    ),
    # ── Auto-reject candidates (weak rubric signals) ──────────────────────────
    (
        "LuxBasic Premium Cotton Tee Rebrand",
        "LuxBasic is launching a direct-to-consumer premium cotton t-shirt — a me-too product "
        "in a saturated, declining basics market. No material innovation, no proprietary "
        "process, no IP. We are essentially repackaging a standard commodity blank with a "
        "logo and influencer marketing. Existing competitors (Everlane, Cuts, Uniqlo) dominate "
        "at lower price points. No prototype beyond generic blanks from Alibaba; manufacturing "
        "is standard cut-and-sew with no differentiation. Vague hype about 'disrupting fashion' "
        "without evidence. No relevant team experience in apparel manufacturing. We believe "
        "our brand story will carry it — unsubstantiated. No revenue model beyond DTC hope.",
        "PDF", "Concept", "Rejected",
    ),
    (
        "HydraGlow LED Smart Water Bottle",
        "HydraGlow adds a generic LED strip and basic hydration reminder app to a standard "
        "stainless steel water bottle — a simple copy of dozens of existing Amazon listings. "
        "No novel technology, no patent, commodity product in a crowded market. Manufacturing "
        "is off-the-shelf OEM from Shenzhen with our logo. Concept only — no working prototype "
        "beyond a 3D-printed mockup. Speculative TAM claims ('every gym-goer needs one'). "
        "Vague marketing language with no validated customer demand or test data. First-time "
        "founders with no hardware or supply chain experience. Unclear revenue model at $12 "
        "retail with no margin analysis.",
        "Image", "Intake", "Rejected",
    ),
    (
        "Perpetua Motion Self-Charging Desk Device",
        "Perpetua Labs claims to have built a desk ornament that generates perpetual motion "
        "energy from ambient vibration — concept only, purely theoretical with no working "
        "prototype. We believe this violates no laws of physics because of our 'quantum "
        "harmonic resonance' design — speculative and unsubstantiated. No manufacturing path, "
        "unclear process, unproven at scale. Heavy on hype, light on evidence. No team with "
        "relevant engineering background; no compliance plan for consumer electronics (no CE, "
        "no FCC). No revenue model — planning to license the 'technology' with no LOIs.",
        "PDF", "Concept", "Rejected",
    ),
    (
        "WrapJoy Monthly Virgin Plastic Gift Box",
        "WrapJoy is a subscription service shipping monthly gift boxes made entirely from virgin "
        "plastic with synthetic fragrance inserts — no sustainability consideration, no "
        "certification, destined for landfill after single use. Greenwash marketing claims "
        "products are 'eco-luxe' without measurable metrics or LCA. Toxic fragrance compounds "
        "with no safety testing or regulatory compliance plan. Standard me-too subscription "
        "box in a saturated market. No prototype beyond a Canva mockup. Vague hype about "
        "'wellness revolution.' Inexperienced team with no CPG track record. No clear path "
        "to revenue — free trial boxes with unclear unit economics.",
        "PDF", "Intake", "Rejected",
    ),
    # ── Unscored (fresh intake) ─────────────────────────────────────────────────
    (
        "KineticHub Modular EV Battery Enclosure",
        "KineticHub is developing a modular, crash-safe battery enclosure system for mid-volume "
        "EV platforms, using a proprietary aluminum-extrusion lattice that enables pack "
        "capacity swaps without full vehicle redesign. Early CAD and FEA simulations complete; "
        "physical prototype build scheduled next quarter. Target customers are EV startups "
        "needing flexible pack architectures. Submission includes industrial design renders "
        "and preliminary BOM — awaiting full rubric scoring at intake review.",
        "PDF", "Intake", "New",
    ),
    (
        "HarvestFold Collapsible Insulated Grocery Crate",
        "HarvestFold combines a collapsible corrugated frame with vacuum-insulated panels for "
        "last-mile grocery delivery — keeps chilled items at temperature for 4+ hours without "
        "refrigerant. Early prototype tested in 12 delivery routes with a regional grocer. "
        "Team submitting initial concept deck and pilot data for ForgeOS intake scoring.",
        "PDF", "Intake", "New",
    ),
]


def add_demo_submissions():
    for name, notes, ftype, stage, status in DEMO_SUBMISSIONS:
        sid = f"FOS-{st.session_state.next_id}"
        st.session_state.next_id += 1
        sub_stub = {"name": name, "notes": notes}
        if status != "New":
            scores = ai_score_submission(sub_stub, rubric)
        else:
            scores = {
                "overall": 0.0, "innovation": 0.0, "feasibility": 0.0,
                "categories": {}, "auto_reject": [], "high_risk": [], "scored_at": "",
            }
        days_ago = random.randint(1, 45)
        base_dt  = datetime.now() - timedelta(days=days_ago)
        # Build stage history: one entry per stage up to the current one
        history = []
        for sn in STAGE_NAMES:
            entry_dt = base_dt + timedelta(days=STAGE_NAMES.index(sn) * 3)
            history.append({"stage": sn, "moved_at": entry_dt.strftime("%Y-%m-%d")})
            if sn == stage:
                break
        st.session_state.submissions.append({
            "id":             sid,
            "name":           name,
            "file_type":      ftype,
            "status":         status,
            "stage":          stage,
            "overall":        scores["overall"],
            "innovation":     scores["innovation"],
            "feasibility":    scores["feasibility"],
            "categories":     scores["categories"],
            "auto_reject":    scores.get("auto_reject", []),
            "high_risk":      scores.get("high_risk", []),
            "scored_at":      scores.get("scored_at", ""),
            "stage_summary":  generate_stage_summary({"name": name, "overall": scores["overall"]}, stage) if stage != "Intake" else "",
            "stage_history":  history,
            "extracted_text": "",
            "file_summaries": [],
            "submitted_at":  base_dt.strftime("%Y-%m-%d"),
            "notes":          notes,
        })


def _build_landing_mockup_html() -> str:
    """Rich CSS dashboard preview for the landing hero."""
    bar_heights = [28, 44, 36, 52, 40, 48, 32, 56, 38, 46, 42, 50]
    bars = "".join(
        f'<div class="landing-mock-bar" style="height:{h}%;opacity:{0.45 + (i % 4) * 0.12};"></div>'
        for i, h in enumerate(bar_heights)
    )
    kanban_cols = [
        ("Concept", "#1f6feb", [("BioWrap Produce Film", "82", "#22c55e"), ("SolarPack Mailer", "74", "#22c55e")]),
        ("Validation", "#0ea5e9", [("FlexiBottle", "68", "#f59e0b")]),
        ("Prototyping", "#238636", [("CarbonCork Trim", "91", "#22c55e")]),
        ("Market Test", "#9e6a03", [("Perpetua Motion Toy", "12", "#f43f5e")]),
    ]
    kanban_html = ""
    for name, color, cards in kanban_cols:
        cards_html = ""
        for title, score, sc in cards:
            cards_html += (
                f'<div class="landing-mock-k-card">{_esc(title)}'
                f'<div class="landing-mock-k-score" style="color:{sc};">Score {score}</div>'
                f"</div>"
            )
        kanban_html += (
            f'<div class="landing-mock-k-col">'
            f'<div class="landing-mock-k-hd" style="border-color:{color};">{_esc(name)}</div>'
            f"{cards_html}</div>"
        )
    return f"""
    <div class="landing-mock-wrap">
    <div class="landing-mockup-shell">
      <div class="landing-mock-topbar">
        <div class="landing-mock-dots">
          <span class="landing-mock-dot" style="background:#f43f5e;"></span>
          <span class="landing-mock-dot" style="background:#f59e0b;"></span>
          <span class="landing-mock-dot" style="background:#22c55e;"></span>
        </div>
        <span style="font-size:10px;color:var(--text-3);font-weight:600;">ForgeOS · Innovation OS</span>
        <span style="font-size:9px;color:var(--green);font-weight:700;">● Live</span>
      </div>
      <div class="landing-mock-body">
        <div class="landing-mock-sidebar">
          <div class="landing-mock-nav-item active">Dashboard</div>
          <div class="landing-mock-nav-item">Submissions</div>
          <div class="landing-mock-nav-item">Shortlist</div>
          <div class="landing-mock-nav-item">Pipeline</div>
          <div class="landing-mock-nav-item">Rubric</div>
        </div>
        <div class="landing-mock-main">
          <div class="landing-mock-forge-bar">
            <div class="landing-mock-breadcrumb">ForgeOS <span>/ Dashboard</span></div>
            <div class="landing-mock-status">
              <span class="landing-mock-status-dot"></span> 19 ideas tracked
            </div>
          </div>
          <div class="landing-mock-kpis">
            <div class="landing-mock-kpi">
              <div class="landing-mock-kpi-lbl">Submissions</div>
              <div class="landing-mock-kpi-val">19</div>
            </div>
            <div class="landing-mock-kpi">
              <div class="landing-mock-kpi-lbl">Avg Score</div>
              <div class="landing-mock-kpi-val" style="color:#22c55e;">71</div>
            </div>
            <div class="landing-mock-kpi">
              <div class="landing-mock-kpi-lbl">High Potential</div>
              <div class="landing-mock-kpi-val" style="color:#3b82f6;">8</div>
            </div>
            <div class="landing-mock-kpi">
              <div class="landing-mock-kpi-lbl">Approved</div>
              <div class="landing-mock-kpi-val" style="color:#a78bfa;">4</div>
            </div>
          </div>
          <div class="landing-mock-chart">{bars}</div>
          <table class="landing-mock-table">
            <thead><tr>
              <th>ID</th><th>Idea</th><th>Score</th><th>Status</th>
            </tr></thead>
            <tbody>
              <tr>
                <td>1001</td><td>BioWrap Produce Film</td>
                <td><span class="landing-mock-badge landing-mock-badge-green">82</span></td>
                <td><span class="landing-mock-badge landing-mock-badge-blue">Scored</span></td>
              </tr>
              <tr>
                <td>1004</td><td>CarbonCork Interior Trim</td>
                <td><span class="landing-mock-badge landing-mock-badge-green">91</span></td>
                <td><span class="landing-mock-badge landing-mock-badge-amber">In Review</span></td>
              </tr>
              <tr>
                <td>1018</td><td>Perpetua Motion Desk Toy</td>
                <td><span class="landing-mock-badge" style="background:rgba(244,63,94,0.15);color:#f43f5e;">12</span></td>
                <td><span class="landing-mock-badge" style="background:rgba(244,63,94,0.15);color:#f43f5e;">Rejected</span></td>
              </tr>
            </tbody>
          </table>
          <div class="landing-mock-pipeline-hd" style="margin-top:10px;">Pipeline Kanban</div>
          <div class="landing-mock-kanban">{kanban_html}</div>
        </div>
      </div>
    </div>
    <div class="landing-mock-float">
      <div class="landing-mock-float-lbl">Rubric v2 · Top Score</div>
      <div class="landing-mock-float-val">91</div>
      <div class="landing-mock-float-sub">CarbonCork Interior Trim</div>
    </div>
    <div class="landing-mock-score-chip">
      <div class="landing-mock-score-chip-num">8</div>
      <div class="landing-mock-score-chip-lbl">High Potential</div>
    </div>
    </div>"""


def _build_landing_hero_html() -> str:
    """Unified hero section — copy + dashboard mockup in one HTML block."""
    return f"""
    <section class="landing-hero-section">
      <div class="landing-hero-glow landing-hero-glow--blue"></div>
      <div class="landing-hero-glow landing-hero-glow--green"></div>
      <div class="landing-hero-glow landing-hero-glow--purple"></div>
      <div class="landing-hero-inner">
        <div class="landing-hero-copy">
          <div class="landing-eyebrow">
            <span class="landing-eyebrow-dot"></span>
            Agentic innovation intelligence
          </div>
          <h1 class="landing-h1">
            ForgeOS — The <em>AI Agentic</em> Innovation OS
          </h1>
          <p class="landing-sub">
            From raw ideas to commercialized products. Intelligently screen thousands
            of submissions, apply a world-class physical-goods rubric, and advance the
            best innovations through the full pipeline.
          </p>
        </div>
        <div class="landing-hero-visual">
          {_build_landing_mockup_html()}
        </div>
      </div>
    </section>"""


def _build_landing_cta_html() -> str:
    """Centered CTA bar — aligned with benefits section below."""
    return """
    <div class="landing-hero-cta-bar">
      <a class="landing-cta-btn" href="?dashboard=1">Launch ForgeOS Dashboard</a>
      <p class="landing-hero-micro">No setup required · Load demo data from the dashboard sidebar</p>
    </div>"""


def _render_landing_page():
    """Full-screen marketing landing — default app entry point."""
    st.markdown("""
    <div class="landing-nav">
      <div class="landing-nav-brand">
        <div class="sb-logo-icon" style="width:32px;height:32px;font-size:16px;">⚙</div>
        ForgeOS
        <span class="landing-nav-badge">Innovation OS</span>
      </div>
      <div style="font-size:12px;color:var(--text-3);font-weight:500;">
        AI Agentic Innovation OS
      </div>
    </div>""", unsafe_allow_html=True)

    st.markdown(_build_landing_hero_html(), unsafe_allow_html=True)
    st.markdown(_build_landing_cta_html(), unsafe_allow_html=True)

    # ── Benefit cards ─────────────────────────────────────────────────────────
    st.markdown('<div id="landing-benefits" class="landing-benefits">', unsafe_allow_html=True)
    st.markdown(
        '<div class="landing-benefits-hd">Built for innovation</div>',
        unsafe_allow_html=True,
    )
    b1, b2, b3, b4 = st.columns(4)
    benefits = [
        (b1, "📎", "Multimodal Intake",
         "Ingest PDFs, images, video, and docs. Vision-aware parsing extracts evidence "
         "for richer, fairer scoring."),
        (b2, "📊", "Rubric Scoring",
         "8-criterion Extensive Rubric v2 with gating rules, LLM or simulated scoring, "
         "and investment-ready memos."),
        (b3, "🔀", "Agentic Pipeline",
         "Seven-stage flow from Intake to Monitoring — advance ideas with AI stage "
         "briefs and kanban visibility."),
        (b4, "⚡", "Speed & Scale",
         "Bulk-score submissions in parallel. Shortlist by category. Filter, sort, "
         "and triage thousands of ideas fast."),
    ]
    for col, icon, title, desc in benefits:
        with col:
            st.markdown(f"""
            <div class="landing-benefit-card">
              <span class="landing-benefit-icon">{icon}</span>
              <div class="landing-benefit-title">{_esc(title)}</div>
              <p class="landing-benefit-desc">{_esc(desc)}</p>
            </div>""", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Footer ────────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="landing-footer">
      <p style="font-size:14px;font-weight:600;color:var(--text-2);margin-bottom:8px;">
        ForgeOS · The AI Agentic Innovation OS
      </p>
      <p>Built as a hackathon MVP · Extensive Rubric v2 · Physical Goods</p>
      <p style="margin-top:12px;font-size:10px;color:var(--text-3);">
        © 2026 ForgeOS · From raw ideas to commercialized products
      </p>
    </div>
    <div style="height:24px;"></div>""", unsafe_allow_html=True)


# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="sb-logo">
        <div class="sb-wordmark">
            <div class="sb-logo-icon">⚙</div>
            ForgeOS
        </div>
        <div class="sb-tag">Innovation OS</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sb-section-label">Main</div>', unsafe_allow_html=True)

    if st.session_state.nav_page not in NAV_OPTIONS:
        st.session_state.nav_page = "Home"
    page = st.radio(
        "nav",
        NAV_OPTIONS,
        index=NAV_OPTIONS.index(st.session_state.nav_page),
        label_visibility="collapsed",
    )
    prev_nav = st.session_state.get("_tracked_nav_page")
    if prev_nav is not None and prev_nav != page:
        _close_criterion_detail_dialog()
    st.session_state._tracked_nav_page = page
    st.session_state.nav_page = page
    # #region agent log
    _debug_log(
        "sidebar navigation resolved",
        {
            "radio_page": page,
            "prev_nav": prev_nav,
            "query_page": st.query_params.get("page"),
            "query_pipeline": st.query_params.get("pipeline"),
            "query_submission": st.query_params.get("submission"),
        },
        hypothesis_id="H1",
        location="app.py:6520",
    )
    # #endregion
    _sync_pipeline_detail_from_query()

    subs = st.session_state.submissions
    _purge_shortlist_orphans()
    sl_total = _shortlist_total_count()
    total    = len(subs)
    scored   = sum(1 for s in subs if s["status"] in ("Scored","In Review","Approved"))
    approved = sum(1 for s in subs if s["status"] == "Approved")
    high_pot = sum(1 for s in subs if s["overall"] >= THRESHOLDS["green"])
    avg_score= round(sum(s["overall"] for s in subs) / max(total, 1), 1)

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-section-label">Overview</div>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="sb-stat-row"><span class="sb-stat-label">Submissions</span><span class="sb-stat-val">{total}</span></div>
    <div class="sb-stat-row"><span class="sb-stat-label">Scored</span><span class="sb-stat-val">{scored}</span></div>
    <div class="sb-stat-row"><span class="sb-stat-label">Approved</span><span class="sb-stat-val">{approved}</span></div>
    <div class="sb-stat-row"><span class="sb-stat-label">High Potential</span><span class="sb-stat-val" style="color:#3fb950">{high_pot}</span></div>
    <div class="sb-stat-row"><span class="sb-stat-label">Avg Score</span><span class="sb-stat-val" style="color:{score_hex(avg_score) if total else '#8b949e'}">{avg_score if total else '—'}</span></div>
    <div class="sb-stat-row"><span class="sb-stat-label">Shortlist</span><span class="sb-stat-val" style="color:#f59e0b">{sl_total}</span></div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    if st.button("Load Demo Data", use_container_width=True):
        if not st.session_state.submissions:
            add_demo_submissions()
            st.rerun()
        else:
            st.warning("Submissions already loaded.")

    st.markdown(f"""
    <div style="padding:16px 14px 12px;border-top:1px solid var(--border);margin-top:12px;">
        <div style="font-size:10px;font-weight:700;color:var(--text-3);text-transform:uppercase;letter-spacing:0.09em;">ForgeOS v0.2</div>
        <div style="font-size:10px;color:var(--text-3);margin-top:3px;">AI Scoring · Extensive Rubric v2</div>
        <div style="margin-top:10px;display:flex;align-items:center;gap:6px;">
          <div style="width:5px;height:5px;border-radius:50%;background:var(--green);box-shadow:0 0 6px var(--green);"></div>
          <span style="font-size:10px;color:var(--text-3);">Engine ready</span>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: HOME (Landing)
# ══════════════════════════════════════════════════════════════════════════════
if page == "Home":
    _render_landing_page()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE: DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Dashboard":

    st.markdown(f"""
    <div class="forge-topbar">
      <div class="forge-topbar-left">
        <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Dashboard</span></div>
        <div class="forge-page-tag">Analytics</div>
      </div>
      <div class="forge-topbar-status">
        <div class="forge-status-dot"></div>
        {total} idea{"s" if total != 1 else ""} tracked
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="page-content">', unsafe_allow_html=True)

    # ── Apollo-style stat strip ───────────────────────────────────────────────
    avg_innov = round(sum(s["innovation"] for s in subs if s["innovation"] > 0) / max(sum(1 for s in subs if s["innovation"] > 0), 1), 1)
    avg_feas  = round(sum(s["feasibility"] for s in subs if s["feasibility"] > 0) / max(sum(1 for s in subs if s["feasibility"] > 0), 1), 1)

    _kpi_cols = st.columns(6)
    _kpi_items = [
        (_kpi_cols[0], "📋", "Total Ideas",      str(total),                                                                  "All ideas tracked",              "#ddeaf8"),
        (_kpi_cols[1], "📊", "Avg Score",         str(avg_score) if total else "—",                                           "Out of 100 points",              score_hex(avg_score) if total else "#4e6680"),
        (_kpi_cols[2], "🚀", "High Potential",    str(high_pot),                                                               f"Score ≥ {THRESHOLDS['green']}",  "#22c55e"),
        (_kpi_cols[3], "✅", "Approved",           str(approved),                                                               "Moving to production",           "#a78bfa"),
        (_kpi_cols[4], "💡", "Avg Innovation",    str(avg_innov) if avg_innov else "—",                                        "Innovation criterion",           score_hex(avg_innov) if avg_innov else "#4e6680"),
        (_kpi_cols[5], "🔧", "Avg Feasibility",   str(avg_feas) if avg_feas else "—",                                         "Manufacturing score",            score_hex(avg_feas) if avg_feas else "#4e6680"),
    ]
    for _col, _icon, _label, _val, _sub, _color in _kpi_items:
        with _col:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-icon">{_icon}</div>
                <div class="kpi-label">{_label}</div>
                <div class="kpi-value" style="color:{_color};">{_val}</div>
                <div class="kpi-sub">{_sub}</div>
            </div>""", unsafe_allow_html=True)

    if subs:
        # ── Charts row ────────────────────────────────────────────────────────
        col_c1, col_c2 = st.columns([3, 2], gap="medium")

        with col_c1:
            st.markdown('<div class="section-hd">Score Distribution</div>', unsafe_allow_html=True)
            df = pd.DataFrame(subs)
            scored_df = df[df["overall"] > 0]
            if not scored_df.empty:
                fig_hist = px.histogram(
                    scored_df, x="overall", nbins=12,
                    color_discrete_sequence=["#3b82f6"],
                )
                fig_hist.update_traces(marker_line_width=0)
                fig_hist.update_layout(
                    paper_bgcolor="#0f1e30", plot_bgcolor="#0f1e30",
                    font={"color": "#4e6680", "family": "Inter", "size": 11},
                    xaxis=dict(gridcolor="#1a2f49", color="#4e6680", title="Overall Score", showgrid=True),
                    yaxis=dict(gridcolor="#1a2f49", color="#4e6680", title="Count", showgrid=True),
                    height=230, margin=dict(l=0, r=0, t=8, b=0),
                    bargap=0.10,
                )
                st.plotly_chart(fig_hist, use_container_width=True)
            else:
                st.markdown('<div style="height:220px;display:flex;align-items:center;justify-content:center;color:#6e7681;font-size:13px;">No scored submissions yet</div>', unsafe_allow_html=True)

        with col_c2:
            st.markdown('<div class="section-hd">Status Breakdown</div>', unsafe_allow_html=True)
            status_counts = pd.DataFrame(subs)["status"].value_counts().reset_index()
            status_counts.columns = ["Status", "Count"]
            fig_pie = px.pie(
                status_counts, names="Status", values="Count",
                color_discrete_sequence=["#3b82f6","#22c55e","#f59e0b","#a78bfa","#f43f5e","#38bdf8","#fb923c"],
                hole=0.6,
            )
            fig_pie.update_layout(
                paper_bgcolor="#0f1e30", plot_bgcolor="#0f1e30",
                font={"color": "#4e6680", "family": "Inter", "size": 11},
                height=230, margin=dict(l=0, r=0, t=8, b=0),
                legend=dict(font=dict(color="#8aa4c0", size=11), bgcolor="#0f1e30", bordercolor="#1a2f49", borderwidth=1),
                showlegend=True,
            )
            fig_pie.update_traces(textfont_color="#ddeaf8", textfont_size=11)
            st.plotly_chart(fig_pie, use_container_width=True)

        # ── Recent submissions table ───────────────────────────────────────
        st.markdown('<div class="section-hd">Recent Submissions</div>', unsafe_allow_html=True)
        recent = sorted(subs, key=lambda x: x["submitted_at"], reverse=True)[:8]

        rows_html = ""
        for sub in recent:
            badge = ""
            if sub["overall"] > 0:
                bc = score_badge_class(sub["overall"])
                badge = f'<span class="badge-score {bc}">{sub["overall"]}</span>'
            else:
                badge = '<span style="color:#6e7681;font-size:12px;">—</span>'
            pc = pill_class(sub["status"])
            stage_color = next((s["color"] for s in STAGES if s["name"] == sub["stage"]), "#8b949e")
            rows_html += f"""
            <tr class="forge-tr">
              <td class="forge-td"><span class="forge-id">{_esc(sub['id'])}</span></td>
              <td class="forge-td forge-td-primary">{_esc(sub['name'])}</td>
              <td class="forge-td">{badge}</td>
              <td class="forge-td"><span class="pill {pc}">{_esc(sub['status'])}</span></td>
              <td class="forge-td"><span style="font-size:11px;color:{stage_color};font-weight:600;">{_esc(sub['stage'])}</span></td>
              <td class="forge-td">{_esc(sub['submitted_at'])}</td>
            </tr>"""

        st.markdown(f"""
        <div class="forge-card" style="padding:0; overflow:hidden;">
        <table class="forge-table">
          <thead>
            <tr>
              <th class="forge-th">ID</th>
              <th class="forge-th">Idea Name</th>
              <th class="forge-th">Score</th>
              <th class="forge-th">Status</th>
              <th class="forge-th">Stage</th>
              <th class="forge-th">Submitted</th>
            </tr>
          </thead>
          <tbody>{rows_html}</tbody>
        </table>
        </div>
        """, unsafe_allow_html=True)

        # ── Pipeline snapshot ──────────────────────────────────────────────
        st.markdown('<div class="section-hd">Pipeline Snapshot</div>', unsafe_allow_html=True)
        stage_counts_map = {s["name"]: 0 for s in STAGES}
        for sub in subs:
            if sub["stage"] in stage_counts_map:
                stage_counts_map[sub["stage"]] += 1

        _STAGE_ICONS = {
            "Intake": "📥", "Concept": "💡", "Validation": "🔬",
            "Prototyping": "🛠", "Market Test": "📈", "Scaling": "⚡", "Monitoring": "📡",
        }
        cols = st.columns(len(STAGES))
        for stage, col in zip(STAGES, cols):
            cnt = stage_counts_map.get(stage["name"], 0)
            _ico = _STAGE_ICONS.get(stage["name"], "●")
            with col:
                st.markdown(f"""
                <div class="stage-flow-card" style="border-top-color:{stage['color']};">
                    <div class="stage-flow-icon">{_ico}</div>
                    <div class="stage-flow-num" style="color:{stage['color']};">{cnt}</div>
                    <div class="stage-flow-name">{stage['name']}</div>
                </div>""", unsafe_allow_html=True)

    else:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">⚙️</div>
            <div class="empty-title">No submissions yet</div>
            <div class="empty-sub">Go to Submissions to upload your first idea,<br>or click Load Demo Data in the sidebar.</div>
        </div>""", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: SUBMISSIONS
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Submissions":

    _sub_scored_ct = sum(1 for s in st.session_state.submissions if s["status"] not in ("New",))
    st.markdown(f"""
    <div class="forge-topbar">
      <div class="forge-topbar-left">
        <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Submissions</span></div>
        <div class="forge-page-tag">Idea Intake</div>
      </div>
      <div class="forge-topbar-status">
        <div class="forge-status-dot"></div>
        {_sub_scored_ct} scored
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="page-content">', unsafe_allow_html=True)

    # ── Flash message ─────────────────────────────────────────────────────────
    if st.session_state.flash_msg:
        ftype, fmsg = st.session_state.flash_msg
        st.session_state.flash_msg = None
        if ftype == "success":
            st.success(fmsg)
        elif ftype == "info":
            st.info(fmsg)
        elif ftype == "warning":
            st.warning(fmsg)
        else:
            st.error(fmsg)

    _maybe_open_shortlist_dialog()
    _maybe_open_memo_dialog()
    _maybe_open_criterion_deep_dive_dialog()

    # ── Upload panel ──────────────────────────────────────────────────────────
    _UPLOAD_TYPES = ["pdf", "png", "jpg", "jpeg", "gif", "webp", "bmp", "mp4", "mov", "avi", "webm", "txt", "md", "csv", "docx"]
    _UPLOAD_HELP = (
        "PDF: multi-method text extraction; auto OCR/vision fallback if text is sparse "
        "(vision requires API key). Images/video: optional LLM vision checkbox below."
    )

    with st.expander("➕  New Submission", expanded=not st.session_state.submissions):
        use_vis = False
        if _llm_ready():
            use_vis = st.checkbox(
                "🔍 Vision analysis for images",
                value=True,
                key="submit_use_vision",
                help="Uses LLM vision to describe image content. Requires Real LLM mode & API key.",
            )

        tab_single, tab_bulk = st.tabs(["Single idea", "Bulk upload"])

        with tab_single:
            col_f1, col_f2 = st.columns([2, 1])
            with col_f1:
                idea_name = st.text_input(
                    "Idea Name",
                    placeholder="e.g. Self-Healing Polymer Coating",
                    key="single_idea_name",
                )
                uploaded = st.file_uploader(
                    "Supporting files",
                    type=_UPLOAD_TYPES,
                    accept_multiple_files=True,
                    help=_UPLOAD_HELP,
                    key="single_idea_files",
                )
                _render_upload_file_previews(uploaded)

            with col_f2:
                notes_txt = st.text_area(
                    "Notes",
                    height=96,
                    placeholder="Brief context…",
                    key="single_idea_notes",
                )

            stage_analysis = _get_single_stage_analysis(idea_name, notes_txt, uploaded, use_vis)
            if uploaded and stage_analysis is None:
                st.info("Files are ready. Press Analyze to allow ForgeOS to inspect them and recommend a starting stage.")
                if st.button(
                    "Analyze",
                    key="analyze_single_stage_recommendation",
                    use_container_width=False,
                ):
                    st.session_state.single_stage_analysis_request_sig = _stage_analysis_cache_key(
                        idea_name,
                        notes_txt,
                        uploaded,
                        use_vis,
                        cache_scope="single",
                    )
                    st.rerun()
            if stage_analysis:
                _render_stage_recommendation_card(stage_analysis)
                if st.button(
                    "Re-analyze",
                    key="refresh_single_stage_recommendation",
                    use_container_width=False,
                ):
                    st.session_state.single_stage_analysis = None
                    st.session_state.single_stage_analysis_request_sig = _stage_analysis_cache_key(
                        idea_name,
                        notes_txt,
                        uploaded,
                        use_vis,
                        cache_scope="single",
                    )
                    st.rerun()

            if st.button("Submit Idea", key="submit_single_idea"):
                _close_criterion_detail_dialog()
                if not idea_name.strip():
                    st.error("Idea name is required.")
                else:
                    if stage_analysis is None:
                        if uploaded:
                            st.session_state.single_stage_analysis_request_sig = _stage_analysis_cache_key(
                                idea_name,
                                notes_txt,
                                uploaded,
                                use_vis,
                                cache_scope="single",
                            )
                        stage_analysis = _get_single_stage_analysis(idea_name, notes_txt, uploaded, use_vis)

                    stage_rec = stage_analysis.get("recommendation") if stage_analysis else None
                    if uploaded and stage_analysis and stage_analysis.get("file_signature") != "notes-only":
                        sid, file_count, total_chars = _append_submission_from_upload(
                            idea_name,
                            notes_txt,
                            uploaded,
                            use_vis,
                            precomputed_extracted=stage_analysis.get("extracted_text", ""),
                            precomputed_summaries=stage_analysis.get("file_summaries", []),
                            stage_recommendation=stage_rec,
                        )
                    else:
                        sid, file_count, _ = _append_submission_from_upload(
                            idea_name,
                            notes_txt,
                            uploaded,
                            use_vis,
                            stage_recommendation=stage_rec,
                        )
                    st.session_state.single_stage_analysis = None
                    st.session_state.single_stage_analysis_request_sig = None
                    file_note = f" ({file_count} file(s) parsed)" if file_count else ""
                    st.success(f"Submission {sid} added{file_note}.")
                    st.rerun()

        with tab_bulk:
            st.caption(
                "Add one row per idea — each with its own name, optional notes, and supporting files. "
                "Empty rows (no name) are skipped on submit."
            )
            bulk_rows: list[dict] = []
            for idx, row_id in enumerate(st.session_state.bulk_upload_row_ids):
                with st.container(border=True):
                    hdr_l, hdr_r = st.columns([5, 1])
                    with hdr_l:
                        st.markdown(f"**Idea {idx + 1}**")
                    with hdr_r:
                        if len(st.session_state.bulk_upload_row_ids) > 1 and st.button(
                            "Remove",
                            key=f"bulk_remove_{row_id}",
                            use_container_width=True,
                        ):
                            st.session_state.bulk_upload_row_ids.remove(row_id)
                            for prefix in ("bulk_name_", "bulk_notes_", "bulk_files_"):
                                st.session_state.pop(f"{prefix}{row_id}", None)
                            st.rerun()

                    b_name_col, b_notes_col = st.columns([1, 1])
                    with b_name_col:
                        row_name = st.text_input(
                            "Idea Name",
                            placeholder="e.g. Compostable Mailer Film",
                            key=f"bulk_name_{row_id}",
                        )
                    with b_notes_col:
                        row_notes = st.text_area(
                            "Notes (optional)",
                            height=68,
                            placeholder="Brief context for this idea…",
                            key=f"bulk_notes_{row_id}",
                        )
                    row_files = st.file_uploader(
                        "Files for this idea",
                        type=_UPLOAD_TYPES,
                        accept_multiple_files=True,
                        key=f"bulk_files_{row_id}",
                        label_visibility="visible",
                    )
                    if row_files:
                        st.caption(
                            f"{len(row_files)} file(s): "
                            + ", ".join(f.name for f in row_files[:4])
                            + (f" + {len(row_files) - 4} more" if len(row_files) > 4 else "")
                        )
                    row_stage_analysis = None
                    row_cache_key = _stage_analysis_cache_key(
                        row_name,
                        row_notes,
                        row_files,
                        use_vis,
                        cache_scope=f"bulk-{row_id}",
                    )
                    if (row_name or "").strip() or (row_notes or "").strip() or row_files:
                        row_stage_analysis = _get_stage_analysis(
                            row_name,
                            row_notes,
                            row_files,
                            use_vis,
                            cache_scope=f"bulk-{row_id}",
                            show_spinner=False,
                            require_explicit_run=bool(row_files),
                        )
                        if row_files and row_stage_analysis is None:
                            if st.button(
                                "Analyze",
                                key=f"bulk_analyze_{row_id}",
                                use_container_width=False,
                            ):
                                requested = set(
                                    st.session_state.bulk_stage_analysis_cache.get("__requested__", set())
                                )
                                requested.add(row_cache_key)
                                st.session_state.bulk_stage_analysis_cache["__requested__"] = requested
                                st.rerun()
                        if row_stage_analysis:
                            _render_stage_recommendation_card(row_stage_analysis)
                    bulk_rows.append({
                        "row_id": row_id,
                        "cache_key": row_cache_key,
                        "name": row_name,
                        "notes": row_notes,
                        "files": row_files,
                        "stage_analysis": row_stage_analysis,
                    })

            bulk_btn_l, bulk_btn_r = st.columns([1, 1])
            with bulk_btn_l:
                if st.button("➕ Add another idea", key="bulk_add_row"):
                    if len(st.session_state.bulk_upload_row_ids) >= 15:
                        st.warning("Maximum 15 ideas per bulk upload batch.")
                    else:
                        st.session_state.bulk_upload_row_ids.append(
                            st.session_state.bulk_upload_next_row_id
                        )
                        st.session_state.bulk_upload_next_row_id += 1
                        st.rerun()
            with bulk_btn_r:
                submit_bulk = st.button(
                    f"Submit all ideas ({len(st.session_state.bulk_upload_row_ids)})",
                    key="submit_bulk_ideas",
                    type="primary",
                    use_container_width=True,
                )

            if submit_bulk:
                _close_criterion_detail_dialog()
                valid_rows = [r for r in bulk_rows if (r["name"] or "").strip()]
                if not valid_rows:
                    st.error("Add at least one idea with a name.")
                else:
                    created: list[str] = []
                    total_files = 0
                    total_chars = 0
                    with st.status(
                        f"Submitting {len(valid_rows)} idea(s)…",
                        expanded=True,
                    ) as bulk_status:
                        slot = st.empty()
                        for i, row in enumerate(valid_rows):
                            slot.markdown(
                                f'<div style="font-size:12px;color:#8b949e;padding:2px 0;">'
                                f'📥 **{i + 1}/{len(valid_rows)}** — {_esc(row["name"].strip())}</div>',
                                unsafe_allow_html=True,
                            )
                            row_analysis = row.get("stage_analysis")
                            if row_analysis is None:
                                if row["files"]:
                                    requested = set(
                                        st.session_state.bulk_stage_analysis_cache.get("__requested__", set())
                                    )
                                    requested.add(row["cache_key"])
                                    st.session_state.bulk_stage_analysis_cache["__requested__"] = requested
                                row_analysis = _get_stage_analysis(
                                    row["name"],
                                    row["notes"],
                                    row["files"],
                                    use_vis,
                                    cache_scope=f"bulk-{row['row_id']}",
                                    show_spinner=False,
                                    require_explicit_run=bool(row["files"]),
                                )
                            stage_rec_row = (
                                row_analysis.get("recommendation") if row_analysis else None
                            )
                            if row["files"] and row_analysis and row_analysis.get("file_signature") != "notes-only":
                                sid, fc, chars = _append_submission_from_upload(
                                    row["name"],
                                    row["notes"],
                                    row["files"],
                                    use_vis,
                                    precomputed_extracted=row_analysis.get("extracted_text", ""),
                                    precomputed_summaries=row_analysis.get("file_summaries", []),
                                    stage_recommendation=stage_rec_row,
                                )
                            else:
                                sid, fc, chars = _append_submission_from_upload(
                                    row["name"],
                                    row["notes"],
                                    row["files"],
                                    use_vis,
                                    stage_recommendation=stage_rec_row,
                                )
                            created.append(sid)
                            total_files += fc
                            total_chars += chars
                        slot.empty()
                        bulk_status.update(
                            label=(
                                f"✅ {len(created)} submission(s) added — "
                                f"{total_files} file(s), {total_chars:,} chars extracted"
                            ),
                            state="complete",
                        )
                    _reset_bulk_upload_form()
                    id_list = ", ".join(created[:6])
                    if len(created) > 6:
                        id_list += f", +{len(created) - 6} more"
                    st.session_state.flash_msg = (
                        "success",
                        f"Added {len(created)} submissions: {id_list}.",
                    )
                    st.rerun()

    # ── Toolbar ───────────────────────────────────────────────────────────────
    if st.session_state.submissions:
        tb1, tb2, tb3, tb4, tb5, tb6 = st.columns([3, 1.5, 1.5, 1.2, 1.1, 1.8])
        with tb1:
            search_q = st.text_input("search", placeholder="Search by name or ID…", label_visibility="collapsed")
        with tb2:
            f_status = st.selectbox("st", ["All Statuses","New","Scored","In Review","Approved","Rejected"], label_visibility="collapsed")
        with tb3:
            f_stage  = st.selectbox("sg", ["All Stages"] + STAGE_NAMES, label_visibility="collapsed")
        with tb4:
            f_sort   = st.selectbox(
                "sort",
                ["Newest", "Score (high first)", "Score (low first)", "Name"],
                label_visibility="collapsed",
            )
        with tb5:
            bulk_concurrency = st.number_input(
                "concurrency",
                min_value=1,
                max_value=10,
                value=st.session_state.get("bulk_concurrency", 5),
                step=1,
                label_visibility="collapsed",
                help="Number of submissions scored simultaneously. Higher = faster, but raises rate-limit risk with Real LLM mode.",
                key="bulk_concurrency",
            )
        with tb6:
            run_bulk = st.button("🤖  Process All with AI", use_container_width=True)

        if run_bulk:
            unscored = [s for s in st.session_state.submissions if s["status"] == "New"]
            if unscored:
                # ── Snapshot ALL config before spawning threads ─────────────────
                # Workers must never read st.session_state — pass everything in.
                mode_label = st.session_state.get("scoring_mode", "Simulated")
                is_llm     = mode_label == "Real LLM"
                snap_key   = _effective_llm_key() if is_llm else ""
                snap_url   = _FORGE_LLM_BASE_URL
                snap_model = _FORGE_LLM_MODEL
                rubric_snap = rubric   # plain dict — safe to read across threads

                # Use user-selected concurrency limit, capped at the number of unscored items
                user_concurrency = int(st.session_state.get("bulk_concurrency", 5))
                max_workers = min(user_concurrency, len(unscored))
                n           = len(unscored)

                # ── Live display slots ──────────────────────────────────────────
                bulk_start_time     = time.time()
                _completion_times_: list[float] = []
                completed_ids_:     set[str] = set()

                # Per-card state: id → {"name", "state": Pending/Scoring/Done, "score"}
                card_states_: dict = {
                    sub["id"]: {"name": sub["name"], "state": "Pending", "score": 0}
                    for sub in unscored
                }

                prog_bar    = st.progress(0.0)
                dash_header = st.empty()
                dash_grid   = st.empty()
                log_slot    = st.empty()

                # Accumulates plain-data log rows; each entry is a dict
                log_entries_: list[dict] = []

                # ── Thread-safe in-flight tracking ──────────────────────────────
                # active_ids: IDs whose worker is currently executing (not queued)
                active_ids  = set()
                active_lock = threading.Lock()

                def _render_bulk_dashboard_():
                    """Re-render the full card grid and ETA header."""
                    elapsed = time.time() - bulk_start_time
                    n_done  = len(completed_ids_)
                    n_left  = n - n_done
                    if n_done > 0:
                        avg_t   = elapsed / n_done
                        eta_sec = int(avg_t * n_left)
                        if eta_sec < 60:
                            eta_str = f"{eta_sec}s remaining"
                        else:
                            m_part  = eta_sec // 60
                            s_part  = eta_sec % 60
                            eta_str = f"{m_part}m {s_part}s remaining"
                    else:
                        eta_str = "calculating…"

                    dash_header.markdown(
                        f'<div style="display:flex;align-items:center;gap:16px;padding:8px 0 2px;">'
                        f'<span style="font-size:12px;color:#8b949e;">'
                        f'<strong style="color:#e6edf3">{n_done}</strong> / {n} scored'
                        f'</span>'
                        f'<span style="font-size:12px;color:#30363d;">|</span>'
                        f'<span style="font-size:12px;color:#8b949e;">&#9201; {eta_str}</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

                    with active_lock:
                        scoring_ids = set(active_ids)

                    cards_html = ""
                    for cid, cs in card_states_.items():
                        raw_name  = cs["name"]
                        name_disp = (raw_name[:20] + "…") if len(raw_name) > 20 else raw_name
                        state     = "Scoring" if cid in scoring_ids else cs["state"]

                        if state == "Done":
                            sc_val = cs["score"]
                            color  = score_hex(sc_val)
                            badge  = (
                                f'<span style="font-size:10px;font-weight:600;color:{color};'
                                f'background:rgba(63,185,80,0.1);padding:2px 7px;border-radius:20px;">Done</span>'
                            )
                            score_disp = (
                                f'<div style="font-size:24px;font-weight:700;color:{color};'
                                f'line-height:1.1;margin-top:4px;">{sc_val}</div>'
                            )
                            card_bg = "rgba(63,185,80,0.04)"
                            border  = color
                        elif state == "Scoring":
                            badge  = (
                                '<span style="font-size:10px;font-weight:600;color:#58a6ff;'
                                'background:rgba(31,111,235,0.15);padding:2px 7px;border-radius:20px;">Scoring…</span>'
                            )
                            score_disp = '<div style="font-size:11px;color:#6e7681;margin-top:4px;">in progress</div>'
                            card_bg    = "rgba(31,111,235,0.05)"
                            border     = "#1f6feb"
                        else:
                            badge  = (
                                '<span style="font-size:10px;font-weight:600;color:#6e7681;'
                                'background:#21262d;padding:2px 7px;border-radius:20px;">Pending</span>'
                            )
                            score_disp = '<div style="font-size:11px;color:#484f58;margin-top:4px;">—</div>'
                            card_bg    = "#0d1117"
                            border     = "#30363d"

                        cards_html += (
                            f'<div style="background:{card_bg};border:1px solid {border};border-radius:8px;'
                            f'padding:10px 12px;min-height:80px;display:flex;flex-direction:column;'
                            f'justify-content:space-between;">'
                            f'<div style="font-size:12px;color:#e6edf3;font-weight:500;'
                            f'line-height:1.3;word-break:break-word;">{name_disp}</div>'
                            f'<div style="margin-top:6px;">{badge}{score_disp}</div>'
                            f'</div>'
                        )

                    dash_grid.markdown(
                        f'<div style="display:grid;grid-template-columns:repeat(auto-fill,minmax(150px,1fr));'
                        f'gap:8px;padding:4px 0 10px;">{cards_html}</div>',
                        unsafe_allow_html=True,
                    )

                def _render_log_():
                    """Re-render the scrollable progress log panel."""
                    if not log_entries_:
                        log_slot.empty()
                        return
                    rows_html = ""
                    for entry in reversed(log_entries_):
                        name_disp = (entry["name"][:28] + "…") if len(entry["name"]) > 28 else entry["name"]
                        sc_val    = entry["score"]
                        sc_color  = score_hex(sc_val)
                        status    = entry["status"]
                        if status == "rate-limited":
                            badge_color = "#f85149"
                            badge_bg    = "rgba(248,81,73,0.12)"
                            badge_label = "Rate-limited"
                        elif status == "fallback":
                            badge_color = "#d29922"
                            badge_bg    = "rgba(210,153,34,0.12)"
                            badge_label = "Fallback"
                        else:
                            badge_color = "#3fb950"
                            badge_bg    = "rgba(63,185,80,0.12)"
                            badge_label = "Scored"
                        rows_html += (
                            f'<div style="display:flex;align-items:center;gap:10px;'
                            f'padding:5px 10px;border-bottom:1px solid #21262d;font-size:12px;">'
                            f'<span style="color:#484f58;min-width:50px;">{entry["ts"]}</span>'
                            f'<span style="color:#e6edf3;flex:1;white-space:nowrap;overflow:hidden;'
                            f'text-overflow:ellipsis;">{name_disp}</span>'
                            f'<span style="color:{sc_color};font-weight:700;min-width:32px;'
                            f'text-align:right;">{sc_val}</span>'
                            f'<span style="font-size:10px;font-weight:600;color:{badge_color};'
                            f'background:{badge_bg};padding:1px 7px;border-radius:20px;'
                            f'min-width:76px;text-align:center;">{badge_label}</span>'
                            f'</div>'
                        )
                    n_done = len(log_entries_)
                    panel_html = (
                        f'<details open style="margin-top:8px;">'
                        f'<summary style="cursor:pointer;font-size:12px;font-weight:600;'
                        f'color:#8b949e;padding:4px 2px;user-select:none;">'
                        f'Progress Log &nbsp;<span style="color:#58a6ff">{n_done}</span> / {n} completed'
                        f'</summary>'
                        f'<div style="max-height:180px;overflow-y:auto;background:#0d1117;'
                        f'border:1px solid #21262d;border-radius:6px;margin-top:4px;">'
                        f'{rows_html}'
                        f'</div>'
                        f'</details>'
                    )
                    log_slot.markdown(panel_html, unsafe_allow_html=True)

                # Initial render — all Pending
                _render_bulk_dashboard_()

                # stop_event: set by the first worker that hits a 429/rate-limit;
                # subsequent workers check this before making their LLM call and
                # fall back to Simulated instead, preserving score quality semantics.
                stop_event  = threading.Event()

                def _score_worker_pure(sub, rubric_data, mode, api_key, base_url, model):
                    """
                    Pure worker: all config passed in — no st.session_state access.
                    Returns (sub_id, score_dict, warning_str_or_None, is_rate_limit).
                    """
                    sub_id = sub["id"]
                    with active_lock:
                        active_ids.add(sub_id)
                    try:
                        if mode == "Real LLM":
                            # Rate-limit gate: if another worker already hit 429,
                            # use Simulated immediately rather than hammering the API.
                            if stop_event.is_set():
                                sc = ai_score_submission(sub, rubric_data)
                                return sub_id, sc, None, False
                            if not api_key:
                                sc = ai_score_submission(sub, rubric_data)
                                return sub_id, sc, "No API key — used Simulated fallback.", False
                            try:
                                sc = _ai_score_llm_pure(sub, rubric_data, api_key, base_url, model)
                                return sub_id, sc, None, False
                            except RuntimeError as exc:
                                warn = str(exc)
                                is_rate = ("429" in warn or "rate" in warn.lower())
                                if is_rate:
                                    stop_event.set()   # signal all other workers
                                sc = ai_score_submission(sub, rubric_data)
                                return sub_id, sc, warn, is_rate
                        else:
                            time.sleep(0.3)   # brief stagger for Simulated visual effect
                            sc = ai_score_submission(sub, rubric_data)
                            return sub_id, sc, None, False
                    finally:
                        with active_lock:
                            active_ids.discard(sub_id)

                id_to_name       = {sub["id"]: sub["name"] for sub in unscored}
                llm_warns        = []
                completed        = 0
                rate_limit_hits  = 0

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    future_to_id = {
                        executor.submit(
                            _score_worker_pure,
                            sub, rubric_snap, mode_label, snap_key, snap_url, snap_model,
                        ): sub["id"]
                        for sub in unscored
                    }

                    for future in as_completed(future_to_id):
                        # Cancelled futures (from stop_event path) raise CancelledError
                        try:
                            sub_id, sc, warn, is_rate = future.result()
                        except Exception:
                            # Will be picked up in the sequential fallback below
                            continue

                        if is_rate:
                            rate_limit_hits += 1
                            # Cancel any futures still queued (not yet started)
                            for f in future_to_id:
                                if not f.running() and not f.done():
                                    f.cancel()

                        # ── Session state update — main thread only ────────────
                        idx = next(
                            j for j, s in enumerate(st.session_state.submissions)
                            if s["id"] == sub_id
                        )
                        st.session_state.submissions[idx].update({
                            "overall":     sc["overall"],
                            "innovation":  sc["innovation"],
                            "feasibility": sc["feasibility"],
                            "categories":  sc["categories"],
                            "auto_reject": sc["auto_reject"],
                            "high_risk":   sc["high_risk"],
                            "scored_at":   sc["scored_at"],
                            "status":      "Scored",
                        })
                        if warn:
                            llm_warns.append(warn)

                        completed += 1
                        prog_bar.progress(completed / n)
                        completed_ids_.add(sub_id)
                        card_states_[sub_id]["state"] = "Done"
                        card_states_[sub_id]["score"] = sc["overall"]
                        log_status = "rate-limited" if is_rate else ("fallback" if warn else "scored")
                        log_entries_.append({
                            "ts":     datetime.now().strftime("%H:%M:%S"),
                            "name":   id_to_name.get(sub_id, sub_id),
                            "score":  sc["overall"],
                            "status": log_status,
                        })
                        _render_bulk_dashboard_()
                        _render_log_()

                # ── Sequential fallback for any still-New after rate-limit ──────
                still_new = [s for s in st.session_state.submissions if s["status"] == "New"]
                if still_new:
                    for sub in still_new:
                        sc = ai_score_submission(sub, rubric_snap)
                        idx = next(
                            j for j, s in enumerate(st.session_state.submissions)
                            if s["id"] == sub["id"]
                        )
                        st.session_state.submissions[idx].update({
                            "overall":     sc["overall"],
                            "innovation":  sc["innovation"],
                            "feasibility": sc["feasibility"],
                            "categories":  sc["categories"],
                            "auto_reject": sc["auto_reject"],
                            "high_risk":   sc["high_risk"],
                            "scored_at":   sc["scored_at"],
                            "status":      "Scored",
                        })
                        completed += 1
                        prog_bar.progress(completed / n)
                        completed_ids_.add(sub["id"])
                        card_states_[sub["id"]]["state"] = "Done"
                        card_states_[sub["id"]]["score"] = sc["overall"]
                        log_entries_.append({
                            "ts":     datetime.now().strftime("%H:%M:%S"),
                            "name":   sub["name"],
                            "score":  sc["overall"],
                            "status": "fallback",
                        })
                        _render_bulk_dashboard_()
                        _render_log_()

                # ── Summary ────────────────────────────────────────────────────
                dash_header.empty()
                dash_grid.empty()
                log_slot.empty()
                prog_bar.progress(1.0)
                _close_criterion_detail_dialog()
                if rate_limit_hits:
                    st.warning(
                        f"{rate_limit_hits} API call(s) hit rate limits — those ideas were scored "
                        "with Simulated fallback. Remaining ideas were processed sequentially. "
                        "Try a smaller batch or wait before re-running.",
                        icon="⚠️",
                    )
                elif llm_warns:
                    st.warning(llm_warns[0], icon="⚠️")
                n_reject       = sum(1 for s in st.session_state.submissions if s.get("auto_reject"))
                parallel_note  = f" ({max_workers} concurrent)" if max_workers > 1 else ""
                st.success(
                    f"Scored {n} submission(s){parallel_note}. "
                    f"{n_reject} triggered auto-reject gates."
                )
                st.rerun()
            else:
                st.info("No unscored submissions.")

        # ── Filter ────────────────────────────────────────────────────────────
        visible = list(st.session_state.submissions)
        if search_q:
            visible = [s for s in visible if search_q.lower() in s["name"].lower() or search_q.lower() in s["id"].lower()]
        if f_status != "All Statuses":
            visible = [s for s in visible if s["status"] == f_status]
        if f_stage != "All Stages":
            visible = [s for s in visible if s["stage"] == f_stage]
        if f_sort == "Score (high first)":
            visible.sort(key=lambda x: x["overall"], reverse=True)
        elif f_sort == "Score (low first)":
            visible.sort(key=lambda x: x["overall"])
        elif f_sort == "Name":
            visible.sort(key=lambda x: x["name"])
        else:
            visible.sort(key=lambda x: x["submitted_at"], reverse=True)

        # ── Table ─────────────────────────────────────────────────────────────
        st.markdown('<div class="section-hd" style="margin-top:16px;">All Submissions</div>', unsafe_allow_html=True)
        _render_submissions_table_header()
        for sub in visible:
            _render_submission_table_row(sub, rubric)

    else:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">📂</div>
            <div class="empty-title">No submissions yet</div>
            <div class="empty-sub">Use the New Submission panel above to add an idea.</div>
        </div>""", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: SHORTLIST — Category folders
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Shortlist":

    _sl_total = _shortlist_total_count()
    st.markdown(f"""
    <div class="forge-topbar">
      <div class="forge-topbar-left">
        <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Shortlist</span></div>
        <div class="forge-page-tag">Saved Ideas</div>
      </div>
      <div class="forge-topbar-status">
        <div class="forge-status-dot"></div>
        {_sl_total} shortlisted
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="page-content">', unsafe_allow_html=True)

    if st.session_state.flash_msg:
        ftype, fmsg = st.session_state.flash_msg
        st.session_state.flash_msg = None
        if ftype == "success":
            st.success(fmsg)
        elif ftype == "info":
            st.info(fmsg)
        elif ftype == "warning":
            st.warning(fmsg)
        else:
            st.error(fmsg)

    _maybe_open_shortlist_dialog()
    _maybe_open_memo_dialog()
    _maybe_open_criterion_deep_dive_dialog()

    view_cat = st.session_state.get("shortlist_view_category")
    prev_sl_view = st.session_state.get("_tracked_sl_view")
    if prev_sl_view != view_cat:
        if prev_sl_view is not None:
            _close_criterion_detail_dialog()
        st.session_state._tracked_sl_view = view_cat

    if view_cat and view_cat in SHORTLIST_CATEGORIES:
        _render_shortlist_category_detail(view_cat, rubric)
    elif _sl_total == 0:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">⭐</div>
            <div class="empty-title">No ideas on your Shortlist yet</div>
            <div class="empty-sub">On the Submissions page, click <strong>⭐ Shortlist</strong><br>
            on any scored idea to save it into a category folder.</div>
        </div>""", unsafe_allow_html=True)
    else:
        _render_shortlist_folder_grid()

    st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: PIPELINE — Linear-style kanban
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Pipeline":

    _pipe_total = len(st.session_state.submissions)
    # ── Flash message ─────────────────────────────────────────────────────────
    if st.session_state.flash_msg:
        ftype, fmsg = st.session_state.flash_msg
        st.session_state.flash_msg = None
        if ftype == "success":
            st.success(fmsg)
        elif ftype == "info":
            st.info(fmsg)
        elif ftype == "warning":
            st.warning(fmsg)
        else:
            st.error(fmsg)

    subs = st.session_state.submissions
    detail_sub_id = st.session_state.get("pipeline_detail_sub_id")
    detail_sub = next((s for s in subs if s["id"] == detail_sub_id), None) if detail_sub_id else None
    # #region agent log
    _debug_log(
        "pipeline branch entered",
        {
            "page": page,
            "detail_sub_id": detail_sub_id,
            "detail_found": bool(detail_sub),
            "query_page": st.query_params.get("page"),
            "query_pipeline": st.query_params.get("pipeline"),
            "query_submission": st.query_params.get("submission"),
        },
        hypothesis_id="H3",
        run_id="post-fix",
        location="app.py:7904",
    )
    # #endregion

    _maybe_open_shortlist_dialog()
    _maybe_open_memo_dialog()
    _maybe_open_criterion_deep_dive_dialog()

    if detail_sub_id and detail_sub:
        _render_pipeline_detail_page(detail_sub, rubric)
    else:
        if detail_sub_id and not detail_sub:
            st.session_state.pipeline_detail_sub_id = None

        st.markdown(f"""
        <div class="forge-topbar">
          <div class="forge-topbar-left">
            <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Pipeline</span></div>
            <div class="forge-page-tag">7-Stage Flow</div>
          </div>
          <div class="forge-topbar-status">
            <div class="forge-status-dot"></div>
            {_pipe_total} in pipeline
          </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="page-content">', unsafe_allow_html=True)

        stage_map = {s["name"]: [] for s in STAGES}
        for sub in subs:
            if sub["stage"] in stage_map:
                stage_map[sub["stage"]].append(sub)

        # ── Kanban board ──────────────────────────────────────────────────────────
        st.markdown('<div class="section-hd">Board View</div>', unsafe_allow_html=True)

        _KBOARD_ICONS = {
            "Intake": "📥", "Concept": "💡", "Validation": "🔬",
            "Prototyping": "🛠", "Market Test": "📈", "Scaling": "⚡", "Monitoring": "📡",
        }
        cols = st.columns(len(STAGES))
        for stage, col in zip(STAGES, cols):
            items = stage_map.get(stage["name"], [])
            with col:
                _kb_icon = _KBOARD_ICONS.get(stage["name"], "●")
                _stage_name_e = _html.escape(_stage_display_name(stage["name"]))
                st.markdown(
                    f'<div class="kanban-col-header">'
                    f'<span class="kanban-col-name" style="color:{stage["color"]}">{_kb_icon} {_stage_name_e}</span>'
                    f'<span class="kanban-count">{len(items)}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                if items:
                    for sub in items:
                        score_badge_html = ""
                        if sub["overall"] > 0:
                            bc = score_badge_class(sub["overall"])
                            score_badge_html = f'<span class="badge-score {bc}" style="font-size:10px;">{sub["overall"]}</span>'
                        pc = pill_class(sub["status"])
                        blabel, bcolor, bbg = forge_badge(sub)
                        fb_html = ""
                        if blabel:
                            fb_html = (
                                f'<span style="font-size:9px;font-weight:700;color:{bcolor};'
                                f'background:{bbg};border:1px solid {bcolor}44;'
                                f'border-radius:3px;padding:1px 5px;margin-left:4px;">'
                                f'{blabel}</span>'
                            )

                        _card_name = _html.escape(sub["name"])
                        _card_id = _html.escape(sub["id"])
                        _card_status = _html.escape(sub["status"])
                        _card_stage = _html.escape(_stage_display_name(sub["stage"]))
                        st.markdown(
                            f'<div class="kanban-card" style="border-left-color:{stage["color"]};">'
                            f'<div class="kanban-card-title">{_card_name}</div>'
                            f'<div class="kanban-card-meta">'
                            f'<span class="forge-id" style="font-size:10px;">{_card_id}</span>'
                            f'{score_badge_html}{fb_html}'
                            f'</div>'
                            f'<div style="margin-top:6px;display:flex;align-items:center;justify-content:space-between;gap:8px;">'
                            f'<span class="pill {pc}" style="font-size:10px;">{_card_status}</span>'
                            f'<span style="font-size:10px;color:#8b949e;">{_card_stage}</span>'
                            f'</div>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )
                        if st.button(
                            "Open profile",
                            key=f"pipe_open_{sub['id']}",
                            use_container_width=True,
                            help=f"Open profile for {sub['name']}",
                        ):
                            _open_pipeline_detail(sub["id"])
                            st.rerun()

                        # ── Inline action buttons ───────────────────────────────
                        cur_idx = STAGE_NAMES.index(sub["stage"]) if sub["stage"] in STAGE_NAMES else -1
                        at_last = cur_idx >= len(STAGE_NAMES) - 1
                        btn_score, btn_adv = st.columns(2)
                        with btn_score:
                            if st.button("Score", key=f"pipe_sc_{sub['id']}", use_container_width=True):
                                mode_label = st.session_state.get("scoring_mode", "Simulated")
                                with st.spinner(f"Scoring using ForgeOS Extensive Rubric v2 ({mode_label})…"):
                                    if mode_label == "Simulated":
                                        time.sleep(0.8)
                                    sc2, warn2 = route_scoring(sub, rubric)
                                    idx = next(i for i, s in enumerate(st.session_state.submissions) if s["id"] == sub["id"])
                                    st.session_state.submissions[idx].update({
                                        "overall":     sc2["overall"],
                                        "innovation":  sc2["innovation"],
                                        "feasibility": sc2["feasibility"],
                                        "categories":  sc2["categories"],
                                        "auto_reject": sc2["auto_reject"],
                                        "high_risk":   sc2["high_risk"],
                                        "scored_at":   sc2["scored_at"],
                                        "status":      "Scored",
                                    })
                                    st.session_state.investment_memos.pop(sub["id"], None)
                                    _clear_criterion_detail_cache(sub["id"])
                                    gate_note = " · Auto-Reject gate triggered" if sc2["auto_reject"] else (" · High-Risk flag raised" if sc2["high_risk"] else "")
                                    fallback_note = f" · ⚠ {warn2}" if warn2 else ""
                                    st.session_state.flash_msg = ("success", f"Scored '{sub['name']}' — Overall: {sc2['overall']}/100{gate_note}{fallback_note}")
                                    st.rerun()
                        with btn_adv:
                            if st.button("Advance", key=f"pipe_adv_{sub['id']}", disabled=at_last, use_container_width=True):
                                new_stage = STAGE_NAMES[cur_idx + 1]
                                with st.spinner(f"Advancing '{sub['name']}' to {new_stage}…"):
                                    time.sleep(0.5)
                                    idx = next(i for i, s in enumerate(st.session_state.submissions) if s["id"] == sub["id"])
                                    summary = generate_stage_summary(st.session_state.submissions[idx], new_stage)
                                    hist = st.session_state.submissions[idx].get("stage_history", [])
                                    hist.append({"stage": new_stage, "moved_at": datetime.now().strftime("%Y-%m-%d")})
                                    st.session_state.submissions[idx]["stage"] = new_stage
                                    st.session_state.submissions[idx]["stage_summary"] = summary
                                    st.session_state.submissions[idx]["stage_history"] = hist
                                    st.session_state.flash_msg = ("info", f"'{sub['name']}' advanced to {new_stage} — AI stage brief generated")
                                st.rerun()
                else:
                    st.markdown('<div class="kanban-empty">No ideas</div>', unsafe_allow_html=True)

        # ── Stage distribution bar chart ──────────────────────────────────────────
        if subs:
            st.markdown('<div class="section-hd" style="margin-top:28px;">Stage Distribution</div>', unsafe_allow_html=True)
            names = [_stage_display_name(s["name"]) for s in STAGES]
            counts = [len(stage_map.get(s["name"], [])) for s in STAGES]
            colors = [s["color"] for s in STAGES]
            fig_bar = go.Figure(go.Bar(
                x=names, y=counts, marker_color=colors,
                text=counts, textposition="auto",
                textfont=dict(color="#e6edf3", size=11),
            ))
            fig_bar.update_layout(
                paper_bgcolor="#0f1e30", plot_bgcolor="#0f1e30",
                font={"color": "#4e6680", "family": "Inter", "size": 11},
                xaxis=dict(gridcolor="#1a2f49", color="#4e6680", showgrid=False),
                yaxis=dict(gridcolor="#1a2f49", color="#4e6680", title="Submissions", showgrid=True),
                height=240, margin=dict(l=0, r=0, t=8, b=0),
                showlegend=False, bargap=0.25,
            )
            st.plotly_chart(fig_bar, use_container_width=True)

            # ── Conversion funnel ──────────────────────────────────────────────
            st.markdown('<div class="section-hd">Conversion Funnel</div>', unsafe_allow_html=True)
            total_subs = max(len(subs), 1)
            for stage in STAGES:
                cnt = len(stage_map.get(stage["name"], []))
                pct = cnt / total_subs
                bar_w = max(int(pct * 100), 2)
                st.markdown(f"""
                <div style="display:flex;align-items:center;gap:10px;margin-bottom:8px;">
                  <span style="width:138px;font-size:11px;color:var(--text-3);text-align:right;flex-shrink:0;font-weight:500;">{_stage_display_name(stage['name'])}</span>
                  <div style="flex:1;background:var(--surface);border-radius:99px;height:16px;overflow:hidden;border:1px solid var(--border);">
                    <div style="width:{bar_w}%;background:{stage['color']};height:100%;border-radius:99px;
                                display:flex;align-items:center;padding-left:8px;min-width:22px;
                                box-shadow:0 0 8px {stage['color']}55;transition:width 0.4s ease;">
                      <span style="font-size:9px;color:rgba(0,0,0,0.75);font-weight:800;">{cnt if cnt else ''}</span>
                    </div>
                  </div>
                  <span style="width:36px;font-size:11px;color:var(--text-3);font-weight:600;font-variant-numeric:tabular-nums;">{int(pct*100)}%</span>
                </div>""", unsafe_allow_html=True)

        else:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-icon">🔀</div>
                <div class="empty-title">Pipeline is empty</div>
                <div class="empty-sub">Add submissions to see them flow through the pipeline.</div>
            </div>""", unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE: RUBRIC SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
elif page == "Rubric Settings":

    st.markdown("""
    <div class="forge-topbar">
      <div class="forge-topbar-left">
        <div class="forge-breadcrumb">ForgeOS <span class="forge-sep">/</span> <span>Rubric Settings</span></div>
        <div class="forge-page-tag">Scoring Config</div>
      </div>
      <div class="forge-topbar-status">
        <div class="forge-status-dot"></div>
        Extensive Rubric v2
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="page-content">', unsafe_allow_html=True)

    if not rubric:
        st.error("No rubric.json found.")
    else:
        criteria_list = rubric.get("criteria", [])
        total_w       = sum(c.get("weight", 0) for c in criteria_list)
        gating        = rubric.get("gating_rules", [])
        cat_palette   = ["#1f6feb","#238636","#9e6a03","#6e40c9","#d18000","#0ea5e9","#ec4899","#f85149"]

        # ── Rubric header ─────────────────────────────────────────────────────
        col_h1, col_h2, col_h3, col_h4 = st.columns(4)
        with col_h1:
            w_ok = abs(total_w - 100) < 1
            st.markdown(f"""
            <div class="forge-card">
              <div class="stat-label">Rubric</div>
              <div style="font-size:13px;font-weight:600;color:#e6edf3;margin:4px 0 2px;">{rubric.get('rubric_name','Innovation Rubric')[:30]}…</div>
              <div style="font-size:11px;color:#8b949e;">{len(criteria_list)} criteria</div>
            </div>""", unsafe_allow_html=True)
        with col_h2:
            st.markdown(f"""
            <div class="forge-card">
              <div class="stat-label">Total Weight</div>
              <div class="stat-value" style="color:{'#3fb950' if w_ok else '#f85149'}">{total_w}%</div>
              <div class="stat-sub">{'Balanced ✓' if w_ok else 'Imbalanced'}</div>
            </div>""", unsafe_allow_html=True)
        with col_h3:
            st.markdown(f"""
            <div class="forge-card">
              <div class="stat-label">Green Zone</div>
              <div class="stat-value" style="color:#3fb950">≥ {THRESHOLDS['green']}</div>
              <div class="stat-sub">High potential, fast-track</div>
            </div>""", unsafe_allow_html=True)
        with col_h4:
            st.markdown(f"""
            <div class="forge-card">
              <div class="stat-label">Auto-Reject Rules</div>
              <div class="stat-value" style="color:#f85149">{len(gating)}</div>
              <div class="stat-sub">Gating rules active</div>
            </div>""", unsafe_allow_html=True)

        # ── Scoring Mode ──────────────────────────────────────────────────────
        st.markdown('<div class="section-hd">Scoring Mode</div>', unsafe_allow_html=True)

        # Determine key source so UI can react without rerunning
        _key_from_env     = bool(_FORGE_LLM_API_KEY)
        _key_from_session = bool(st.session_state.get("session_llm_key", "").strip())
        _key_available    = _key_from_env or _key_from_session

        # ── In-app API key entry (shown only when env var is absent) ───────────
        if not _key_from_env:
            _key_on_disk = bool(_load_saved_key())

            st.markdown("""
            <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;
                        padding:14px 16px;margin-bottom:14px;">
              <div style="font-size:11px;font-weight:700;color:#8b949e;text-transform:uppercase;
                          letter-spacing:0.06em;margin-bottom:6px;">API Key</div>
              <div style="font-size:12px;color:#8b949e;margin-bottom:10px;">
                Paste an OpenAI-compatible API key to enable Real LLM scoring.
                Use <strong style="color:#e6edf3;">Save key</strong> to remember it across refreshes —
                it is stored in a local file on this server and never committed to version control.
                It is only sent to the configured LLM endpoint
                (<code style="color:#58a6ff;">api.x.ai/v1</code> by default).
              </div>
            </div>""", unsafe_allow_html=True)

            ki1, ki2 = st.columns([5, 1])
            with ki1:
                entered_key = st.text_input(
                    "API Key",
                    value=st.session_state.session_llm_key,
                    type="password",
                    placeholder="sk-… or xai-…",
                    label_visibility="collapsed",
                    help="OpenAI-compatible API key. Click 'Save key' to persist across refreshes.",
                )
            with ki2:
                save_clicked = st.button(
                    "Save key",
                    use_container_width=True,
                    disabled=not entered_key.strip(),
                    help="Write key to disk so it survives page refreshes.",
                )

            # Update session state when input changes
            if entered_key != st.session_state.session_llm_key:
                st.session_state.session_llm_key = entered_key.strip()
                _key_from_session = bool(st.session_state.session_llm_key)
                _key_available    = _key_from_session

            # Save to disk when button clicked
            if save_clicked and entered_key.strip():
                _save_key_to_disk(entered_key.strip())
                st.session_state.session_llm_key = entered_key.strip()
                _key_from_session = True
                _key_available    = True
                _key_on_disk      = True
                st.session_state.flash_msg = ("success", "API key saved — it will be remembered after page refresh.")
                st.rerun()

            # Status line + clear button
            if _key_from_session:
                key_preview = st.session_state.session_llm_key[:6] + "••••••••"
                persist_note = (
                    "saved to disk · survives refresh"
                    if _key_on_disk
                    else "session only · click Save key to persist"
                )
                st_row1, st_row2 = st.columns([4, 1])
                with st_row1:
                    st.markdown(
                        f'<div style="font-size:11px;color:#3fb950;margin-top:6px;">'
                        f'✓ Key active ({key_preview}) · {persist_note}</div>',
                        unsafe_allow_html=True,
                    )
                with st_row2:
                    if st.button("Clear", help="Remove key from memory and disk.", use_container_width=True):
                        _clear_saved_key()
                        st.session_state.session_llm_key = ""
                        st.session_state.scoring_mode    = "Simulated"
                        st.session_state.flash_msg = ("info", "API key cleared from memory and disk.")
                        st.rerun()
            else:
                st.markdown(
                    '<div style="font-size:11px;color:#6e7681;margin-top:4px;">'
                    'No key entered — Real LLM mode unavailable.</div>',
                    unsafe_allow_html=True,
                )

        llm_option_label = "Real LLM" if _key_available else "Real LLM (no key)"
        mode_options = ["Simulated", llm_option_label]
        current_mode = st.session_state.get("scoring_mode", "Simulated")
        current_idx  = 1 if current_mode == "Real LLM" else 0

        scol1, scol2 = st.columns([2, 3])
        with scol1:
            chosen = st.radio(
                "scoring_mode_radio",
                options=mode_options,
                index=current_idx,
                label_visibility="collapsed",
                help="Simulated uses keyword-weighted heuristics. Real LLM sends data to the configured AI endpoint.",
                disabled=not _key_available and current_mode != "Real LLM",
            )
            new_mode = "Real LLM" if chosen.startswith("Real LLM") and _key_available else "Simulated"
            if new_mode != current_mode:
                st.session_state.scoring_mode = new_mode
                st.rerun()
        with scol2:
            if new_mode == "Real LLM":
                key_note = "API key from environment variable." if _key_from_env else "API key entered for this session — not persisted."
                st.markdown(f"""
                <div style="background:#0d2b1a;border:1px solid #238636;border-radius:8px;padding:10px 14px;">
                  <div style="font-size:11px;font-weight:700;color:#3fb950;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px;">Real LLM Active</div>
                  <div style="font-size:12px;color:#b0b8c4;">Model: <code style="color:#58a6ff;">{_FORGE_LLM_MODEL}</code></div>
                  <div style="font-size:12px;color:#b0b8c4;">Endpoint: <code style="color:#58a6ff;">{_FORGE_LLM_BASE_URL}</code></div>
                  <div style="font-size:11px;color:#8b949e;margin-top:4px;">{key_note} Uploaded PDFs are extracted and sent as context.</div>
                </div>""", unsafe_allow_html=True)
            else:
                if not _key_available and not _key_from_env:
                    st.markdown("""
                    <div style="background:#1c1207;border:1px solid #9e6a03;border-radius:8px;padding:10px 14px;">
                      <div style="font-size:11px;font-weight:700;color:#d29922;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px;">Real LLM Unavailable</div>
                      <div style="font-size:12px;color:#8b949e;">Paste an API key above to unlock Real LLM scoring, or set the <code style="color:#58a6ff;">FORGE_LLM_API_KEY</code> environment variable for a persistent connection.</div>
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;padding:10px 14px;">
                      <div style="font-size:11px;font-weight:700;color:#8b949e;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px;">Simulated Mode</div>
                      <div style="font-size:12px;color:#8b949e;">Uses keyword signals and rubric-anchored heuristics. Reproducible — same idea always scores the same. No API key required.</div>
                    </div>""", unsafe_allow_html=True)

        # ── Tabs ──────────────────────────────────────────────────────────────
        tab_crit, tab_gate, tab_chart = st.tabs(["Criteria", "Gating Rules", "Weight Chart"])

        with tab_crit:
            for i, crit in enumerate(criteria_list):
                color    = cat_palette[i % len(cat_palette)]
                anchors  = crit.get("scoring_anchors", {})
                subs_f   = crit.get("sub_factors", [])
                redflags = crit.get("red_flags", [])
                evidence = crit.get("evidence_required", "")

                with st.expander(f"{crit['criterion']}  ·  {crit.get('weight', 0)}%"):
                    st.markdown(f'<p style="color:#8d96a3;font-size:12px;margin:0 0 12px 0;">{crit.get("description","")}</p>', unsafe_allow_html=True)

                    cl, cr = st.columns([3, 2])

                    with cl:
                        if anchors:
                            st.markdown('<div style="font-size:10px;text-transform:uppercase;letter-spacing:0.08em;color:#8b949e;font-weight:700;margin-bottom:8px;">Scoring Anchors</div>', unsafe_allow_html=True)
                            anchor_map = {"1-3": ("anchor-low","#f85149"), "4-6": ("anchor-mid","#d29922"), "7-10": ("anchor-high","#3fb950")}
                            for band, desc in anchors.items():
                                cls_n, clr = anchor_map.get(band, ("", "#8b949e"))
                                st.markdown(f"""
                                <div class="anchor-band {cls_n}">
                                  <span style="color:{clr};font-weight:700;font-size:11px;">{band}</span>
                                  <span style="color:#8d96a3;font-size:11px;margin-left:8px;">{desc}</span>
                                </div>""", unsafe_allow_html=True)

                        if evidence:
                            st.markdown(f"""
                            <div style="margin-top:12px;padding:8px 12px;background:#0d1117;border:1px solid #21262d;border-radius:6px;">
                              <div style="font-size:10px;text-transform:uppercase;letter-spacing:0.08em;color:#8b949e;font-weight:700;margin-bottom:4px;">Evidence Required</div>
                              <div style="font-size:12px;color:#8d96a3;">{evidence}</div>
                            </div>""", unsafe_allow_html=True)

                    with cr:
                        if subs_f:
                            st.markdown('<div style="font-size:10px;text-transform:uppercase;letter-spacing:0.08em;color:#8b949e;font-weight:700;margin-bottom:8px;">Sub-Factors</div>', unsafe_allow_html=True)
                            tags = "".join(f'<span class="subfactor-tag">▸ {sf}</span>' for sf in subs_f)
                            st.markdown(f'<div style="display:flex;flex-wrap:wrap;gap:4px;">{tags}</div>', unsafe_allow_html=True)

                        if redflags:
                            st.markdown('<div style="font-size:10px;text-transform:uppercase;letter-spacing:0.08em;color:#8b949e;font-weight:700;margin:12px 0 8px;">Red Flags</div>', unsafe_allow_html=True)
                            flags_html = "".join(f'<div class="redflag-item">⚑ {rf}</div>' for rf in redflags)
                            st.markdown(f'<div style="background:#2b0f0f18;border:1px solid #6e181822;border-radius:6px;padding:10px 12px;">{flags_html}</div>', unsafe_allow_html=True)

        with tab_gate:
            if gating:
                st.markdown('<div style="height:8px"></div>', unsafe_allow_html=True)
                for rule in gating:
                    st.markdown(f'<div class="gate-rule"><span style="color:#f85149;font-size:14px;">⛔</span> {rule}</div>', unsafe_allow_html=True)
            else:
                st.markdown('<p style="color:#8b949e;font-size:13px;">No gating rules defined.</p>', unsafe_allow_html=True)

        with tab_chart:
            names_c  = [c["criterion"] for c in criteria_list]
            weights_c = [c.get("weight", 0) for c in criteria_list]
            fig_w = go.Figure(go.Bar(
                x=names_c, y=weights_c,
                marker_color=[cat_palette[i % len(cat_palette)] for i in range(len(names_c))],
                text=[f"{w}%" for w in weights_c],
                textposition="auto",
                textfont=dict(color="#e6edf3", size=11),
            ))
            fig_w.update_layout(
                paper_bgcolor="#0f1e30", plot_bgcolor="#0f1e30",
                font={"color": "#4e6680", "family": "Inter", "size": 11},
                xaxis=dict(gridcolor="#1a2f49", color="#4e6680", tickangle=-30, showgrid=False),
                yaxis=dict(gridcolor="#1a2f49", color="#4e6680", title="Weight (%)", showgrid=True),
                height=300, margin=dict(l=0, r=0, t=8, b=80),
            )
            st.plotly_chart(fig_w, use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)
